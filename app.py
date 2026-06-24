import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
import sqlite3
import json
import os
from datetime import datetime, date, timedelta
from typing import Optional
import calendar

# ── Опциональные экспортные библиотеки ─────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ═══════════════════════════════════════════════════════════════════════════
# БАЗА ДАННЫХ
# ═══════════════════════════════════════════════════════════════════════════

DB_PATH = "it_cost.db"

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT
    );
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        last_name TEXT NOT NULL,
        first_name TEXT NOT NULL,
        middle_name TEXT,
        email TEXT,
        position TEXT,
        roles TEXT DEFAULT '[]',
        workspace_ids TEXT DEFAULT '[]',
        password TEXT DEFAULT 'password'
    );
    CREATE TABLE IF NOT EXISTS workspaces (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        subdomain TEXT,
        admin_user_id INTEGER,
        members TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        inn TEXT,
        type TEXT,
        name TEXT,
        director_name TEXT,
        email TEXT,
        phone TEXT
    );
    CREATE TABLE IF NOT EXISTS employees (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        last_name TEXT NOT NULL,
        first_name TEXT NOT NULL,
        middle_name TEXT,
        position TEXT,
        salary REAL DEFAULT 0,
        tax_rate REAL DEFAULT 30.2,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS contractors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        last_name TEXT NOT NULL,
        first_name TEXT NOT NULL,
        middle_name TEXT,
        contract_type TEXT DEFAULT 'ГПХ',
        tax_rate REAL DEFAULT 0,
        unit TEXT DEFAULT 'часы',
        rate REAL DEFAULT 0,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS subcontractors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        inn TEXT,
        type TEXT DEFAULT 'ООО',
        name TEXT NOT NULL,
        director_name TEXT,
        email TEXT,
        phone TEXT,
        unit TEXT DEFAULT 'часы',
        rate REAL DEFAULT 0,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS equipment (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        description TEXT,
        ownership TEXT DEFAULT 'Собственное',
        rental_cost REAL DEFAULT 0,
        unit TEXT DEFAULT 'часы',
        rate REAL DEFAULT 0,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        workspace_id INTEGER,
        name TEXT NOT NULL,
        date_start TEXT,
        date_end TEXT,
        description TEXT,
        customer_id INTEGER,
        tax_rate REAL DEFAULT 0,
        status TEXT DEFAULT 'Активный',
        resources TEXT DEFAULT '[]'
    );
    """)
    # Default settings
    defaults = {
        'company_name': 'Моя компания',
        'director_name': 'Иванов Иван Иванович',
        'director_position': 'Генеральный директор',
        'logo_path': '',
        'phone': '',
        'email': '',
    }
    for k, v in defaults.items():
        c.execute("INSERT OR IGNORE INTO settings(key,value) VALUES(?,?)", (k, v))
    conn.commit()
    conn.close()

def get_setting(key, default=''):
    conn = get_db()
    row = conn.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
    conn.close()
    return row['value'] if row else default

def set_setting(key, value):
    conn = get_db()
    conn.execute("INSERT OR REPLACE INTO settings(key,value) VALUES(?,?)", (key, value))
    conn.commit()
    conn.close()

# ═══════════════════════════════════════════════════════════════════════════
# СТИЛИ / ПАЛИТРА
# ═══════════════════════════════════════════════════════════════════════════

COLORS = {
    'bg':        '#0F1117',
    'surface':   '#1A1D27',
    'surface2':  '#232635',
    'border':    '#2E3147',
    'accent':    '#4F8EF7',
    'accent2':   '#7C5CFC',
    'success':   '#2ECC71',
    'warning':   '#F39C12',
    'danger':    '#E74C3C',
    'text':      '#E8EAF6',
    'text_dim':  '#7B82A8',
    'white':     '#FFFFFF',
    'hover':     '#2A2E45',
}

FONTS = {
    'title':   ('Segoe UI', 22, 'bold'),
    'heading': ('Segoe UI', 13, 'bold'),
    'normal':  ('Segoe UI', 10),
    'small':   ('Segoe UI', 9),
    'mono':    ('Consolas', 10),
    'big':     ('Segoe UI', 28, 'bold'),
}

def apply_ttk_styles():
    style = ttk.Style()
    style.theme_use('clam')
    bg, surf, acc, txt = COLORS['bg'], COLORS['surface'], COLORS['accent'], COLORS['text']
    dim = COLORS['text_dim']

    style.configure('TFrame', background=bg)
    style.configure('Surface.TFrame', background=surf)
    style.configure('Surface2.TFrame', background=COLORS['surface2'])

    style.configure('TLabel', background=bg, foreground=txt, font=FONTS['normal'])
    style.configure('Dim.TLabel', background=bg, foreground=dim, font=FONTS['small'])
    style.configure('Surface.TLabel', background=surf, foreground=txt, font=FONTS['normal'])
    style.configure('Heading.TLabel', background=bg, foreground=txt, font=FONTS['heading'])
    style.configure('Title.TLabel', background=bg, foreground=txt, font=FONTS['title'])
    style.configure('Accent.TLabel', background=bg, foreground=acc, font=FONTS['heading'])

    style.configure('TEntry', fieldbackground=COLORS['surface2'], background=COLORS['surface2'],
                    foreground=txt, insertcolor=txt, borderwidth=0, relief='flat', font=FONTS['normal'])
    style.map('TEntry', fieldbackground=[('focus', COLORS['surface2'])])

    style.configure('TCombobox', fieldbackground=COLORS['surface2'], background=COLORS['surface2'],
                    foreground=txt, selectbackground=acc, font=FONTS['normal'])
    style.map('TCombobox', fieldbackground=[('readonly', COLORS['surface2'])])

    style.configure('Treeview', background=COLORS['surface2'], fieldbackground=COLORS['surface2'],
                    foreground=txt, rowheight=32, font=FONTS['normal'], borderwidth=0)
    style.configure('Treeview.Heading', background=COLORS['surface'], foreground=dim,
                    font=('Segoe UI', 9, 'bold'), relief='flat')
    style.map('Treeview', background=[('selected', acc)], foreground=[('selected', COLORS['white'])])

    style.configure('TScrollbar', background=COLORS['surface2'], troughcolor=bg,
                    arrowcolor=dim, borderwidth=0)
    style.configure('TNotebook', background=bg, borderwidth=0)
    style.configure('TNotebook.Tab', background=COLORS['surface'], foreground=dim,
                    padding=[16, 8], font=FONTS['normal'])
    style.map('TNotebook.Tab', background=[('selected', bg)], foreground=[('selected', acc)])

    style.configure('TCheckbutton', background=bg, foreground=txt, font=FONTS['normal'])
    style.configure('Surface.TCheckbutton', background=surf, foreground=txt, font=FONTS['normal'])


# ═══════════════════════════════════════════════════════════════════════════
# ВСПОМОГАТЕЛЬНЫЕ ВИДЖЕТЫ
# ═══════════════════════════════════════════════════════════════════════════

class RoundedButton(tk.Canvas):
    def __init__(self, parent, text, command=None, color=None, text_color=None,
                 width=140, height=36, radius=8, **kw):
        # ttk widgets don't support '-bg'; fall back gracefully
        try:
            parent_bg = parent.cget('bg')
        except Exception:
            try:
                parent_bg = parent.cget('background')
            except Exception:
                parent_bg = COLORS['bg']
        super().__init__(parent, width=width, height=height,
                         bg=parent_bg,
                         highlightthickness=0, cursor='hand2', **kw)
        self.color = color or COLORS['accent']
        self.text_color = text_color or COLORS['white']
        self.text = text
        self.command = command
        self.w, self.h, self.r = width, height, radius
        self._draw()
        self.bind('<Button-1>', self._on_click)
        self.bind('<Enter>', self._on_enter)
        self.bind('<Leave>', self._on_leave)

    def _draw(self, hover=False):
        self.delete('all')
        c = self._lighten(self.color) if hover else self.color
        self._rounded_rect(1, 1, self.w-1, self.h-1, self.r, c)
        self.create_text(self.w//2, self.h//2, text=self.text,
                         fill=self.text_color, font=FONTS['normal'])

    def _rounded_rect(self, x1, y1, x2, y2, r, fill):
        self.create_arc(x1, y1, x1+2*r, y1+2*r, start=90, extent=90, fill=fill, outline=fill)
        self.create_arc(x2-2*r, y1, x2, y1+2*r, start=0, extent=90, fill=fill, outline=fill)
        self.create_arc(x1, y2-2*r, x1+2*r, y2, start=180, extent=90, fill=fill, outline=fill)
        self.create_arc(x2-2*r, y2-2*r, x2, y2, start=270, extent=90, fill=fill, outline=fill)
        self.create_rectangle(x1+r, y1, x2-r, y2, fill=fill, outline=fill)
        self.create_rectangle(x1, y1+r, x2, y2-r, fill=fill, outline=fill)

    def _lighten(self, hex_color):
        r = min(255, int(hex_color[1:3], 16) + 20)
        g = min(255, int(hex_color[3:5], 16) + 20)
        b = min(255, int(hex_color[5:7], 16) + 20)
        return f'#{r:02x}{g:02x}{b:02x}'

    def _on_click(self, e):
        if self.command: self.command()

    def _on_enter(self, e): self._draw(hover=True)
    def _on_leave(self, e): self._draw(hover=False)


class FormField(ttk.Frame):
    """Labeled entry with dark styling."""
    def __init__(self, parent, label, var=None, width=280, **kw):
        super().__init__(parent, style='Surface.TFrame', **kw)
        ttk.Label(self, text=label, style='Dim.TLabel',
                  background=COLORS['surface']).pack(anchor='w', pady=(0, 3))
        self.var = var or tk.StringVar()
        e = ttk.Entry(self, textvariable=self.var, width=width//8)
        e.pack(fill='x', ipady=5)
        self._style_entry(e)

    def _style_entry(self, e):
        e.configure(style='TEntry')

    def get(self): return self.var.get()
    def set(self, v): self.var.set(v)


class ScrollableFrame(ttk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        canvas = tk.Canvas(self, bg=COLORS['bg'], highlightthickness=0)
        sb = ttk.Scrollbar(self, orient='vertical', command=canvas.yview)
        self.inner = ttk.Frame(canvas)
        self.inner.bind('<Configure>', lambda e: canvas.configure(
            scrollregion=canvas.bbox('all')))
        canvas.create_window((0, 0), window=self.inner, anchor='nw')
        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        canvas.bind('<MouseWheel>', lambda e: canvas.yview_scroll(-1*(e.delta//120), 'units'))


def card_frame(parent, **kw):
    """Surface card with border."""
    f = tk.Frame(parent, bg=COLORS['surface'], bd=0, highlightthickness=1,
                 highlightbackground=COLORS['border'], **kw)
    return f


def section_label(parent, text):
    lf = tk.Frame(parent, bg=COLORS['bg'])
    lf.pack(fill='x', pady=(18, 6))
    tk.Label(lf, text=text, bg=COLORS['bg'], fg=COLORS['accent'],
             font=FONTS['heading']).pack(side='left')
    tk.Frame(lf, bg=COLORS['border'], height=1).pack(side='left', fill='x', expand=True, padx=(10, 0))


def dark_entry(parent, textvariable=None, width=20, **kw):
    e = tk.Entry(parent, textvariable=textvariable, width=width,
                 bg=COLORS['surface2'], fg=COLORS['text'], insertbackground=COLORS['text'],
                 relief='flat', bd=0, font=FONTS['normal'], highlightthickness=1,
                 highlightbackground=COLORS['border'], highlightcolor=COLORS['accent'],
                 disabledbackground=COLORS['surface2'], disabledforeground=COLORS['text_dim'], **kw)
    return e


def dark_combo(parent, values, textvariable=None, width=18, **kw):
    cb = ttk.Combobox(parent, values=values, textvariable=textvariable,
                      width=width, state='readonly', **kw)
    return cb


# ═══════════════════════════════════════════════════════════════════════════
# ДИАЛОГ — БАЗОВЫЙ
# ═══════════════════════════════════════════════════════════════════════════

class BaseDialog(tk.Toplevel):
    def __init__(self, parent, title, width=560, height=600):
        super().__init__(parent)
        self.title(title)
        self.configure(bg=COLORS['bg'])
        self.geometry(f'{width}x{height}')
        self.resizable(True, True)
        self.grab_set()
        self.result = None
        self._build()

    def _build(self):
        pass  # Override

    def _ok(self):
        if self._validate():
            self.result = self._collect()
            self.destroy()

    def _validate(self): return True
    def _collect(self): return {}

    def _footer(self, parent):
        f = tk.Frame(parent, bg=COLORS['bg'])
        f.pack(fill='x', pady=12, padx=16)
        RoundedButton(f, 'Отмена', command=self.destroy,
                      color=COLORS['surface2'], width=110).pack(side='right', padx=(8, 0))
        RoundedButton(f, 'Сохранить', command=self._ok,
                      color=COLORS['accent'], width=120).pack(side='right')
        return f


# ═══════════════════════════════════════════════════════════════════════════
# БАЗОВЫЙ РЕЕСТР (TABLE + CRUD)
# ═══════════════════════════════════════════════════════════════════════════

class RegistryFrame(ttk.Frame):
    """Generic CRUD table panel."""
    COLUMNS = []   # list of (key, display, width)
    TITLE = 'Реестр'

    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self.configure(style='TFrame')
        self._build()
        self.load()

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg=COLORS['bg'])
        hdr.pack(fill='x', padx=20, pady=(16, 8))
        tk.Label(hdr, text=self.TITLE, bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(side='left')
        btn_f = tk.Frame(hdr, bg=COLORS['bg'])
        btn_f.pack(side='right')
        RoundedButton(btn_f, '＋ Добавить', command=self._add,
                      color=COLORS['accent'], width=130).pack(side='left', padx=4)
        RoundedButton(btn_f, '✎ Изменить', command=self._edit,
                      color=COLORS['surface2'], width=130).pack(side='left', padx=4)
        RoundedButton(btn_f, '✕ Удалить', command=self._delete,
                      color=COLORS['danger'], width=120).pack(side='left', padx=4)

        # Search
        sf = tk.Frame(self, bg=COLORS['bg'])
        sf.pack(fill='x', padx=20, pady=(0, 8))
        tk.Label(sf, text='🔍', bg=COLORS['bg'], fg=COLORS['text_dim'],
                 font=FONTS['normal']).pack(side='left', padx=(0, 6))
        self._search_var = tk.StringVar()
        self._search_var.trace_add('write', lambda *_: self.load())
        dark_entry(sf, textvariable=self._search_var, width=40).pack(side='left')

        # Tree
        tf = tk.Frame(self, bg=COLORS['bg'])
        tf.pack(fill='both', expand=True, padx=20, pady=(0, 16))
        cols = [c[0] for c in self.COLUMNS]
        self.tree = ttk.Treeview(tf, columns=cols, show='headings', selectmode='browse')
        for key, disp, w in self.COLUMNS:
            self.tree.heading(key, text=disp)
            self.tree.column(key, width=w, minwidth=60)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self.tree.bind('<Double-1>', lambda e: self._edit())

    def load(self):
        q = self._search_var.get().lower() if hasattr(self, '_search_var') else ''
        self.tree.delete(*self.tree.get_children())
        for row in self._fetch_rows():
            vals = [str(row[c[0]] or '') for c in self.COLUMNS]
            if q and not any(q in v.lower() for v in vals):
                continue
            self.tree.insert('', 'end', iid=str(row['id']), values=vals)

    def _fetch_rows(self): return []
    def _add(self): pass
    def _edit(self): pass

    def _delete(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите запись для удаления')
            return
        if messagebox.askyesno('Удаление', 'Удалить выбранную запись?'):
            self._do_delete(int(sel[0]))
            self.load()

    def _do_delete(self, row_id): pass

    def _selected_id(self):
        sel = self.tree.selection()
        return int(sel[0]) if sel else None


# ═══════════════════════════════════════════════════════════════════════════
# РЕЕСТР СОТРУДНИКОВ
# ═══════════════════════════════════════════════════════════════════════════

class EmployeeDialog(BaseDialog):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Сотрудник', width=500, height=500)

    def _build(self):
        tk.Label(self, text='Сотрудник', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, var, row_n):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            e = dark_entry(f, textvariable=var, width=32)
            e.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_last = tk.StringVar(value=self.data.get('last_name', ''))
        self.v_first = tk.StringVar(value=self.data.get('first_name', ''))
        self.v_mid = tk.StringVar(value=self.data.get('middle_name', ''))
        self.v_pos = tk.StringVar(value=self.data.get('position', ''))
        self.v_sal = tk.StringVar(value=str(self.data.get('salary', '')))
        self.v_tax = tk.StringVar(value=str(self.data.get('tax_rate', '30.2')))

        row('Фамилия *', self.v_last, 0)
        row('Имя *', self.v_first, 1)
        row('Отчество', self.v_mid, 2)
        row('Должность', self.v_pos, 3)
        row('Оклад (руб/мес) *', self.v_sal, 4)
        row('Налоговая ставка % *', self.v_tax, 5)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_last.get() or not self.v_first.get():
            messagebox.showwarning('Ошибка', 'Заполните Фамилию и Имя'); return False
        try:
            float(self.v_sal.get()); float(self.v_tax.get())
        except ValueError:
            messagebox.showwarning('Ошибка', 'Оклад и ставка — числа'); return False
        return True

    def _collect(self):
        return dict(last_name=self.v_last.get(), first_name=self.v_first.get(),
                    middle_name=self.v_mid.get(), position=self.v_pos.get(),
                    salary=float(self.v_sal.get()), tax_rate=float(self.v_tax.get()))


class EmployeesFrame(RegistryFrame):
    TITLE = 'Реестр сотрудников'
    COLUMNS = [('id', '#', 40), ('fio', 'ФИО', 200), ('position', 'Должность', 160),
               ('salary', 'Оклад', 100), ('tax_rate', 'Нал. ставка %', 100)]

    def _fetch_rows(self):
        conn = get_db()
        rows = conn.execute("SELECT *, last_name||' '||first_name||COALESCE(' '||middle_name,'') as fio FROM employees ORDER BY last_name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = EmployeeDialog(self)
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("INSERT INTO employees(last_name,first_name,middle_name,position,salary,tax_rate) VALUES(?,?,?,?,?,?)",
                         (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                          d.result['position'], d.result['salary'], d.result['tax_rate']))
            conn.commit(); conn.close(); self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid: messagebox.showwarning('Выбор', 'Выберите запись'); return
        conn = get_db()
        row = conn.execute("SELECT * FROM employees WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = EmployeeDialog(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("UPDATE employees SET last_name=?,first_name=?,middle_name=?,position=?,salary=?,tax_rate=? WHERE id=?",
                         (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                          d.result['position'], d.result['salary'], d.result['tax_rate'], rid))
            conn.commit(); conn.close(); self.load()

    def _do_delete(self, rid):
        conn = get_db()
        conn.execute("DELETE FROM employees WHERE id=?", (rid,))
        conn.commit(); conn.close()


# ═══════════════════════════════════════════════════════════════════════════
# РЕЕСТР ИСПОЛНИТЕЛЕЙ (ФИЗИЧЕСКИХ ЛИЦ)
# ═══════════════════════════════════════════════════════════════════════════

class ContractorDialog(BaseDialog):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Исполнитель', width=500, height=520)

    def _build(self):
        tk.Label(self, text='Исполнитель (физ. лицо)', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_last = tk.StringVar(value=self.data.get('last_name', ''))
        self.v_first = tk.StringVar(value=self.data.get('first_name', ''))
        self.v_mid = tk.StringVar(value=self.data.get('middle_name', ''))
        self.v_ctype = tk.StringVar(value=self.data.get('contract_type', 'ГПХ'))
        self.v_tax = tk.StringVar(value=str(self.data.get('tax_rate', '13')))
        self.v_unit = tk.StringVar(value=self.data.get('unit', 'часы'))
        self.v_rate = tk.StringVar(value=str(self.data.get('rate', '')))

        row('Фамилия *', dark_entry(f, textvariable=self.v_last, width=28), 0)
        row('Имя *', dark_entry(f, textvariable=self.v_first, width=28), 1)
        row('Отчество', dark_entry(f, textvariable=self.v_mid, width=28), 2)

        cb_type = dark_combo(f, ['ГПХ', 'НПД'], textvariable=self.v_ctype, width=12)
        cb_type.bind('<<ComboboxSelected>>', self._on_type)
        row('Тип оформления *', cb_type, 3)

        self.tax_entry = dark_entry(f, textvariable=self.v_tax, width=12)
        row('Налоговая ставка %', self.tax_entry, 4)
        row('Единица', dark_combo(f, ['часы', 'дни', 'полная стоимость'],
                                  textvariable=self.v_unit, width=16), 5)
        row('Ставка за ед.', dark_entry(f, textvariable=self.v_rate, width=16), 6)
        f.columnconfigure(1, weight=1)
        self._on_type()
        self._footer(self)

    def _on_type(self, e=None):
        if self.v_ctype.get() == 'НПД':
            self.v_tax.set('0')
            self.tax_entry.config(state='disabled')
        else:
            self.tax_entry.config(state='normal')

    def _validate(self):
        if not self.v_last.get() or not self.v_first.get():
            messagebox.showwarning('Ошибка', 'Заполните Фамилию и Имя'); return False
        return True

    def _collect(self):
        return dict(last_name=self.v_last.get(), first_name=self.v_first.get(),
                    middle_name=self.v_mid.get(), contract_type=self.v_ctype.get(),
                    tax_rate=float(self.v_tax.get() or 0),
                    unit=self.v_unit.get(), rate=float(self.v_rate.get() or 0))


class ContractorsFrame(RegistryFrame):
    TITLE = 'Реестр исполнителей (физ. лица)'
    COLUMNS = [('id', '#', 40), ('fio', 'ФИО', 200), ('contract_type', 'Тип', 80),
               ('unit', 'Ед.', 80), ('rate', 'Ставка', 100), ('tax_rate', 'Нал. %', 80)]

    def _fetch_rows(self):
        conn = get_db()
        rows = conn.execute("SELECT *, last_name||' '||first_name||COALESCE(' '||middle_name,'') as fio FROM contractors ORDER BY last_name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = ContractorDialog(self)
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("INSERT INTO contractors(last_name,first_name,middle_name,contract_type,tax_rate,unit,rate) VALUES(?,?,?,?,?,?,?)",
                         tuple(d.result[k] for k in ['last_name','first_name','middle_name','contract_type','tax_rate','unit','rate']))
            conn.commit(); conn.close(); self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid: messagebox.showwarning('Выбор', 'Выберите запись'); return
        conn = get_db()
        row = conn.execute("SELECT * FROM contractors WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = ContractorDialog(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("UPDATE contractors SET last_name=?,first_name=?,middle_name=?,contract_type=?,tax_rate=?,unit=?,rate=? WHERE id=?",
                         (*[d.result[k] for k in ['last_name','first_name','middle_name','contract_type','tax_rate','unit','rate']], rid))
            conn.commit(); conn.close(); self.load()

    def _do_delete(self, rid):
        conn = get_db()
        conn.execute("DELETE FROM contractors WHERE id=?", (rid,))
        conn.commit(); conn.close()


# ═══════════════════════════════════════════════════════════════════════════
# РЕЕСТР СУБПОДРЯДЧИКОВ
# ═══════════════════════════════════════════════════════════════════════════

class SubcontractorDialog(BaseDialog):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Субподрядчик', width=500, height=560)

    def _build(self):
        tk.Label(self, text='Субподрядчик (юр. лицо / ИП)', bg=COLORS['bg'],
                 fg=COLORS['text'], font=FONTS['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_inn = tk.StringVar(value=self.data.get('inn', ''))
        self.v_type = tk.StringVar(value=self.data.get('type', 'ООО'))
        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_dir = tk.StringVar(value=self.data.get('director_name', ''))
        self.v_email = tk.StringVar(value=self.data.get('email', ''))
        self.v_phone = tk.StringVar(value=self.data.get('phone', ''))
        self.v_unit = tk.StringVar(value=self.data.get('unit', 'часы'))
        self.v_rate = tk.StringVar(value=str(self.data.get('rate', '')))

        row('ИНН', dark_entry(f, textvariable=self.v_inn, width=20), 0)
        row('Тип', dark_combo(f, ['ООО', 'АО', 'ИП', 'ПАО'], textvariable=self.v_type, width=12), 1)
        row('Название *', dark_entry(f, textvariable=self.v_name, width=28), 2)
        row('Руководитель', dark_entry(f, textvariable=self.v_dir, width=28), 3)
        row('Email', dark_entry(f, textvariable=self.v_email, width=28), 4)
        row('Телефон', dark_entry(f, textvariable=self.v_phone, width=20), 5)
        row('Единица', dark_combo(f, ['часы', 'дни', 'полная стоимость'],
                                  textvariable=self.v_unit, width=16), 6)
        row('Ставка за ед.', dark_entry(f, textvariable=self.v_rate, width=16), 7)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Заполните Название'); return False
        return True

    def _collect(self):
        return dict(inn=self.v_inn.get(), type=self.v_type.get(), name=self.v_name.get(),
                    director_name=self.v_dir.get(), email=self.v_email.get(),
                    phone=self.v_phone.get(), unit=self.v_unit.get(),
                    rate=float(self.v_rate.get() or 0))


class SubcontractorsFrame(RegistryFrame):
    TITLE = 'Реестр субподрядчиков (юр. лица / ИП)'
    COLUMNS = [('id', '#', 40), ('type', 'Тип', 60), ('name', 'Название', 220),
               ('director_name', 'Руководитель', 160), ('phone', 'Телефон', 120)]

    def _fetch_rows(self):
        conn = get_db()
        rows = conn.execute("SELECT * FROM subcontractors ORDER BY name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = SubcontractorDialog(self)
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("INSERT INTO subcontractors(inn,type,name,director_name,email,phone,unit,rate) VALUES(?,?,?,?,?,?,?,?)",
                         tuple(d.result[k] for k in ['inn','type','name','director_name','email','phone','unit','rate']))
            conn.commit(); conn.close(); self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid: messagebox.showwarning('Выбор', 'Выберите запись'); return
        conn = get_db()
        row = conn.execute("SELECT * FROM subcontractors WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = SubcontractorDialog(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("UPDATE subcontractors SET inn=?,type=?,name=?,director_name=?,email=?,phone=?,unit=?,rate=? WHERE id=?",
                         (*[d.result[k] for k in ['inn','type','name','director_name','email','phone','unit','rate']], rid))
            conn.commit(); conn.close(); self.load()

    def _do_delete(self, rid):
        conn = get_db()
        conn.execute("DELETE FROM subcontractors WHERE id=?", (rid,))
        conn.commit(); conn.close()


# ═══════════════════════════════════════════════════════════════════════════
# РЕЕСТР ОБОРУДОВАНИЯ
# ═══════════════════════════════════════════════════════════════════════════

class EquipmentDialog(BaseDialog):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Оборудование', width=500, height=480)

    def _build(self):
        tk.Label(self, text='Оборудование', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_desc = tk.StringVar(value=self.data.get('description', ''))
        self.v_own = tk.StringVar(value=self.data.get('ownership', 'Собственное'))
        self.v_rent = tk.StringVar(value=str(self.data.get('rental_cost', '0')))
        self.v_unit = tk.StringVar(value=self.data.get('unit', 'часы'))
        self.v_rate = tk.StringVar(value=str(self.data.get('rate', '')))

        row('Название *', dark_entry(f, textvariable=self.v_name, width=28), 0)
        row('Описание', dark_entry(f, textvariable=self.v_desc, width=28), 1)
        row('Тип владения', dark_combo(f, ['Собственное', 'В аренде'],
                                       textvariable=self.v_own, width=16), 2)
        row('Эксплуат. стоимость', dark_entry(f, textvariable=self.v_rent, width=16), 3)
        row('Единица', dark_combo(f, ['часы', 'дни', 'полная стоимость'],
                                  textvariable=self.v_unit, width=16), 4)
        row('Ставка за ед.', dark_entry(f, textvariable=self.v_rate, width=16), 5)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Введите Название'); return False
        return True

    def _collect(self):
        return dict(name=self.v_name.get(), description=self.v_desc.get(),
                    ownership=self.v_own.get(), rental_cost=float(self.v_rent.get() or 0),
                    unit=self.v_unit.get(), rate=float(self.v_rate.get() or 0))


class EquipmentFrame(RegistryFrame):
    TITLE = 'Реестр оборудования'
    COLUMNS = [('id', '#', 40), ('name', 'Название', 200), ('ownership', 'Тип', 100),
               ('unit', 'Ед.', 80), ('rate', 'Ставка', 100)]

    def _fetch_rows(self):
        conn = get_db()
        rows = conn.execute("SELECT * FROM equipment ORDER BY name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = EquipmentDialog(self)
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("INSERT INTO equipment(name,description,ownership,rental_cost,unit,rate) VALUES(?,?,?,?,?,?)",
                         tuple(d.result[k] for k in ['name','description','ownership','rental_cost','unit','rate']))
            conn.commit(); conn.close(); self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid: messagebox.showwarning('Выбор', 'Выберите запись'); return
        conn = get_db()
        row = conn.execute("SELECT * FROM equipment WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = EquipmentDialog(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("UPDATE equipment SET name=?,description=?,ownership=?,rental_cost=?,unit=?,rate=? WHERE id=?",
                         (*[d.result[k] for k in ['name','description','ownership','rental_cost','unit','rate']], rid))
            conn.commit(); conn.close(); self.load()

    def _do_delete(self, rid):
        conn = get_db()
        conn.execute("DELETE FROM equipment WHERE id=?", (rid,))
        conn.commit(); conn.close()


# ═══════════════════════════════════════════════════════════════════════════
# РЕЕСТР ЗАКАЗЧИКОВ
# ═══════════════════════════════════════════════════════════════════════════

class CustomerDialog(BaseDialog):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Заказчик', width=500, height=500)

    def _build(self):
        tk.Label(self, text='Заказчик', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_inn = tk.StringVar(value=self.data.get('inn', ''))
        self.v_type = tk.StringVar(value=self.data.get('type', 'Юридическое лицо'))
        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_dir = tk.StringVar(value=self.data.get('director_name', ''))
        self.v_email = tk.StringVar(value=self.data.get('email', ''))
        self.v_phone = tk.StringVar(value=self.data.get('phone', ''))

        row('ИНН', dark_entry(f, textvariable=self.v_inn, width=20), 0)
        row('Тип', dark_combo(f, ['Физическое лицо', 'Индивидуальный предприниматель', 'Юридическое лицо'],
                              textvariable=self.v_type, width=28), 1)
        row('Название / ФИО *', dark_entry(f, textvariable=self.v_name, width=28), 2)
        row('ФИО руководителя', dark_entry(f, textvariable=self.v_dir, width=28), 3)
        row('Email', dark_entry(f, textvariable=self.v_email, width=28), 4)
        row('Телефон', dark_entry(f, textvariable=self.v_phone, width=20), 5)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Введите Название / ФИО'); return False
        return True

    def _collect(self):
        return dict(inn=self.v_inn.get(), type=self.v_type.get(), name=self.v_name.get(),
                    director_name=self.v_dir.get(), email=self.v_email.get(), phone=self.v_phone.get())


class CustomersFrame(RegistryFrame):
    TITLE = 'Реестр заказчиков'
    COLUMNS = [('id', '#', 40), ('type', 'Тип', 120), ('name', 'Название / ФИО', 200),
               ('director_name', 'Руководитель', 160), ('email', 'Email', 160)]

    def _fetch_rows(self):
        conn = get_db()
        rows = conn.execute("SELECT * FROM customers ORDER BY name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = CustomerDialog(self)
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("INSERT INTO customers(inn,type,name,director_name,email,phone) VALUES(?,?,?,?,?,?)",
                         tuple(d.result[k] for k in ['inn','type','name','director_name','email','phone']))
            conn.commit(); conn.close(); self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid: messagebox.showwarning('Выбор', 'Выберите запись'); return
        conn = get_db()
        row = conn.execute("SELECT * FROM customers WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = CustomerDialog(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute("UPDATE customers SET inn=?,type=?,name=?,director_name=?,email=?,phone=? WHERE id=?",
                         (*[d.result[k] for k in ['inn','type','name','director_name','email','phone']], rid))
            conn.commit(); conn.close(); self.load()

    def _do_delete(self, rid):
        conn = get_db()
        conn.execute("DELETE FROM customers WHERE id=?", (rid,))
        conn.commit(); conn.close()


# ═══════════════════════════════════════════════════════════════════════════
# УПРАВЛЕНИЕ ПРОЕКТОМ — РЕСУРСЫ
# ═══════════════════════════════════════════════════════════════════════════

def _workdays_in_month(yr, mo):
    """Количество рабочих дней (без выходных) в указанном месяце."""
    _, days_in_month = calendar.monthrange(yr, mo)
    return sum(1 for d in range(1, days_in_month + 1)
               if calendar.weekday(yr, mo, d) < 5)


def workdays_count(date_start, date_end):
    """Считает количество рабочих дней (Пн-Пт) в интервале [date_start, date_end].
    Принимает строки 'YYYY-MM-DD' или объекты date. Возвращает int."""
    if not date_start or not date_end:
        return 0
    try:
        d1 = (datetime.strptime(date_start, '%Y-%m-%d').date()
              if isinstance(date_start, str) else date_start)
        d2 = (datetime.strptime(date_end, '%Y-%m-%d').date()
              if isinstance(date_end, str) else date_end)
    except Exception:
        return 0
    if d2 < d1:
        return 0
    days = (d2 - d1).days + 1
    full_weeks, rem = divmod(days, 7)
    count = full_weeks * 5
    for i in range(rem):
        d = d1 + timedelta(days=full_weeks * 7 + i)
        if d.weekday() < 5:
            count += 1
    return count


def calc_employee_cost(emp, units, date_start, date_end):
    """Расчёт себестоимости работы сотрудника по формуле из ТЗ.

    Точный (помесячный) вариант — основной, используется при наличии интервала дат:
        sr = Σ_по_месяцам cwd_m * ((zp + zp*ns) / cwdm_m)
    где cwd_m — отработанные (рабочие) дни сотрудника в данном месяце интервала,
    cwdm_m — общее количество рабочих дней в этом месяце (без выходных).
    Стоимость дня в каждом месяце своя, т.к. зарплата делится на разное число
    рабочих дней по месяцам (в "коротких" месяцах день стоит дороже).

    Если интервал дат не указан — используется явный fallback из ТЗ
    ("облегчённый вариант"): sr = cwd * ((zp + zp*ns) / cwd), где cwd берётся
    из количества единиц (`units`), введённого пользователем вручную.
    Согласно формуле в ТЗ cwd встречается и в числителе, и в знаменателе —
    поэтому при использовании этого fallback результат не зависит от
    количества дней и равен (zp + zp*ns) целиком; так и реализовано, без
    произвольных констант вроде "/22", которых в ТЗ нет.
    """
    zp = emp['salary']
    ns = emp['tax_rate'] / 100

    if not date_start or not date_end:
        # "Облегчённый вариант" из ТЗ буквально: sr = cwd * ((zp+zp*ns)/cwd)
        # Тождественно равно (zp + zp*ns) независимо от cwd.
        if units and units > 0:
            return zp + zp * ns
        return 0.0

    try:
        d1 = datetime.strptime(date_start, '%Y-%m-%d').date()
        d2 = datetime.strptime(date_end, '%Y-%m-%d').date()
    except Exception:
        if units and units > 0:
            return zp + zp * ns
        return 0.0

    if d2 < d1:
        return 0.0

    total = 0.0
    cur = d1
    while cur <= d2:
        yr, mo = cur.year, cur.month
        cwdm = _workdays_in_month(yr, mo)  # рабочих дней в месяце всего
        if cwdm == 0:
            cwdm = 1  # защита от деления на 0 (теоретически невозможно)
        day_rate = (zp + zp * ns) / cwdm
        _, days_in_month = calendar.monthrange(yr, mo)
        mo_end = date(yr, mo, days_in_month)
        seg_end = min(d2, mo_end)
        # отработанные (рабочие) дни сотрудника в пределах этого месяца
        cwd = sum(1 for d in range(cur.day, seg_end.day + 1)
                  if calendar.weekday(yr, mo, d) < 5)
        total += day_rate * cwd
        if mo == 12:
            cur = date(yr + 1, 1, 1)
        else:
            cur = date(yr, mo + 1, 1)
        if cur > d2:
            break
    return total


def calc_contractor_cost(con, units):
    cez = units
    suz = con['rate']
    ns = con['tax_rate'] / 100
    if con['contract_type'] == 'НПД':
        return cez * suz
    return cez * suz + (cez * suz) * ns


def calc_equipment_cost(eq, units):
    return eq['rate'] * units


def calc_subcontractor_cost(sub, units):
    return sub['rate'] * units


class ResourceDialog(BaseDialog):
    """Добавление/редактирование ресурса в проекте."""
    def __init__(self, parent, project_dates=None, data=None):
        self.project_dates = project_dates or ('', '')
        self.data = data or {}
        self._emp_map = {}; self._con_map = {}; self._sub_map = {}; self._eq_map = {}
        super().__init__(parent, 'Ресурс проекта', width=560, height=620)

    def _build(self):
        tk.Label(self, text='Ресурс проекта', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(16, 8))
        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='w', pady=5)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=5)

        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_rtype = tk.StringVar(value=self.data.get('resource_type', 'Сотрудник'))
        self.v_exec = tk.StringVar(value='')
        self.v_service = tk.StringVar(value=self.data.get('service', ''))
        self.v_ds = tk.StringVar(value=self.data.get('date_start', self.project_dates[0]))
        self.v_de = tk.StringVar(value=self.data.get('date_end', self.project_dates[1]))
        self.v_units = tk.StringVar(value=str(self.data.get('units', '1')))
        self.v_margin = tk.StringVar(value=str(self.data.get('margin_pct', '0')))

        row('Название в документах *', dark_entry(f, textvariable=self.v_name, width=30), 0)

        self.cb_type = dark_combo(f, ['Сотрудник', 'Исполнитель', 'Субподрядчик', 'Оборудование'],
                                  textvariable=self.v_rtype, width=20)
        self.cb_type.bind('<<ComboboxSelected>>', self._on_type)
        row('Тип ресурса *', self.cb_type, 1)

        self.cb_exec = dark_combo(f, [], textvariable=self.v_exec, width=26)
        row('Исполнитель *', self.cb_exec, 2)

        row('Название услуги', dark_entry(f, textvariable=self.v_service, width=30), 3)
        self.e_ds = dark_entry(f, textvariable=self.v_ds, width=14)
        row('Дата начала', self.e_ds, 4)
        self.e_de = dark_entry(f, textvariable=self.v_de, width=14)
        row('Дата окончания', self.e_de, 5)
        self.e_units = dark_entry(f, textvariable=self.v_units, width=10)
        self.lbl_units_row = tk.Label(f, text='Кол-во единиц (часы/дни)', bg=COLORS['bg'],
                                      fg=COLORS['text_dim'], font=FONTS['small'])
        self.lbl_units_row.grid(row=6, column=0, sticky='w', pady=5)
        self.e_units.grid(row=6, column=1, sticky='ew', padx=(10, 0), pady=5)
        self.lbl_units_hint = tk.Label(f, text='', bg=COLORS['bg'],
                                       fg=COLORS['text_dim'], font=FONTS['small'])
        self.lbl_units_hint.grid(row=6, column=1, sticky='e', padx=(0, 4))
        self._can_margin = Session.can_see_margin()
        self.lbl_margin_row = tk.Label(f, text='Маржинальность %', bg=COLORS['bg'],
                                       fg=COLORS['text_dim'], font=FONTS['small'])
        self.e_margin = dark_entry(f, textvariable=self.v_margin, width=10)
        if self._can_margin:
            self.lbl_margin_row.grid(row=7, column=0, sticky='w', pady=5)
            self.e_margin.grid(row=7, column=1, sticky='ew', padx=(10, 0), pady=5)
        # Если прав нет — поле просто не размещается на форме (не grid),
        # значение margin_pct из существующих данных сохраняется как есть
        # при сохранении ресурса, но недоступно для просмотра/изменения.

        # Cost preview
        self.lbl_cost = tk.Label(f, text='Себестоимость: —', bg=COLORS['bg'],
                                 fg=COLORS['success'], font=FONTS['heading'])
        self.lbl_cost.grid(row=8, column=0, columnspan=2, sticky='w', pady=(10, 0))

        f.columnconfigure(1, weight=1)
        self._load_executor_list()

        # Restore executor selection
        if self.data.get('executor_id'):
            self.v_exec.set(str(self.data.get('executor_display', '')))

        self.v_units.trace_add('write', lambda *_: self._update_cost())
        self.v_margin.trace_add('write', lambda *_: self._update_cost())
        self.v_ds.trace_add('write', lambda *_: self._sync_units())
        self.v_de.trace_add('write', lambda *_: self._sync_units())
        self.cb_exec.bind('<<ComboboxSelected>>', lambda e: self._update_cost())
        self._sync_units()
        self._update_cost()
        self._footer(self)

    def _on_type(self, e=None):
        self._load_executor_list()
        self._sync_units()
        self._update_cost()

    def _sync_units(self):
        """Для типа 'Сотрудник' количество единиц = кол-во рабочих дней из
        интервала и поле блокируется для ручного ввода (см. ТЗ: стоимость
        дня работы умножается на количество рабочих дней из указанного в
        управлении проектом интервала — интервал, а не ручной ввод, задаёт
        количество дней)."""
        if self.v_rtype.get() == 'Сотрудник':
            wd = workdays_count(self.v_ds.get(), self.v_de.get())
            self.v_units.set(str(wd))
            self.e_units.config(state='disabled')
            self.lbl_units_row.config(text='Кол-во рабочих дней (из интервала)')
            self.lbl_units_hint.config(text='авто' if wd else '')
        else:
            self.e_units.config(state='normal')
            self.lbl_units_row.config(text='Кол-во единиц (часы/дни)')
            self.lbl_units_hint.config(text='')

    def _load_executor_list(self):
        rtype = self.v_rtype.get()
        conn = get_db()
        names = []
        if rtype == 'Сотрудник':
            rows = conn.execute("SELECT id, last_name||' '||first_name as n FROM employees").fetchall()
            self._emp_map = {r['n']: dict(r) for r in rows}
            names = list(self._emp_map.keys())
        elif rtype == 'Исполнитель':
            rows = conn.execute("SELECT id, last_name||' '||first_name as n, contract_type, tax_rate, rate, unit FROM contractors").fetchall()
            self._con_map = {r['n']: dict(r) for r in rows}
            names = list(self._con_map.keys())
        elif rtype == 'Субподрядчик':
            rows = conn.execute("SELECT id, name as n, rate, unit FROM subcontractors").fetchall()
            self._sub_map = {r['n']: dict(r) for r in rows}
            names = list(self._sub_map.keys())
        elif rtype == 'Оборудование':
            rows = conn.execute("SELECT id, name as n, rate, unit FROM equipment").fetchall()
            self._eq_map = {r['n']: dict(r) for r in rows}
            names = list(self._eq_map.keys())
        conn.close()
        self.cb_exec['values'] = names
        if names: self.cb_exec.current(0)

    def _get_cost(self):
        try:
            units = float(self.v_units.get() or 0)
        except ValueError:
            return 0.0
        rtype = self.v_rtype.get()
        name = self.v_exec.get()
        if rtype == 'Сотрудник' and name in self._emp_map:
            conn = get_db()
            emp = conn.execute("SELECT * FROM employees WHERE id=?",
                               (self._emp_map[name]['id'],)).fetchone()
            conn.close()
            return calc_employee_cost(dict(emp), units, self.v_ds.get(), self.v_de.get())
        elif rtype == 'Исполнитель' and name in self._con_map:
            conn = get_db()
            con = conn.execute("SELECT * FROM contractors WHERE id=?",
                               (self._con_map[name]['id'],)).fetchone()
            conn.close()
            return calc_contractor_cost(dict(con), units)
        elif rtype == 'Субподрядчик' and name in self._sub_map:
            conn = get_db()
            sub = conn.execute("SELECT * FROM subcontractors WHERE id=?",
                               (self._sub_map[name]['id'],)).fetchone()
            conn.close()
            return calc_subcontractor_cost(dict(sub), units)
        elif rtype == 'Оборудование' and name in self._eq_map:
            conn = get_db()
            eq = conn.execute("SELECT * FROM equipment WHERE id=?",
                              (self._eq_map[name]['id'],)).fetchone()
            conn.close()
            return calc_equipment_cost(dict(eq), units)
        return 0.0

    def _update_cost(self):
        try:
            cost = self._get_cost()
            if self._can_margin:
                try:
                    mp = float(self.v_margin.get() or 0) / 100
                except ValueError:
                    mp = 0.0
                total = cost * (1 + mp)
                self.lbl_cost.config(
                    text=f'Себестоимость: {cost:,.2f} ₽   |   Итого (с маржой): {total:,.2f} ₽')
            else:
                self.lbl_cost.config(text=f'Себестоимость: {cost:,.2f} ₽')
        except Exception:
            pass

    def _get_exec_id(self):
        rtype = self.v_rtype.get()
        name = self.v_exec.get()
        m = {'Сотрудник': self._emp_map, 'Исполнитель': self._con_map,
             'Субподрядчик': self._sub_map, 'Оборудование': self._eq_map}
        mapping = m.get(rtype, {})
        return mapping.get(name, {}).get('id')

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Введите название ресурса'); return False
        if not self.v_exec.get():
            messagebox.showwarning('Ошибка', 'Выберите исполнителя'); return False
        if self.v_rtype.get() == 'Сотрудник' and (not self.v_ds.get() or not self.v_de.get()):
            messagebox.showwarning(
                'Ошибка',
                'Для типа «Сотрудник» укажите дату начала и окончания —\n'
                'количество дней и стоимость рассчитываются по интервалу.')
            return False
        return True

    def _collect(self):
        cost = self._get_cost()
        try:
            mp = float(self.v_margin.get() or 0)
        except ValueError:
            mp = 0.0
        total = cost * (1 + mp / 100)
        rtype = self.v_rtype.get()
        if rtype == 'Сотрудник':
            unit_label = 'дни'
        else:
            m = {'Исполнитель': self._con_map, 'Субподрядчик': self._sub_map,
                 'Оборудование': self._eq_map}.get(rtype, {})
            unit_label = m.get(self.v_exec.get(), {}).get('unit', '')
        return dict(
            name=self.v_name.get(),
            resource_type=rtype,
            executor_id=self._get_exec_id(),
            executor_display=self.v_exec.get(),
            service=self.v_service.get(),
            date_start=self.v_ds.get(),
            date_end=self.v_de.get(),
            units=float(self.v_units.get() or 1),
            unit_label=unit_label,
            margin_pct=mp,
            cost=cost,
            total=total,
        )


# ═══════════════════════════════════════════════════════════════════════════
# СТРАНИЦА ПРОЕКТА
# ═══════════════════════════════════════════════════════════════════════════

class ProjectPage(tk.Toplevel):
    def __init__(self, parent, project_id=None, workspace_id=None, on_save=None):
        super().__init__(parent)
        self.project_id = project_id
        self.workspace_id = workspace_id
        self.on_save = on_save
        self.resources = []
        self.configure(bg=COLORS['bg'])
        self.title('Управление проектом')
        self.geometry('960x780')
        self._load_project()
        self._build()
        self.grab_set()

    def _load_project(self):
        if self.project_id:
            conn = get_db()
            row = conn.execute("SELECT * FROM projects WHERE id=?", (self.project_id,)).fetchone()
            conn.close()
            self.proj = dict(row)
            self.resources = json.loads(self.proj.get('resources', '[]'))
        else:
            self.proj = {'name': '', 'date_start': '', 'date_end': '',
                         'description': '', 'customer_id': None, 'tax_rate': 0, 'status': 'Активный'}

    def _build(self):
        # Notebook tabs
        nb = ttk.Notebook(self)
        nb.pack(fill='both', expand=True, padx=0, pady=0)

        tab1 = ttk.Frame(nb)
        tab2 = ttk.Frame(nb)
        nb.add(tab1, text='  Реквизиты проекта  ')
        nb.add(tab2, text='  Управление ресурсами  ')

        self._build_details(tab1)
        self._build_resources(tab2)
        self._build_bottom()

    def _build_details(self, parent):
        sf = ScrollableFrame(parent)
        sf.pack(fill='both', expand=True)
        f = sf.inner
        f.configure(style='TFrame')

        pad = {'padx': 20, 'pady': 6}

        section_label(f, 'Основная информация')

        row_f = ttk.Frame(f)
        row_f.pack(fill='x', **pad)
        ttk.Label(row_f, text='Название *').pack(side='left', padx=(0, 8))
        self.v_name = tk.StringVar(value=self.proj.get('name', ''))
        dark_entry(f if False else row_f, textvariable=self.v_name, width=44).pack(side='left')

        row2 = ttk.Frame(f)
        row2.pack(fill='x', **pad)
        ttk.Label(row2, text='Начало').pack(side='left', padx=(0, 8))
        self.v_ds = tk.StringVar(value=self.proj.get('date_start', ''))
        dark_entry(row2, textvariable=self.v_ds, width=14).pack(side='left', padx=(0, 20))
        ttk.Label(row2, text='Окончание').pack(side='left', padx=(0, 8))
        self.v_de = tk.StringVar(value=self.proj.get('date_end', ''))
        dark_entry(row2, textvariable=self.v_de, width=14).pack(side='left')

        row3 = ttk.Frame(f)
        row3.pack(fill='x', **pad)
        ttk.Label(row3, text='Статус').pack(side='left', padx=(0, 8))
        self.v_status = tk.StringVar(value=self.proj.get('status', 'Активный'))
        dark_combo(row3, ['Активный', 'Завершён', 'Приостановлен'],
                   textvariable=self.v_status, width=18).pack(side='left')

        section_label(f, 'Описание')
        desc_f = tk.Frame(f, bg=COLORS['bg'])
        desc_f.pack(fill='x', padx=20, pady=4)
        self.txt_desc = tk.Text(desc_f, height=4, bg=COLORS['surface2'], fg=COLORS['text'],
                                font=FONTS['normal'], relief='flat', bd=0, insertbackground=COLORS['text'])
        self.txt_desc.insert('1.0', self.proj.get('description', ''))
        self.txt_desc.pack(fill='x')

        section_label(f, 'Финансы')

        row4 = ttk.Frame(f)
        row4.pack(fill='x', **pad)
        ttk.Label(row4, text='Заказчик').pack(side='left', padx=(0, 8))

        conn = get_db()
        cust_rows = conn.execute("SELECT id, name FROM customers ORDER BY name").fetchall()
        conn.close()
        self._cust_map = {r['name']: r['id'] for r in cust_rows}
        cust_names = ['(не выбран)'] + list(self._cust_map.keys())
        self.v_cust = tk.StringVar()
        if self.proj.get('customer_id'):
            conn = get_db()
            cr = conn.execute("SELECT name FROM customers WHERE id=?",
                              (self.proj['customer_id'],)).fetchone()
            conn.close()
            if cr: self.v_cust.set(cr['name'])
        if not self.v_cust.get(): self.v_cust.set('(не выбран)')
        dark_combo(row4, cust_names, textvariable=self.v_cust, width=30).pack(side='left')

        row5 = ttk.Frame(f)
        row5.pack(fill='x', **pad)
        ttk.Label(row5, text='Налоговая ставка %').pack(side='left', padx=(0, 8))
        self.v_tax = tk.StringVar(value=str(self.proj.get('tax_rate', '0')))
        dark_entry(row5, textvariable=self.v_tax, width=8).pack(side='left')

        section_label(f, 'Итоги (авторасчёт)')
        self.lbl_summary = tk.Label(f, text='', bg=COLORS['bg'], fg=COLORS['text'],
                                    font=FONTS['normal'], justify='left')
        self.lbl_summary.pack(anchor='w', padx=20, pady=4)
        self._update_summary()

    def _build_resources(self, parent):
        hdr = tk.Frame(parent, bg=COLORS['bg'])
        hdr.pack(fill='x', padx=20, pady=(16, 8))
        tk.Label(hdr, text='Ресурсы проекта', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['heading']).pack(side='left')
        RoundedButton(hdr, '＋ Ресурс', command=self._add_resource,
                      color=COLORS['accent'], width=120).pack(side='right', padx=4)
        RoundedButton(hdr, '✕ Удалить', command=self._del_resource,
                      color=COLORS['danger'], width=110).pack(side='right', padx=4)
        RoundedButton(hdr, '✎ Изменить', command=self._edit_resource,
                      color=COLORS['surface2'], width=120).pack(side='right', padx=4)

        can_margin = Session.can_see_margin()
        if can_margin:
            cols = ('name', 'type', 'executor', 'service', 'units', 'cost', 'margin', 'total')
            hdrs = ('Название', 'Тип', 'Исполнитель', 'Услуга', 'Ед.', 'Себест.', 'Маржа %', 'Итого')
            widths = (160, 100, 140, 120, 60, 100, 80, 110)
        else:
            cols = ('name', 'type', 'executor', 'service', 'units', 'cost')
            hdrs = ('Название', 'Тип', 'Исполнитель', 'Услуга', 'Ед.', 'Себест.')
            widths = (160, 100, 140, 120, 60, 100)
        tf = tk.Frame(parent, bg=COLORS['bg'])
        tf.pack(fill='both', expand=True, padx=20, pady=(0, 10))
        self.res_tree = ttk.Treeview(tf, columns=cols, show='headings')
        for c, h, w in zip(cols, hdrs, widths):
            self.res_tree.heading(c, text=h)
            self.res_tree.column(c, width=w, minwidth=50)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.res_tree.yview)
        self.res_tree.configure(yscrollcommand=sb.set)
        self.res_tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self.res_tree.bind('<Double-1>', lambda e: self._edit_resource())
        self._reload_res_tree()

    def _reload_res_tree(self):
        self.res_tree.delete(*self.res_tree.get_children())
        can_margin = Session.can_see_margin()
        for i, r in enumerate(self.resources):
            base_vals = [r.get('name', ''), r.get('resource_type', ''), r.get('executor_display', ''),
                         r.get('service', ''), r.get('units', ''), f"{r.get('cost', 0):,.2f}"]
            if can_margin:
                base_vals += [f"{r.get('margin_pct', 0):.1f}%", f"{r.get('total', 0):,.2f}"]
            self.res_tree.insert('', 'end', iid=str(i), values=tuple(base_vals))
        self._update_summary()

    def _add_resource(self):
        d = ResourceDialog(self, project_dates=(self.v_ds.get(), self.v_de.get()))
        self.wait_window(d)
        if d.result:
            self.resources.append(d.result)
            self._reload_res_tree()

    def _edit_resource(self):
        sel = self.res_tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите ресурс'); return
        idx = int(sel[0])
        d = ResourceDialog(self, project_dates=(self.v_ds.get(), self.v_de.get()),
                           data=self.resources[idx])
        self.wait_window(d)
        if d.result:
            self.resources[idx] = d.result
            self._reload_res_tree()

    def _del_resource(self):
        sel = self.res_tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите ресурс'); return
        idx = int(sel[0])
        if messagebox.askyesno('Удаление', 'Удалить ресурс?'):
            self.resources.pop(idx)
            self._reload_res_tree()

    def _update_summary(self):
        cost_sum = sum(r.get('cost', 0) for r in self.resources)
        total_sum = sum(r.get('total', 0) for r in self.resources)
        try:
            ns = float(getattr(self, 'v_tax', tk.StringVar()).get() or 0) / 100
        except Exception:
            ns = 0.0
        if Session.can_see_margin():
            final = total_sum * (1 + ns)
            profit = total_sum - cost_sum
            txt = (f'Себестоимость:          {cost_sum:>15,.2f} ₽\n'
                   f'Стоимость с маржой:     {total_sum:>15,.2f} ₽\n'
                   f'Налог ({ns*100:.0f}%):              {total_sum*ns:>15,.2f} ₽\n'
                   f'Итоговая стоимость:     {final:>15,.2f} ₽\n'
                   f'Чистая прибыль:         {profit:>15,.2f} ₽')
        else:
            # Без прав на маржу не показываем ни итоговую стоимость (она
            # включает маржу), ни прибыль — только себестоимость и налог
            # на неё, чтобы не раскрывать наценку даже косвенно.
            txt = (f'Себестоимость:          {cost_sum:>15,.2f} ₽\n'
                   f'Налог на себестоимость ({ns*100:.0f}%): {cost_sum*ns:>15,.2f} ₽\n'
                   f'(стоимость с маржой и прибыль видны только руководителю)')
        if hasattr(self, 'lbl_summary'):
            self.lbl_summary.config(text=txt, font=FONTS['mono'])

    def _build_bottom(self):
        bf = tk.Frame(self, bg=COLORS['surface'], pady=10)
        bf.pack(fill='x', side='bottom')
        inner = tk.Frame(bf, bg=COLORS['surface'])
        inner.pack(padx=20)
        RoundedButton(inner, '💾 Сохранить', command=self._save,
                      color=COLORS['accent'], width=150).pack(side='left', padx=6)
        RoundedButton(inner, '📄 НМА', command=self._export_nma,
                      color=COLORS['surface2'], width=130).pack(side='left', padx=6)
        RoundedButton(inner, '📋 КП', command=self._export_kp,
                      color=COLORS['surface2'], width=130).pack(side='left', padx=6)
        RoundedButton(inner, '✕ Закрыть', command=self.destroy,
                      color=COLORS['danger'], width=120).pack(side='right', padx=6)

    def _save(self):
        name = self.v_name.get().strip()
        if not name:
            messagebox.showwarning('Ошибка', 'Введите название проекта'); return
        cid = self._cust_map.get(self.v_cust.get())
        conn = get_db()
        if self.project_id:
            conn.execute("""UPDATE projects SET name=?,date_start=?,date_end=?,description=?,
                            customer_id=?,tax_rate=?,status=?,resources=? WHERE id=?""",
                         (name, self.v_ds.get(), self.v_de.get(),
                          self.txt_desc.get('1.0', 'end').strip(),
                          cid, float(self.v_tax.get() or 0), self.v_status.get(),
                          json.dumps(self.resources, ensure_ascii=False), self.project_id))
        else:
            conn.execute("""INSERT INTO projects(workspace_id,name,date_start,date_end,description,
                            customer_id,tax_rate,status,resources) VALUES(?,?,?,?,?,?,?,?,?)""",
                         (self.workspace_id, name, self.v_ds.get(), self.v_de.get(),
                          self.txt_desc.get('1.0', 'end').strip(),
                          cid, float(self.v_tax.get() or 0), self.v_status.get(),
                          json.dumps(self.resources, ensure_ascii=False)))
        conn.commit(); conn.close()
        messagebox.showinfo('Готово', 'Проект сохранён')
        if self.on_save: self.on_save()

    # ── Export helpers ─────────────────────────────────────────────────────
    def _choose_export_format(self, title):
        d = tk.Toplevel(self)
        d.title(title)
        d.configure(bg=COLORS['bg'])
        d.geometry('300x180')
        d.grab_set()
        result = tk.StringVar(value='')
        tk.Label(d, text='Формат экспорта', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['heading']).pack(pady=(20, 10))
        bf = tk.Frame(d, bg=COLORS['bg'])
        bf.pack()
        for fmt, label, avail in [('pdf', 'PDF', HAS_PDF), ('xlsx', 'Excel', HAS_XLSX), ('docx', 'Word', HAS_DOCX)]:
            state = 'normal' if avail else 'disabled'
            b = tk.Button(bf, text=label, width=8, bg=COLORS['surface2'], fg=COLORS['text'],
                          relief='flat', font=FONTS['normal'], state=state,
                          command=lambda f=fmt: (result.set(f), d.destroy()))
            b.pack(side='left', padx=6)
        d.wait_window()
        return result.get()

    def _export_nma(self):
        fmt = self._choose_export_format('Экспорт НМА')
        if not fmt: return
        name = self.v_name.get() or 'Проект'
        resources = self.resources
        cost_sum = sum(r.get('cost', 0) for r in resources)
        path = filedialog.asksaveasfilename(
            defaultextension=f'.{fmt}',
            filetypes=[(fmt.upper(), f'*.{fmt}')],
            initialfile=f'НМА_{name}.{fmt}')
        if not path: return
        if fmt == 'xlsx':
            export_nma_xlsx(path, name, self.v_ds.get(), self.v_de.get(), cost_sum, resources)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'docx':
            export_nma_docx(path, name, self.v_ds.get(), self.v_de.get(), cost_sum, resources)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'pdf':
            export_nma_pdf(path, name, self.v_ds.get(), self.v_de.get(), cost_sum, resources)

    def _export_kp(self):
        if self.v_cust.get() == '(не выбран)' or not self.v_cust.get():
            messagebox.showwarning('КП', 'Выберите заказчика для формирования КП'); return
        fmt = self._choose_export_format('Экспорт КП')
        if not fmt: return
        name = self.v_name.get() or 'Проект'
        cid = self._cust_map.get(self.v_cust.get())
        conn = get_db()
        cust = dict(conn.execute("SELECT * FROM customers WHERE id=?", (cid,)).fetchone()) if cid else {}
        conn.close()
        total_sum = sum(r.get('total', 0) for r in self.resources)
        try:
            ns = float(self.v_tax.get() or 0) / 100
        except Exception:
            ns = 0.0
        final = total_sum * (1 + ns)
        path = filedialog.asksaveasfilename(
            defaultextension=f'.{fmt}',
            filetypes=[(fmt.upper(), f'*.{fmt}')],
            initialfile=f'КП_{name}.{fmt}')
        if not path: return
        company = get_setting('company_name')
        director = get_setting('director_name')
        dir_pos = get_setting('director_position')
        if fmt == 'xlsx':
            export_kp_xlsx(path, name, cust, self.v_ds.get(), self.v_de.get(),
                           final, self.resources, company, director, dir_pos)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'docx':
            export_kp_docx(path, name, cust, self.v_ds.get(), self.v_de.get(),
                           final, self.resources, company, director, dir_pos)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'pdf':
            export_kp_pdf(path, name, cust, self.v_ds.get(), self.v_de.get(),
                          final, self.resources, company, director, dir_pos)


# ═══════════════════════════════════════════════════════════════════════════
# РАБОЧАЯ ОБЛАСТЬ
# ═══════════════════════════════════════════════════════════════════════════

class WorkspaceFrame(ttk.Frame):
    def __init__(self, parent, workspace_id, **kw):
        super().__init__(parent, **kw)
        self.workspace_id = workspace_id
        self.configure(style='TFrame')
        self._build()
        self.load()

    def _build(self):
        conn = get_db()
        ws = conn.execute("SELECT * FROM workspaces WHERE id=?", (self.workspace_id,)).fetchone()
        conn.close()
        ws_name = ws['name'] if ws else '—'

        hdr = tk.Frame(self, bg=COLORS['bg'])
        hdr.pack(fill='x', padx=20, pady=(16, 8))
        tk.Label(hdr, text=f'📁  {ws_name}', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(side='left')
        btn_f = tk.Frame(hdr, bg=COLORS['bg'])
        btn_f.pack(side='right')
        RoundedButton(btn_f, '＋ Проект', command=self._add_project,
                      color=COLORS['accent'], width=120).pack(side='left', padx=4)
        RoundedButton(btn_f, '✎ Открыть', command=self._open_project,
                      color=COLORS['surface2'], width=120).pack(side='left', padx=4)
        RoundedButton(btn_f, '✕ Удалить', command=self._delete_project,
                      color=COLORS['danger'], width=110).pack(side='left', padx=4)
        RoundedButton(btn_f, '👥 Участники', command=self._manage_members,
                      color=COLORS['accent2'], width=130).pack(side='left', padx=4)

        cols = ('id', 'name', 'status', 'date_start', 'date_end', 'cost')
        hdrs = ('#', 'Название проекта', 'Статус', 'Начало', 'Окончание', 'Итоговая стоимость')
        widths = (40, 240, 100, 100, 100, 140)
        tf = tk.Frame(self, bg=COLORS['bg'])
        tf.pack(fill='both', expand=True, padx=20, pady=(0, 16))
        self.tree = ttk.Treeview(tf, columns=cols, show='headings')
        for c, h, w in zip(cols, hdrs, widths):
            self.tree.heading(c, text=h)
            self.tree.column(c, width=w, minwidth=50)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self.tree.bind('<Double-1>', lambda e: self._open_project())

    def load(self):
        self.tree.delete(*self.tree.get_children())
        conn = get_db()
        rows = conn.execute("SELECT * FROM projects WHERE workspace_id=? ORDER BY id DESC",
                            (self.workspace_id,)).fetchall()
        conn.close()
        for row in rows:
            resources = json.loads(row['resources'] or '[]')
            total = sum(r.get('total', 0) for r in resources)
            try:
                ns = (row['tax_rate'] or 0) / 100
            except Exception:
                ns = 0
            final = total * (1 + ns)
            self.tree.insert('', 'end', iid=str(row['id']), values=(
                row['id'], row['name'], row['status'],
                row['date_start'] or '', row['date_end'] or '',
                f'{final:,.2f} ₽'
            ))

    def _add_project(self):
        p = ProjectPage(self, workspace_id=self.workspace_id, on_save=self.load)

    def _open_project(self):
        sel = self.tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите проект'); return
        p = ProjectPage(self, project_id=int(sel[0]), on_save=self.load)

    def _delete_project(self):
        sel = self.tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите проект'); return
        if messagebox.askyesno('Удаление', 'Удалить проект со всеми ресурсами?'):
            conn = get_db()
            conn.execute("DELETE FROM projects WHERE id=?", (int(sel[0]),))
            conn.commit(); conn.close()
            self.load()

    def _manage_members(self):
        d = WorkspaceMembersDialog(self, self.workspace_id)
        self.wait_window(d)


# ═══════════════════════════════════════════════════════════════════════════
# НАСТРОЙКИ
# ═══════════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════════
# НАСТРОЙКИ
# ═══════════════════════════════════════════════════════════════════════════

class SettingsFrame(ttk.Frame):
    FIELDS = [
        ('company_name',      'Название компании'),
        ('director_name',     'ФИО руководителя'),
        ('director_position', 'Должность руководителя'),
        ('phone',             'Контактный телефон'),
        ('email',             'Контактный email'),
    ]

    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self.configure(style='TFrame')
        self._logo_img = None   # PhotoImage reference (must be kept alive)
        self._build()

    def _build(self):
        sf = ScrollableFrame(self)
        sf.pack(fill='both', expand=True)
        f = sf.inner
        f.configure(style='TFrame')

        tk.Label(f, text='Настройки системы', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(24, 4), padx=24, anchor='w')
        tk.Label(f, text='Данные компании для документов',
                 bg=COLORS['bg'], fg=COLORS['text_dim'], font=FONTS['small']).pack(
            pady=(0, 12), padx=24, anchor='w')

        # ── Данные компании ────────────────────────────────────────────────
        card = card_frame(f)
        card.pack(fill='x', padx=24, pady=4)
        inner = tk.Frame(card, bg=COLORS['surface'])
        inner.pack(fill='x', padx=20, pady=16)

        self.vars = {}
        for key, label in self.FIELDS:
            tk.Label(inner, text=label, bg=COLORS['surface'], fg=COLORS['text_dim'],
                     font=FONTS['small']).pack(anchor='w', pady=(8, 2))
            v = tk.StringVar(value=get_setting(key))
            self.vars[key] = v
            dark_entry(inner, textvariable=v, width=50).pack(fill='x', ipady=5)

        # ── Логотип ────────────────────────────────────────────────────────
        tk.Label(f, text='Логотип предприятия', bg=COLORS['bg'], fg=COLORS['accent'],
                 font=FONTS['heading']).pack(anchor='w', padx=24, pady=(20, 6))

        logo_card = card_frame(f)
        logo_card.pack(fill='x', padx=24, pady=4)
        logo_inner = tk.Frame(logo_card, bg=COLORS['surface'])
        logo_inner.pack(fill='x', padx=20, pady=16)

        # Превью
        self.logo_preview = tk.Label(logo_inner, bg=COLORS['surface'],
                                     text='[Логотип не выбран]',
                                     fg=COLORS['text_dim'], font=FONTS['small'],
                                     width=20, height=6, relief='flat',
                                     highlightthickness=1,
                                     highlightbackground=COLORS['border'])
        self.logo_preview.pack(anchor='w', pady=(0, 10))
        self._refresh_logo_preview()

        btn_row = tk.Frame(logo_inner, bg=COLORS['surface'])
        btn_row.pack(anchor='w')
        RoundedButton(btn_row, '📂 Выбрать файл', command=self._pick_logo,
                      color=COLORS['accent'], width=160).pack(side='left', padx=(0, 10))
        RoundedButton(btn_row, '✕ Удалить лого', command=self._clear_logo,
                      color=COLORS['danger'], width=150).pack(side='left')

        tk.Label(logo_inner,
                 text='Поддерживаемые форматы: PNG, JPG, GIF  |  Рекомендуемый размер: до 400×150 px',
                 bg=COLORS['surface'], fg=COLORS['text_dim'], font=FONTS['small']).pack(
            anchor='w', pady=(8, 0))

        # ── Сохранить ──────────────────────────────────────────────────────
        RoundedButton(f, '💾 Сохранить настройки', command=self._save,
                      color=COLORS['accent'], width=210).pack(pady=20, padx=24, anchor='w')

    # ── Логотип ────────────────────────────────────────────────────────────
    def _refresh_logo_preview(self):
        path = get_setting('logo_path')
        if path and os.path.exists(path):
            try:
                from PIL import Image, ImageTk
                img = Image.open(path)
                img.thumbnail((200, 100), Image.LANCZOS)
                self._logo_img = ImageTk.PhotoImage(img)
                self.logo_preview.config(image=self._logo_img, text='',
                                         width=200, height=100)
                return
            except ImportError:
                pass
            # Fallback без PIL — показываем путь
            self.logo_preview.config(image='', text=os.path.basename(path),
                                     width=30, height=3)
        else:
            self.logo_preview.config(image='', text='[Логотип не выбран]',
                                     width=20, height=6)
            self._logo_img = None

    def _pick_logo(self):
        path = filedialog.askopenfilename(
            title='Выберите логотип',
            filetypes=[('Изображения', '*.png *.jpg *.jpeg *.gif *.bmp'),
                       ('Все файлы', '*.*')])
        if not path: return
        set_setting('logo_path', path)
        self._refresh_logo_preview()

    def _clear_logo(self):
        set_setting('logo_path', '')
        self._refresh_logo_preview()

    def _save(self):
        for key, v in self.vars.items():
            set_setting(key, v.get())
        messagebox.showinfo('Готово', 'Настройки сохранены')


# ═══════════════════════════════════════════════════════════════════════════
# РЕЕСТР ПОЛЬЗОВАТЕЛЕЙ
# ═══════════════════════════════════════════════════════════════════════════

ALL_ROLES = ['Глобальный администратор', 'Коммерческий директор', 'Бухгалтер', 'Кадровик']

# Роли, которым доступна маржинальность (наценка) по ТЗ:
# "Маржинальность... отображается только руководителю". Руководящие роли —
# Глобальный администратор и Коммерческий директор, которые принимают
# решения по ценообразованию. Бухгалтер и Кадровик её не видят.
MARGIN_VISIBLE_ROLES = {'Глобальный администратор', 'Коммерческий директор'}


class Session:
    """Хранит текущего авторизованного пользователя на время работы приложения."""
    current_user = None  # dict с полями users-таблицы, либо None

    @classmethod
    def set_user(cls, user_row):
        cls.current_user = dict(user_row) if user_row is not None else None

    @classmethod
    def get_roles(cls):
        if not cls.current_user:
            return []
        roles = cls.current_user.get('roles', '[]')
        if isinstance(roles, str):
            try:
                roles = json.loads(roles)
            except Exception:
                roles = []
        return roles or []

    @classmethod
    def can_see_margin(cls):
        """True, если у текущего пользователя есть роль, которой доступна
        маржинальность. Если пользователь не авторизован (сессия без логина),
        по умолчанию считаем доступ открытым — чтобы не блокировать работу
        в однопользовательском режиме без настроенных пользователей."""
        if not cls.current_user:
            return True
        roles = set(cls.get_roles())
        return bool(roles & MARGIN_VISIBLE_ROLES)

    @classmethod
    def display_name(cls):
        if not cls.current_user:
            return 'Гость'
        u = cls.current_user
        return f"{u.get('last_name','')} {u.get('first_name','')}".strip() or 'Пользователь'


class LoginDialog(tk.Toplevel):
    """Окно входа в систему. Список пользователей берётся из реестра
    пользователей; пароль сверяется с полем password (см. таблицу users)."""
    def __init__(self, parent):
        super().__init__(parent)
        self.title('Вход в систему')
        self.configure(bg=COLORS['bg'])
        self.geometry('380x320')
        self.resizable(False, False)
        self.result = False
        self.protocol('WM_DELETE_WINDOW', self._on_close)
        self._build()
        self.grab_set()

    def _build(self):
        tk.Label(self, text='IT Cost', bg=COLORS['bg'], fg=COLORS['accent'],
                 font=('Segoe UI', 22, 'bold')).pack(pady=(28, 0))
        tk.Label(self, text='Вход в систему', bg=COLORS['bg'], fg=COLORS['text_dim'],
                 font=FONTS['small']).pack(pady=(0, 20))

        conn = get_db()
        rows = conn.execute(
            "SELECT id, last_name||' '||first_name as n, password FROM users ORDER BY last_name"
        ).fetchall()
        conn.close()
        self._user_map = {r['n']: dict(r) for r in rows}

        f = tk.Frame(self, bg=COLORS['bg'])
        f.pack(fill='x', padx=30)

        tk.Label(f, text='Пользователь', bg=COLORS['bg'], fg=COLORS['text_dim'],
                 font=FONTS['small']).pack(anchor='w')
        self.v_user = tk.StringVar()
        names = list(self._user_map.keys())
        cb = dark_combo(f, names, textvariable=self.v_user, width=30)
        cb.pack(fill='x', pady=(2, 12))
        if names:
            cb.current(0)

        tk.Label(f, text='Пароль', bg=COLORS['bg'], fg=COLORS['text_dim'],
                 font=FONTS['small']).pack(anchor='w')
        self.v_pwd = tk.StringVar()
        dark_entry(f, textvariable=self.v_pwd, width=30, show='•').pack(fill='x', pady=(2, 4))

        if not names:
            tk.Label(self, text='Нет пользователей — вход выполняется без авторизации.',
                     bg=COLORS['bg'], fg=COLORS['warning'], font=FONTS['small'],
                     wraplength=320, justify='left').pack(padx=30, pady=(8, 0))

        btn_f = tk.Frame(self, bg=COLORS['bg'])
        btn_f.pack(pady=20)
        RoundedButton(btn_f, 'Войти без авторизации', command=self._skip,
                      color=COLORS['surface2'], width=180).pack(side='left', padx=6)
        RoundedButton(btn_f, 'Войти', command=self._login,
                      color=COLORS['accent'], width=110).pack(side='left', padx=6)

    def _login(self):
        name = self.v_user.get()
        u = self._user_map.get(name)
        if not u:
            self._skip()
            return
        if (u.get('password') or 'password') != self.v_pwd.get():
            messagebox.showwarning('Ошибка', 'Неверный пароль')
            return
        conn = get_db()
        full = conn.execute("SELECT * FROM users WHERE id=?", (u['id'],)).fetchone()
        conn.close()
        Session.set_user(full)
        self.result = True
        self.destroy()

    def _skip(self):
        Session.set_user(None)
        self.result = True
        self.destroy()

    def _on_close(self):
        self.result = False
        self.destroy()


class UserDialog(BaseDialog):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Пользователь', width=540, height=620)

    def _build(self):
        tk.Label(self, text='Учётная запись', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(pady=(20, 10))
        sf = ScrollableFrame(self)
        sf.pack(fill='both', expand=True)
        f = sf.inner
        f.configure(style='TFrame')

        def row(label, widget, row_n, colspan=1):
            tk.Label(f, text=label, bg=COLORS['bg'], fg=COLORS['text_dim'],
                     font=FONTS['small']).grid(row=row_n, column=0, sticky='nw', pady=6, padx=(20, 8))
            widget.grid(row=row_n, column=1, sticky='ew', pady=6, padx=(0, 20), columnspan=colspan)

        self.v_last  = tk.StringVar(value=self.data.get('last_name', ''))
        self.v_first = tk.StringVar(value=self.data.get('first_name', ''))
        self.v_mid   = tk.StringVar(value=self.data.get('middle_name', ''))
        self.v_email = tk.StringVar(value=self.data.get('email', ''))
        self.v_pos   = tk.StringVar(value=self.data.get('position', ''))
        self.v_pwd   = tk.StringVar(value=self.data.get('password', 'password'))

        row('Фамилия *', dark_entry(f, textvariable=self.v_last, width=30), 0)
        row('Имя *',     dark_entry(f, textvariable=self.v_first, width=30), 1)
        row('Отчество',  dark_entry(f, textvariable=self.v_mid, width=30), 2)
        row('Email',     dark_entry(f, textvariable=self.v_email, width=30), 3)
        row('Должность', dark_entry(f, textvariable=self.v_pos, width=30), 4)
        row('Пароль',    dark_entry(f, textvariable=self.v_pwd, width=30), 5)

        # Roles checkboxes
        tk.Label(f, text='Роли', bg=COLORS['bg'], fg=COLORS['text_dim'],
                 font=FONTS['small']).grid(row=6, column=0, sticky='nw', pady=6, padx=(20, 8))
        roles_f = tk.Frame(f, bg=COLORS['bg'])
        roles_f.grid(row=6, column=1, sticky='ew', pady=6, padx=(0, 20))
        cur_roles = json.loads(self.data.get('roles', '[]')) if isinstance(self.data.get('roles'), str) else (self.data.get('roles') or [])
        self.role_vars = {}
        for role in ALL_ROLES:
            v = tk.BooleanVar(value=role in cur_roles)
            self.role_vars[role] = v
            tk.Checkbutton(roles_f, text=role, variable=v,
                           bg=COLORS['bg'], fg=COLORS['text'], selectcolor=COLORS['surface2'],
                           activebackground=COLORS['bg'], activeforeground=COLORS['text'],
                           font=FONTS['small']).pack(anchor='w')

        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_last.get() or not self.v_first.get():
            messagebox.showwarning('Ошибка', 'Заполните Фамилию и Имя'); return False
        return True

    def _collect(self):
        roles = [r for r, v in self.role_vars.items() if v.get()]
        return dict(
            last_name=self.v_last.get(), first_name=self.v_first.get(),
            middle_name=self.v_mid.get(), email=self.v_email.get(),
            position=self.v_pos.get(), password=self.v_pwd.get(),
            roles=json.dumps(roles, ensure_ascii=False),
        )


class UsersFrame(RegistryFrame):
    TITLE = 'Реестр пользователей'
    COLUMNS = [('id', '#', 40), ('fio', 'ФИО', 220), ('email', 'Email', 180),
               ('position', 'Должность', 140), ('roles_display', 'Роли', 200)]

    def _fetch_rows(self):
        conn = get_db()
        rows = conn.execute(
            "SELECT *, last_name||' '||first_name||COALESCE(' '||middle_name,'') as fio "
            "FROM users ORDER BY last_name").fetchall()
        conn.close()
        result = []
        for r in rows:
            d = dict(r)
            try:
                roles = json.loads(d.get('roles', '[]'))
                d['roles_display'] = ', '.join(roles) if roles else '—'
            except Exception:
                d['roles_display'] = '—'
            result.append(d)
        return result

    def _add(self):
        d = UserDialog(self)
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute(
                "INSERT INTO users(last_name,first_name,middle_name,email,position,password,roles) "
                "VALUES(?,?,?,?,?,?,?)",
                (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                 d.result['email'], d.result['position'], d.result['password'], d.result['roles']))
            conn.commit(); conn.close(); self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid: messagebox.showwarning('Выбор', 'Выберите запись'); return
        conn = get_db()
        row = conn.execute("SELECT * FROM users WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = UserDialog(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = get_db()
            conn.execute(
                "UPDATE users SET last_name=?,first_name=?,middle_name=?,email=?,position=?,password=?,roles=? "
                "WHERE id=?",
                (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                 d.result['email'], d.result['position'], d.result['password'], d.result['roles'], rid))
            conn.commit(); conn.close(); self.load()

    def _do_delete(self, rid):
        conn = get_db()
        conn.execute("DELETE FROM users WHERE id=?", (rid,))
        conn.commit(); conn.close()


# ═══════════════════════════════════════════════════════════════════════════
# УПРАВЛЕНИЕ УЧАСТНИКАМИ РАБОЧЕЙ ОБЛАСТИ
# ═══════════════════════════════════════════════════════════════════════════

class WorkspaceMembersDialog(tk.Toplevel):
    """Диалог управления участниками рабочей области."""
    RIGHTS = ['Просмотр', 'Редактирование проектов']

    def __init__(self, parent, workspace_id):
        super().__init__(parent)
        self.workspace_id = workspace_id
        self.title('Участники рабочей области')
        self.configure(bg=COLORS['bg'])
        self.geometry('640x520')
        self.grab_set()
        self._load()
        self._build()

    def _load(self):
        conn = get_db()
        ws = conn.execute("SELECT * FROM workspaces WHERE id=?", (self.workspace_id,)).fetchone()
        self.ws = dict(ws) if ws else {}
        try:
            self.members = json.loads(self.ws.get('members', '[]'))
        except Exception:
            self.members = []
        self.all_users = conn.execute(
            "SELECT id, last_name||' '||first_name as fio FROM users ORDER BY last_name").fetchall()
        conn.close()

    def _build(self):
        tk.Label(self, text=f'Участники: {self.ws.get("name", "")}',
                 bg=COLORS['bg'], fg=COLORS['text'], font=FONTS['title']).pack(pady=(16, 4), padx=20, anchor='w')

        # Add user panel
        add_f = tk.Frame(self, bg=COLORS['bg'])
        add_f.pack(fill='x', padx=20, pady=8)
        tk.Label(add_f, text='Добавить:', bg=COLORS['bg'], fg=COLORS['text_dim'],
                 font=FONTS['small']).pack(side='left', padx=(0, 8))
        self.v_add_user = tk.StringVar()
        user_names = [f"{u['fio']} (#{u['id']})" for u in self.all_users]
        self.user_id_map = {f"{u['fio']} (#{u['id']})": u['id'] for u in self.all_users}
        self.cb_user = dark_combo(add_f, user_names, textvariable=self.v_add_user, width=28)
        self.cb_user.pack(side='left', padx=(0, 10))
        self.v_right = tk.StringVar(value='Просмотр')
        dark_combo(add_f, self.RIGHTS, textvariable=self.v_right, width=22).pack(side='left', padx=(0, 10))
        RoundedButton(add_f, '＋ Добавить', command=self._add_member,
                      color=COLORS['accent'], width=120).pack(side='left')

        # Members list
        tf = tk.Frame(self, bg=COLORS['bg'])
        tf.pack(fill='both', expand=True, padx=20, pady=(0, 8))
        self.tree = ttk.Treeview(tf, columns=('uid', 'fio', 'right'), show='headings')
        self.tree.heading('uid', text='#'); self.tree.column('uid', width=40)
        self.tree.heading('fio', text='ФИО'); self.tree.column('fio', width=280)
        self.tree.heading('right', text='Права'); self.tree.column('right', width=200)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self._reload_tree()

        # Bottom buttons
        bf = tk.Frame(self, bg=COLORS['bg'])
        bf.pack(fill='x', padx=20, pady=10)
        RoundedButton(bf, '✕ Удалить выбранного', command=self._remove_member,
                      color=COLORS['danger'], width=180).pack(side='left', padx=(0, 10))
        RoundedButton(bf, '💾 Сохранить', command=self._save,
                      color=COLORS['accent'], width=140).pack(side='left')
        RoundedButton(bf, 'Закрыть', command=self.destroy,
                      color=COLORS['surface2'], width=110).pack(side='right')

    def _reload_tree(self):
        self.tree.delete(*self.tree.get_children())
        conn = get_db()
        for m in self.members:
            uid = m.get('user_id')
            right = m.get('right', 'Просмотр')
            u = conn.execute(
                "SELECT last_name||' '||first_name as fio FROM users WHERE id=?", (uid,)).fetchone()
            fio = u['fio'] if u else f'ID={uid}'
            self.tree.insert('', 'end', iid=str(uid), values=(uid, fio, right))
        conn.close()

    def _add_member(self):
        name = self.v_add_user.get()
        uid = self.user_id_map.get(name)
        if not uid:
            messagebox.showwarning('Ошибка', 'Выберите пользователя'); return
        if any(m['user_id'] == uid for m in self.members):
            messagebox.showwarning('Ошибка', 'Пользователь уже добавлен'); return
        self.members.append({'user_id': uid, 'right': self.v_right.get()})
        self._reload_tree()

    def _remove_member(self):
        sel = self.tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите участника'); return
        uid = int(sel[0])
        self.members = [m for m in self.members if m['user_id'] != uid]
        self._reload_tree()

    def _save(self):
        conn = get_db()
        conn.execute("UPDATE workspaces SET members=? WHERE id=?",
                     (json.dumps(self.members, ensure_ascii=False), self.workspace_id))
        conn.commit(); conn.close()
        messagebox.showinfo('Готово', 'Состав участников сохранён')


# ═══════════════════════════════════════════════════════════════════════════
# ГЛАВНОЕ ОКНО
# ═══════════════════════════════════════════════════════════════════════════

class MainApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('IT Cost — Система расчёта стоимости IT-проекта')
        self.geometry('1280x800')
        self.minsize(1000, 650)
        self.configure(bg=COLORS['bg'])
        apply_ttk_styles()
        self._build()

    def _build(self):
        # Sidebar
        sidebar = tk.Frame(self, bg=COLORS['surface'], width=220)
        sidebar.pack(side='left', fill='y')
        sidebar.pack_propagate(False)

        # Logo area
        logo_f = tk.Frame(sidebar, bg=COLORS['surface'])
        logo_f.pack(fill='x', pady=(24, 16))
        tk.Label(logo_f, text='IT Cost', bg=COLORS['surface'],
                 fg=COLORS['accent'], font=('Segoe UI', 20, 'bold')).pack(padx=20, anchor='w')
        tk.Label(logo_f, text='Расчёт стоимости проектов', bg=COLORS['surface'],
                 fg=COLORS['text_dim'], font=FONTS['small']).pack(padx=20, anchor='w')

        tk.Frame(sidebar, bg=COLORS['border'], height=1).pack(fill='x', padx=16, pady=8)

        # Content area
        self.content = tk.Frame(self, bg=COLORS['bg'])
        self.content.pack(side='right', fill='both', expand=True)

        self._current_frame = None
        self._nav_btns = {}

        sections = [
            ('🏠', 'Главная', self._show_home),
            ('─', None, None),
            ('📁', 'Рабочие области', self._show_workspaces),
            ('─', None, None),
            ('👥', 'Сотрудники', lambda: self._show_registry(EmployeesFrame)),
            ('🧑‍💼', 'Исполнители', lambda: self._show_registry(ContractorsFrame)),
            ('🏢', 'Субподрядчики', lambda: self._show_registry(SubcontractorsFrame)),
            ('🖥️', 'Оборудование', lambda: self._show_registry(EquipmentFrame)),
            ('📋', 'Заказчики', lambda: self._show_registry(CustomersFrame)),
            ('─', None, None),
            ('🔑', 'Пользователи', self._show_users),
            ('⚙️', 'Настройки', self._show_settings),
        ]

        for icon, label, cmd in sections:
            if icon == '─':
                tk.Frame(sidebar, bg=COLORS['border'], height=1).pack(fill='x', padx=16, pady=4)
                continue
            btn = tk.Button(sidebar, text=f'  {icon}  {label}', anchor='w',
                            bg=COLORS['surface'], fg=COLORS['text'], activebackground=COLORS['hover'],
                            activeforeground=COLORS['text'], relief='flat', font=FONTS['normal'],
                            cursor='hand2', pady=10, command=cmd)
            btn.pack(fill='x', padx=8)
            btn.bind('<Enter>', lambda e, b=btn: b.config(bg=COLORS['hover']))
            btn.bind('<Leave>', lambda e, b=btn: b.config(
                bg=COLORS['accent'] if b == self._active_btn else COLORS['surface']))
            self._nav_btns[label] = btn

        self._active_btn = None
        self._show_home()

        # Индикатор текущего пользователя (для понимания, почему маржа
        # видна/скрыта — см. Session.can_see_margin())
        user_f = tk.Frame(sidebar, bg=COLORS['surface'])
        user_f.pack(fill='x', side='bottom', pady=(8, 16), padx=16)
        tk.Frame(sidebar, bg=COLORS['border'], height=1).pack(fill='x', side='bottom', padx=16, pady=(0, 8))
        roles = Session.get_roles()
        roles_txt = ', '.join(roles) if roles else 'без роли'
        margin_txt = '👁 маржа видна' if Session.can_see_margin() else '🔒 маржа скрыта'
        tk.Label(user_f, text=Session.display_name(), bg=COLORS['surface'],
                 fg=COLORS['text'], font=FONTS['small']).pack(anchor='w')
        tk.Label(user_f, text=roles_txt, bg=COLORS['surface'],
                 fg=COLORS['text_dim'], font=('Segoe UI', 8)).pack(anchor='w')
        tk.Label(user_f, text=margin_txt, bg=COLORS['surface'],
                 fg=COLORS['text_dim'], font=('Segoe UI', 8)).pack(anchor='w', pady=(4, 0))

    def _activate_btn(self, label):
        if self._active_btn:
            self._active_btn.config(bg=COLORS['surface'])
        btn = self._nav_btns.get(label)
        if btn:
            btn.config(bg=COLORS['accent'])
            self._active_btn = btn

    def _switch(self, frame):
        if self._current_frame:
            self._current_frame.destroy()
        self._current_frame = frame
        frame.pack(fill='both', expand=True)

    def _show_home(self):
        self._activate_btn('Главная')
        f = ttk.Frame(self.content)
        company = get_setting('company_name', 'Моя компания')
        tk.Label(f, text=f'Добро пожаловать', bg=COLORS['bg'],
                 fg=COLORS['text_dim'], font=FONTS['heading']).pack(pady=(60, 4))
        tk.Label(f, text=company, bg=COLORS['bg'],
                 fg=COLORS['text'], font=('Segoe UI', 32, 'bold')).pack()
        tk.Label(f, text='Система расчёта стоимости IT-проектов',
                 bg=COLORS['bg'], fg=COLORS['text_dim'], font=FONTS['normal']).pack(pady=(4, 40))

        conn = get_db()
        n_projects = conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        n_employees = conn.execute("SELECT COUNT(*) FROM employees").fetchone()[0]
        n_customers = conn.execute("SELECT COUNT(*) FROM customers").fetchone()[0]
        conn.close()

        stats_f = tk.Frame(f, bg=COLORS['bg'])
        stats_f.pack()
        for val, lbl in [(n_projects, 'Проектов'), (n_employees, 'Сотрудников'), (n_customers, 'Заказчиков')]:
            c = card_frame(stats_f)
            c.pack(side='left', padx=12, ipadx=20, ipady=12)
            tk.Label(c, text=str(val), bg=COLORS['surface'],
                     fg=COLORS['accent'], font=FONTS['big']).pack()
            tk.Label(c, text=lbl, bg=COLORS['surface'],
                     fg=COLORS['text_dim'], font=FONTS['small']).pack()
        self._switch(f)

    def _show_workspaces(self):
        self._activate_btn('Рабочие области')
        f = ttk.Frame(self.content)

        hdr = tk.Frame(f, bg=COLORS['bg'])
        hdr.pack(fill='x', padx=20, pady=(16, 8))
        tk.Label(hdr, text='Рабочие области', bg=COLORS['bg'], fg=COLORS['text'],
                 font=FONTS['title']).pack(side='left')

        nb_frame = tk.Frame(f, bg=COLORS['bg'])
        nb_frame.pack(fill='both', expand=True)

        # Tabs for each workspace + "Управление"
        nb = ttk.Notebook(nb_frame)
        nb.pack(fill='both', expand=True, padx=20)

        self._ws_nb = nb
        self._refresh_workspaces(nb, f)
        self._switch(f)

    def _refresh_workspaces(self, nb, parent):
        for tab in nb.tabs():
            nb.forget(tab)

        # Management tab
        mgmt = ttk.Frame(nb)
        nb.add(mgmt, text='  ⚙ Управление  ')
        self._build_ws_mgmt(mgmt, nb, parent)

        # Workspace tabs
        conn = get_db()
        wss = conn.execute("SELECT * FROM workspaces ORDER BY name").fetchall()
        conn.close()
        for ws in wss:
            wf = WorkspaceFrame(nb, workspace_id=ws['id'])
            nb.add(wf, text=f'  📁 {ws["name"]}  ')

    def _build_ws_mgmt(self, parent, nb, top):
        hdr = tk.Frame(parent, bg=COLORS['bg'])
        hdr.pack(fill='x', padx=20, pady=(16, 8))
        tk.Label(hdr, text='Управление рабочими областями', bg=COLORS['bg'],
                 fg=COLORS['text'], font=FONTS['heading']).pack(side='left')
        RoundedButton(hdr, '＋ Создать', command=lambda: self._add_workspace(nb, top),
                      color=COLORS['accent'], width=120).pack(side='right', padx=4)
        RoundedButton(hdr, '✎ Изменить', command=lambda: self._edit_workspace(tree, nb, top),
                      color=COLORS['surface2'], width=120).pack(side='right', padx=4)
        RoundedButton(hdr, '✕ Удалить', command=lambda: self._del_workspace(tree, nb, top),
                      color=COLORS['danger'], width=110).pack(side='right', padx=4)

        tf = tk.Frame(parent, bg=COLORS['bg'])
        tf.pack(fill='both', expand=True, padx=20)
        tree = ttk.Treeview(tf, columns=('id', 'name', 'subdomain'), show='headings')
        tree.heading('id', text='#'); tree.column('id', width=40)
        tree.heading('name', text='Название'); tree.column('name', width=200)
        tree.heading('subdomain', text='Поддомен'); tree.column('subdomain', width=200)
        sb = ttk.Scrollbar(tf, orient='vertical', command=tree.yview)
        tree.configure(yscrollcommand=sb.set)
        tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')

        conn = get_db()
        for ws in conn.execute("SELECT * FROM workspaces ORDER BY name").fetchall():
            tree.insert('', 'end', iid=str(ws['id']), values=(ws['id'], ws['name'], ws['subdomain'] or ''))
        conn.close()

    def _add_workspace(self, nb, top):
        name = simpledialog.askstring('Новая область', 'Название рабочей области:', parent=self)
        if not name: return
        sub = simpledialog.askstring('Поддомен', 'Поддомен (например: workspace1):', parent=self)
        conn = get_db()
        conn.execute("INSERT INTO workspaces(name,subdomain) VALUES(?,?)", (name, sub or ''))
        conn.commit(); conn.close()
        self._refresh_workspaces(nb, top)

    def _edit_workspace(self, tree, nb, top):
        sel = tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите область'); return
        wid = int(sel[0])
        conn = get_db()
        ws = conn.execute("SELECT * FROM workspaces WHERE id=?", (wid,)).fetchone()
        conn.close()
        name = simpledialog.askstring('Изменить область', 'Название:', initialvalue=ws['name'], parent=self)
        if not name: return
        sub = simpledialog.askstring('Поддомен', 'Поддомен:', initialvalue=ws['subdomain'] or '', parent=self)
        conn = get_db()
        conn.execute("UPDATE workspaces SET name=?,subdomain=? WHERE id=?", (name, sub or '', wid))
        conn.commit(); conn.close()
        self._refresh_workspaces(nb, top)

    def _del_workspace(self, tree, nb, top):
        sel = tree.selection()
        if not sel: messagebox.showwarning('Выбор', 'Выберите область'); return
        if messagebox.askyesno('Удаление', 'Удалить рабочую область и все проекты?'):
            wid = int(sel[0])
            conn = get_db()
            conn.execute("DELETE FROM projects WHERE workspace_id=?", (wid,))
            conn.execute("DELETE FROM workspaces WHERE id=?", (wid,))
            conn.commit(); conn.close()
            self._refresh_workspaces(nb, top)

    def _show_registry(self, cls):
        name_map = {EmployeesFrame: 'Сотрудники', ContractorsFrame: 'Исполнители',
                    SubcontractorsFrame: 'Субподрядчики', EquipmentFrame: 'Оборудование',
                    CustomersFrame: 'Заказчики'}
        self._activate_btn(name_map.get(cls, ''))
        f = cls(self.content)
        self._switch(f)

    def _show_users(self):
        self._activate_btn('Пользователи')
        f = UsersFrame(self.content)
        self._switch(f)

    def _show_settings(self):
        self._activate_btn('Настройки')
        f = SettingsFrame(self.content)
        self._switch(f)


# ═══════════════════════════════════════════════════════════════════════════
# ЭКСПОРТ XLSX
# ═══════════════════════════════════════════════════════════════════════════

def export_nma_xlsx(path, project_name, ds, de, cost_sum, resources):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'НМА'
    hdr_fill = PatternFill(fgColor='1A2550', fill_type='solid')
    hdr_font = Font(color='FFFFFF', bold=True)
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    company   = get_setting('company_name')
    director  = get_setting('director_name')
    dir_pos   = get_setting('director_position')
    phone     = get_setting('phone')
    email_val = get_setting('email')
    logo_path = get_setting('logo_path')

    start_row = 1
    # Логотип
    if logo_path and os.path.exists(logo_path):
        try:
            from openpyxl.drawing.image import Image as XLImage
            img = XLImage(logo_path)
            img.height = 60
            img.width  = 160
            ws.add_image(img, 'A1')
            ws.row_dimensions[1].height = 50
            start_row = 4
        except Exception:
            pass

    r = start_row
    ws.cell(r, 1, 'ФОРМА СТОИМОСТИ НЕМАТЕРИАЛЬНОГО АКТИВА').font = Font(bold=True, size=14)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    r += 1
    ws.cell(r, 1, company); r += 1
    if phone:     ws.cell(r, 1, f'Тел.: {phone}'); r += 1
    if email_val: ws.cell(r, 1, f'Email: {email_val}'); r += 1
    ws.cell(r, 1, f'Проект: {project_name}'); r += 1
    ws.cell(r, 1, f'Период: {ds} — {de}'); r += 1
    c = ws.cell(r, 1, f'Итоговая стоимость НМА: {cost_sum:,.2f} ₽')
    c.font = Font(bold=True); r += 1
    ws.append([])  # пустая строка
    r += 1

    headers = ['Название', 'Тип', 'Исполнитель', 'Единиц', 'Себестоимость']
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(r, ci, h)
        cell.fill = hdr_fill; cell.font = hdr_font; cell.border = border
    r += 1

    for res in resources:
        for ci, val in enumerate([res.get('name'), res.get('resource_type'),
                                   res.get('executor_display'), res.get('units'),
                                   res.get('cost', 0)], 1):
            cell = ws.cell(r, ci, val)
            cell.border = border
        r += 1

    for ci, val in enumerate(['', '', '', 'ИТОГО', cost_sum], 1):
        cell = ws.cell(r, ci, val)
        cell.font = Font(bold=True)
    r += 2

    ws.cell(r, 1, f'{dir_pos}  _______________  {director}')

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 16
    ws.column_dimensions['C'].width = 24
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 18
    wb.save(path)


def export_kp_xlsx(path, project_name, cust, ds, de, final, resources, company, director, dir_pos):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'КП'
    hdr_fill = PatternFill(fgColor='1A2550', fill_type='solid')
    hdr_font = Font(color='FFFFFF', bold=True)
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    phone_val = get_setting('phone')
    email_val = get_setting('email')
    logo_path = get_setting('logo_path')

    start_row = 1
    if logo_path and os.path.exists(logo_path):
        try:
            from openpyxl.drawing.image import Image as XLImage
            img = XLImage(logo_path)
            img.height = 60; img.width = 160
            ws.add_image(img, 'A1')
            ws.row_dimensions[1].height = 50
            start_row = 4
        except Exception:
            pass

    r = start_row
    ws.cell(r, 1, 'КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ').font = Font(bold=True, size=16)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    r += 1
    ws.cell(r, 1, company); r += 1
    if phone_val:  ws.cell(r, 1, f'Тел.: {phone_val}'); r += 1
    if email_val:  ws.cell(r, 1, f'Email: {email_val}'); r += 1
    ws.cell(r, 1, f'Заказчик: {cust.get("name", "")} | {cust.get("director_name", "")}'); r += 1
    ws.cell(r, 1, f'Email заказчика: {cust.get("email", "")}'); r += 1
    ws.cell(r, 1, f'Проект: {project_name}'); r += 1
    ws.cell(r, 1, f'Дата: {datetime.now().strftime("%d.%m.%Y")}'); r += 1
    c = ws.cell(r, 1, f'Итоговая стоимость: {final:,.2f} ₽')
    c.font = Font(bold=True, size=12); r += 1
    ws.append([]); r += 1

    headers = ['Название услуги', 'Кол-во', 'Ед.', 'Начало', 'Окончание', 'Стоимость (с маржой)']
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(r, ci, h)
        cell.fill = hdr_fill; cell.font = hdr_font; cell.border = border
    r += 1

    for res in resources:
        service_name = res.get('service') or res.get('name')
        for ci, val in enumerate([service_name, res.get('units'), res.get('unit_label', ''),
                                   res.get('date_start'), res.get('date_end'), res.get('total', 0)], 1):
            ws.cell(r, ci, val).border = border
        r += 1

    for ci, val in enumerate(['', '', '', '', 'ИТОГО', final], 1):
        ws.cell(r, ci, val).font = Font(bold=True)
    r += 2

    ws.cell(r, 1, f'{dir_pos}  _______________  {director}')

    for col, w in zip('ABCDEF', [30, 10, 10, 12, 12, 20]):
        ws.column_dimensions[col].width = w
    wb.save(path)


# ═══════════════════════════════════════════════════════════════════════════
# ЭКСПОРТ DOCX
# ═══════════════════════════════════════════════════════════════════════════

def export_nma_docx(path, project_name, ds, de, cost_sum, resources):
    doc = Document()

    company   = get_setting('company_name')
    director  = get_setting('director_name')
    dir_pos   = get_setting('director_position')
    phone_val = get_setting('phone')
    email_val = get_setting('email')
    logo_path = get_setting('logo_path')

    # Логотип
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Cm(5))
        except Exception:
            pass

    # Шапка компании
    doc.add_heading('ФОРМА СТОИМОСТИ НЕМАТЕРИАЛЬНОГО АКТИВА', 0)
    doc.add_paragraph(company)
    if phone_val: doc.add_paragraph(f'Тел.: {phone_val}')
    if email_val: doc.add_paragraph(f'Email: {email_val}')
    doc.add_paragraph()
    doc.add_paragraph(f'Проект: {project_name}')
    doc.add_paragraph(f'Период: {ds} — {de}')
    p = doc.add_paragraph()
    p.add_run(f'Итоговая стоимость НМА: {cost_sum:,.2f} ₽').bold = True
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Название', 'Тип', 'Исполнитель', 'Единиц', 'Себестоимость']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for r in resources:
        row = table.add_row().cells
        row[0].text = str(r.get('name', ''))
        row[1].text = str(r.get('resource_type', ''))
        row[2].text = str(r.get('executor_display', ''))
        row[3].text = str(r.get('units', ''))
        row[4].text = f"{r.get('cost', 0):,.2f}"

    total_row = table.add_row().cells
    total_row[3].text = 'ИТОГО'
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[4].text = f'{cost_sum:,.2f} ₽'
    total_row[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(f'{dir_pos}  _______________  {director}')
    doc.save(path)


def export_kp_docx(path, project_name, cust, ds, de, final, resources, company, director, dir_pos):
    doc = Document()

    phone_val = get_setting('phone')
    email_val = get_setting('email')
    logo_path = get_setting('logo_path')

    # Логотип
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Cm(5))
        except Exception:
            pass

    doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', 0)
    doc.add_paragraph(company)
    if phone_val: doc.add_paragraph(f'Тел.: {phone_val}')
    if email_val: doc.add_paragraph(f'Email: {email_val}')
    doc.add_paragraph()
    doc.add_paragraph(f'Заказчик: {cust.get("name", "")} | {cust.get("director_name", "")}')
    doc.add_paragraph(f'Email заказчика: {cust.get("email", "")}')
    doc.add_paragraph(f'Проект: {project_name}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    p = doc.add_paragraph()
    p.add_run(f'Итоговая стоимость: {final:,.2f} ₽').bold = True
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Услуга', 'Кол-во', 'Ед.', 'Начало', 'Окончание', 'Стоимость']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for r in resources:
        row = table.add_row().cells
        row[0].text = str(r.get('service') or r.get('name', ''))
        row[1].text = str(r.get('units', ''))
        row[2].text = str(r.get('unit_label', ''))
        row[3].text = str(r.get('date_start', ''))
        row[4].text = str(r.get('date_end', ''))
        row[5].text = f"{r.get('total', 0):,.2f}"

    t_row = table.add_row().cells
    t_row[4].text = 'ИТОГО'
    t_row[4].paragraphs[0].runs[0].bold = True
    t_row[5].text = f'{final:,.2f} ₽'
    t_row[5].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f'{dir_pos}  _______________  {director}')
    doc.save(path)


# ═══════════════════════════════════════════════════════════════════════════
# ЭКСПОРТ PDF
# ═══════════════════════════════════════════════════════════════════════════

def register_cyrillic_font():
    """Регистрация шрифта с поддержкой кириллицы"""
    try:
        # Если уже зарегистрирован — не повторяем
        pdfmetrics.getFont('Arial')
        return True
    except Exception:
        pass
    try:
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",                               # Windows
            "C:/Windows/Fonts/Arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",          # Linux
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",             # macOS
            os.path.expanduser("~/Library/Fonts/Arial.ttf"),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'arial.ttf'),
            "arial.ttf",
        ]
        # matplotlib — почти всегда есть в pip-окружениях
        try:
            import importlib.util
            spec = importlib.util.find_spec('matplotlib')
            if spec:
                mpl_dir = os.path.dirname(spec.origin)
                font_paths.insert(0, os.path.join(mpl_dir, 'mpl-data', 'fonts', 'ttf', 'DejaVuSans.ttf'))
        except Exception:
            pass
        for path in font_paths:
            if path and os.path.exists(path):
                pdfmetrics.registerFont(TTFont('Arial', path))
                pdfmetrics.registerFontFamily('Arial', normal='Arial')
                return True
        return False
    except Exception:
        return False


def export_nma_pdf(path, project_name, ds, de, cost_sum, resources):
    if not HAS_PDF:
        messagebox.showerror("Ошибка", "ReportLab не установлен")
        return

    register_cyrillic_font()

    company   = get_setting('company_name')
    director  = get_setting('director_name')
    dir_pos   = get_setting('director_position')
    phone_val = get_setting('phone')
    email_val = get_setting('email')
    logo_path = get_setting('logo_path')

    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    try:
        styles.add(ParagraphStyle(name='Russian',     fontName='Arial', fontSize=11, leading=14))
        styles.add(ParagraphStyle(name='RussianBold', fontName='Arial', fontSize=12, leading=14, alignment=1))
        styles.add(ParagraphStyle(name='RussianSm',   fontName='Arial', fontSize=9,  leading=12, textColor=colors.grey))
    except Exception:
        pass

    story = []

    # Логотип
    if logo_path and os.path.exists(logo_path):
        try:
            from reportlab.platypus import Image as RLImage
            logo = RLImage(logo_path, width=4*cm, height=2*cm, kind='proportional')
            story.append(logo)
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.append(Paragraph("ФОРМА СТОИМОСТИ НЕМАТЕРИАЛЬНОГО АКТИВА", styles['RussianBold']))
    story.append(Spacer(1, 8))
    story.append(Paragraph(company, styles['Russian']))
    if phone_val: story.append(Paragraph(f"Тел.: {phone_val}", styles.get('RussianSm', styles['Russian'])))
    if email_val: story.append(Paragraph(f"Email: {email_val}", styles.get('RussianSm', styles['Russian'])))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Проект:</b> {project_name}", styles['Russian']))
    story.append(Paragraph(f"<b>Период:</b> {ds} — {de}", styles['Russian']))
    story.append(Paragraph(f"<b>Итоговая стоимость НМА:</b> {cost_sum:,.2f} руб.", styles['RussianBold']))
    story.append(Spacer(1, 20))

    data = [['Название', 'Тип', 'Исполнитель', 'Единиц', 'Себестоимость']]
    for r in resources:
        data.append([
            r.get('name', ''),
            r.get('resource_type', ''),
            r.get('executor_display', ''),
            str(r.get('units', '')),
            f"{r.get('cost', 0):,.2f}",
        ])
    data.append(['', '', '', 'ИТОГО', f"{cost_sum:,.2f}"])

    table = Table(data, colWidths=[200, 80, 140, 60, 90])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A2550')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('FONTNAME', (0, -1), (-1, -1), 'Arial'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F5F7FA')]),
    ]))
    story.append(table)
    story.append(Spacer(1, 40))
    story.append(Paragraph(f"{dir_pos}  _______________  {director}", styles['Russian']))
    doc.build(story)
    messagebox.showinfo("Успешно", f"PDF сохранён:\n{path}")


def export_kp_pdf(path, project_name, cust, ds, de, final, resources, company, director, dir_pos):
    if not HAS_PDF:
        messagebox.showerror("Ошибка", "ReportLab не установлен")
        return

    register_cyrillic_font()

    phone_val = get_setting('phone')
    email_val = get_setting('email')
    logo_path = get_setting('logo_path')

    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    try:
        styles.add(ParagraphStyle(name='Russian',     fontName='Arial', fontSize=11, leading=14))
        styles.add(ParagraphStyle(name='RussianBold', fontName='Arial', fontSize=12, leading=14, alignment=1))
        styles.add(ParagraphStyle(name='RussianSm',   fontName='Arial', fontSize=9,  leading=12, textColor=colors.grey))
    except Exception:
        pass

    story = []

    # Логотип
    if logo_path and os.path.exists(logo_path):
        try:
            from reportlab.platypus import Image as RLImage
            logo = RLImage(logo_path, width=4*cm, height=2*cm, kind='proportional')
            story.append(logo)
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.append(Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", styles['RussianBold']))
    story.append(Spacer(1, 8))
    story.append(Paragraph(company, styles['Russian']))
    if phone_val: story.append(Paragraph(f"Тел.: {phone_val}", styles.get('RussianSm', styles['Russian'])))
    if email_val: story.append(Paragraph(f"Email: {email_val}", styles.get('RussianSm', styles['Russian'])))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Заказчик:</b> {cust.get('name', '')} | {cust.get('director_name', '')}", styles['Russian']))
    story.append(Paragraph(f"<b>Проект:</b> {project_name}", styles['Russian']))
    story.append(Paragraph(f"<b>Дата:</b> {datetime.now().strftime('%d.%m.%Y')}", styles['Russian']))
    story.append(Paragraph(f"<b>Итоговая стоимость:</b> {final:,.2f} руб.", styles['RussianBold']))
    story.append(Spacer(1, 20))

    data = [['Услуга', 'Кол-во', 'Ед.', 'Начало', 'Окончание', 'Стоимость']]
    for r in resources:
        data.append([
            r.get('service') or r.get('name', ''),
            str(r.get('units', '')),
            r.get('unit_label', ''),
            r.get('date_start', ''),
            r.get('date_end', ''),
            f"{r.get('total', 0):,.2f}",
        ])
    data.append(['', '', '', '', 'ИТОГО', f"{final:,.2f}"])

    table = Table(data, colWidths=[160, 110, 50, 65, 65, 90])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A2550')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F5F7FA')]),
    ]))
    story.append(table)
    story.append(Spacer(1, 40))
    story.append(Paragraph(f"{dir_pos} ____________________ {director}", styles['Russian']))
    doc.build(story)
    messagebox.showinfo("Успешно", f"PDF сохранён:\n{path}")


# ═══════════════════════════════════════════════════════════════════════════
# ТОЧКА ВХОДА
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    init_db()
    _root = tk.Tk()
    _root.withdraw()
    apply_ttk_styles()
    _login = LoginDialog(_root)
    _root.wait_window(_login)
    _root.destroy()
    if _login.result:
        app = MainApp()
        app.mainloop()