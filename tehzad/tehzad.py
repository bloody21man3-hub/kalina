"""СметаПро — платформа управления проектными сметами (Tkinter, SQLite)."""

APP_NAME = 'СметаПро'
APP_TAGLINE = 'Платформа управления проектными сметами'

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
import sqlite3
import json
import os
from datetime import datetime, date, timedelta
import calendar

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

DB_PATH = "baza.db"

# ---------------------------------------------------------------------------
# База 
# ---------------------------------------------------------------------------

def connect_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db():
    conn = connect_db()
    c = conn.cursor()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT
    );
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        last_name TEXT NOT NULL,
        first_name TEXT NOT NULL,
        middle_name TEXT,
        email TEXT,
        position TEXT,
        roles TEXT DEFAULT '[]',
        workspace_ids TEXT DEFAULT '[]',
        password TEXT DEFAULT 'password'
    );
    CREATE TABLE IF NOT EXISTS workspaces (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        subdomain TEXT,
        admin_user_id INTEGER,
        members TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        inn TEXT,
        type TEXT,
        name TEXT,
        director_name TEXT,
        email TEXT,
        phone TEXT
    );
    CREATE TABLE IF NOT EXISTS employees (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        last_name TEXT NOT NULL,
        first_name TEXT NOT NULL,
        middle_name TEXT,
        position TEXT,
        salary REAL DEFAULT 0,
        tax_rate REAL DEFAULT 30.2,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS contractors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        last_name TEXT NOT NULL,
        first_name TEXT NOT NULL,
        middle_name TEXT,
        contract_type TEXT DEFAULT 'ГПХ',
        tax_rate REAL DEFAULT 0,
        unit TEXT DEFAULT 'часы',
        rate REAL DEFAULT 0,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS subcontractors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        inn TEXT,
        type TEXT DEFAULT 'ООО',
        name TEXT NOT NULL,
        director_name TEXT,
        email TEXT,
        phone TEXT,
        unit TEXT DEFAULT 'часы',
        rate REAL DEFAULT 0,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS equipment (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        description TEXT,
        ownership TEXT DEFAULT 'Собственное',
        rental_cost REAL DEFAULT 0,
        unit TEXT DEFAULT 'часы',
        rate REAL DEFAULT 0,
        services TEXT DEFAULT '[]'
    );
    CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        workspace_id INTEGER,
        name TEXT NOT NULL,
        date_start TEXT,
        date_end TEXT,
        description TEXT,
        customer_id INTEGER,
        tax_rate REAL DEFAULT 0,
        status TEXT DEFAULT 'Активный',
        resources TEXT DEFAULT '[]'
    );
    """)
    defaults = {
        'company_name': 'Моя компания',
        'director_name': 'Иванов Иван Иванович',
        'director_position': 'Генеральный директор',
        'logo_path': '',
        'phone': '',
        'email': '',
    }
    for k, v in defaults.items():
        c.execute("INSERT OR IGNORE INTO settings(key,value) VALUES(?,?)", (k, v))
    conn.commit()
    conn.close()


def read_setting(key, default=''):
    conn = connect_db()
    row = conn.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
    conn.close()
    return row['value'] if row else default


def write_setting(key, value):
    conn = connect_db()
    conn.execute("INSERT OR REPLACE INTO settings(key,value) VALUES(?,?)", (key, value))
    conn.commit()
    conn.close()


# ---------------------------------------------------------------------------
# Тема оформления
# ---------------------------------------------------------------------------

PALETTE = {
    'bg': '#F0F2F5',
    'surface': '#FFFFFF',
    'surface2': '#E4E8ED',
    'border': '#CBD2D9',
    'primary': '#2563EB',
    'primary_light': '#3B82F6',
    'accent': '#D97706',
    'accent_light': '#F59E0B',
    'success': '#059669',
    'warning': '#D97706',
    'danger': '#DC2626',
    'text': '#1E293B',
    'text_dim': '#64748B',
    'white': '#FFFFFF',
    'hover': '#E2E8F0',
    'sidebar': '#1E293B',
    'sidebar_text': '#CBD5E1',
    'sidebar_active': '#334155',
    'sidebar_section': '#64748B',
    'header': '#FFFFFF',
    'header_text': '#1E293B',
    'stat_bg': '#EFF6FF',
}

TYPEFACE = {
    'title': ('Verdana', 20, 'bold'),
    'heading': ('Verdana', 12, 'bold'),
    'normal': ('Verdana', 10),
    'small': ('Verdana', 9),
    'mono': ('Consolas', 10),
    'big': ('Verdana', 24, 'bold'),
    'brand': ('Verdana', 22, 'bold'),
    'section': ('Verdana', 8, 'bold'),
}


def configure_theme():
    style = ttk.Style()
    style.theme_use('clam')
    bg = PALETTE['bg']
    surf = PALETTE['surface']
    pri = PALETTE['primary']
    txt = PALETTE['text']
    dim = PALETTE['text_dim']

    style.configure('TFrame', background=bg)
    style.configure('Surface.TFrame', background=surf)
    style.configure('Surface2.TFrame', background=PALETTE['surface2'])
    style.configure('Header.TFrame', background=PALETTE['header'])

    style.configure('TLabel', background=bg, foreground=txt, font=TYPEFACE['normal'])
    style.configure('Dim.TLabel', background=bg, foreground=dim, font=TYPEFACE['small'])
    style.configure('Surface.TLabel', background=surf, foreground=txt, font=TYPEFACE['normal'])
    style.configure('Heading.TLabel', background=bg, foreground=txt, font=TYPEFACE['heading'])
    style.configure('Title.TLabel', background=bg, foreground=txt, font=TYPEFACE['title'])
    style.configure('Accent.TLabel', background=bg, foreground=pri, font=TYPEFACE['heading'])
    style.configure('Header.TLabel', background=PALETTE['header'], foreground=PALETTE['header_text'],
                    font=TYPEFACE['heading'])

    style.configure('TEntry', fieldbackground=surf, background=surf,
                    foreground=txt, insertcolor=txt, borderwidth=1, relief='solid',
                    font=TYPEFACE['normal'])
    style.map('TEntry', fieldbackground=[('focus', PALETTE['surface2'])])

    style.configure('TCombobox', fieldbackground=surf, background=surf,
                    foreground=txt, selectbackground=pri, font=TYPEFACE['normal'])
    style.map('TCombobox', fieldbackground=[('readonly', surf)])

    style.configure('Treeview', background=surf, fieldbackground=surf,
                    foreground=txt, rowheight=30, font=TYPEFACE['normal'], borderwidth=1)
    style.configure('Treeview.Heading', background=PALETTE['surface2'], foreground=txt,
                    font=('Verdana', 9, 'bold'), relief='flat')
    style.map('Treeview', background=[('selected', pri)], foreground=[('selected', PALETTE['white'])])

    style.configure('TScrollbar', background=PALETTE['surface2'], troughcolor=bg,
                    arrowcolor=dim, borderwidth=0)
    style.configure('TNotebook', background=bg, borderwidth=0)
    style.configure('TNotebook.Tab', background=PALETTE['surface2'], foreground=dim,
                    padding=[14, 7], font=TYPEFACE['normal'])
    style.map('TNotebook.Tab',
              background=[('selected', surf)],
              foreground=[('selected', pri)])

    style.configure('TCheckbutton', background=bg, foreground=txt, font=TYPEFACE['normal'])
    style.configure('Surface.TCheckbutton', background=surf, foreground=txt, font=TYPEFACE['normal'])

    style.configure('Primary.TButton', background=pri, foreground=PALETTE['white'],
                    font=TYPEFACE['normal'], padding=[12, 6], borderwidth=0)
    style.map('Primary.TButton',
              background=[('active', PALETTE['primary_light']), ('disabled', dim)])

    style.configure('Secondary.TButton', background=PALETTE['surface2'], foreground=txt,
                    font=TYPEFACE['normal'], padding=[12, 6], borderwidth=0)
    style.map('Secondary.TButton', background=[('active', PALETTE['border'])])

    style.configure('Danger.TButton', background=PALETTE['danger'], foreground=PALETTE['white'],
                    font=TYPEFACE['normal'], padding=[12, 6], borderwidth=0)
    style.map('Danger.TButton', background=[('active', '#A93226')])

    style.configure('Side.TButton', background=PALETTE['sidebar'], foreground=PALETTE['sidebar_text'],
                    font=TYPEFACE['normal'], padding=[10, 8], borderwidth=0, anchor='w')
    style.map('Side.TButton',
              background=[('active', PALETTE['sidebar_active'])],
              foreground=[('active', PALETTE['white'])])

    style.configure('SideActive.TButton', background=PALETTE['sidebar_active'],
                    foreground=PALETTE['white'], font=('Verdana', 10, 'bold'),
                    padding=[10, 8], borderwidth=0, anchor='w')

    style.configure('Accent.TButton', background=PALETTE['accent'], foreground=PALETTE['white'],
                    font=TYPEFACE['normal'], padding=[12, 6], borderwidth=0)
    style.map('Accent.TButton', background=[('active', PALETTE['accent_light'])])


# ---------------------------------------------------------------------------
# Вспомогательные виджеты
# ---------------------------------------------------------------------------

def styled_entry(parent, textvariable=None, width=20, **kw):
    e = tk.Entry(parent, textvariable=textvariable, width=width,
                 bg=PALETTE['surface'], fg=PALETTE['text'], insertbackground=PALETTE['text'],
                 relief='solid', bd=1, font=TYPEFACE['normal'],
                 highlightthickness=1, highlightbackground=PALETTE['border'],
                 highlightcolor=PALETTE['primary'],
                 disabledbackground=PALETTE['surface2'], disabledforeground=PALETTE['text_dim'], **kw)
    return e


def styled_combo(parent, values, textvariable=None, width=18, **kw):
    return ttk.Combobox(parent, values=values, textvariable=textvariable,
                        width=width, state='readonly', **kw)


def panel_box(parent, **kw):
    return tk.Frame(parent, bg=PALETTE['surface'], bd=0, highlightthickness=1,
                    highlightbackground=PALETTE['border'], **kw)


def divider_heading(parent, text):
    row = tk.Frame(parent, bg=PALETTE['bg'])
    row.pack(fill='x', pady=(16, 6))
    tk.Label(row, text=text, bg=PALETTE['bg'], fg=PALETTE['primary'],
             font=TYPEFACE['heading']).pack(side='left')
    tk.Frame(row, bg=PALETTE['border'], height=1).pack(side='left', fill='x', expand=True, padx=(10, 0))


def page_card(parent, title, subtitle=''):
    """Карточка-обёртка для содержимого раздела."""
    wrap = tk.Frame(parent, bg=PALETTE['bg'])
    wrap.pack(fill='both', expand=True, padx=24, pady=20)
    card = panel_box(wrap)
    card.pack(fill='both', expand=True)
    head = tk.Frame(card, bg=PALETTE['surface2'])
    head.pack(fill='x')
    tk.Label(head, text=title, bg=PALETTE['surface2'], fg=PALETTE['text'],
             font=TYPEFACE['heading']).pack(side='left', padx=16, pady=12)
    if subtitle:
        tk.Label(head, text=subtitle, bg=PALETTE['surface2'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(side='left', padx=(0, 16), pady=12)
    body = tk.Frame(card, bg=PALETTE['surface'])
    body.pack(fill='both', expand=True)
    return wrap, body


def nav_tile(parent, title, desc, command):
    """Плитка быстрого перехода на главной."""
    t = tk.Frame(parent, bg=PALETTE['surface'], cursor='hand2',
                 highlightthickness=1, highlightbackground=PALETTE['border'])
    t.pack(side='left', padx=(0, 12), pady=4, ipadx=14, ipady=12)
    tk.Label(t, text=title, bg=PALETTE['surface'], fg=PALETTE['primary'],
             font=TYPEFACE['heading']).pack(anchor='w', padx=12, pady=(10, 2))
    tk.Label(t, text=desc, bg=PALETTE['surface'], fg=PALETTE['text_dim'],
             font=TYPEFACE['small'], wraplength=160, justify='left').pack(
        anchor='w', padx=12, pady=(0, 10))
    for w in (t, *t.winfo_children()):
        w.bind('<Button-1>', lambda e, cmd=command: cmd())
        w.bind('<Enter>', lambda e, f=t: f.config(highlightbackground=PALETTE['primary']))
        w.bind('<Leave>', lambda e, f=t: f.config(highlightbackground=PALETTE['border']))
    return t


def action_btn(parent, text, command, kind='primary', **kw):
    styles = {
        'primary': 'Primary.TButton', 'secondary': 'Secondary.TButton',
        'danger': 'Danger.TButton', 'accent': 'Accent.TButton',
    }
    return ttk.Button(parent, text=text, command=command,
                      style=styles.get(kind, 'Primary.TButton'), **kw)


class ScrollPanel(ttk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        canvas = tk.Canvas(self, bg=PALETTE['bg'], highlightthickness=0)
        sb = ttk.Scrollbar(self, orient='vertical', command=canvas.yview)
        self.inner = ttk.Frame(canvas)
        self.inner.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0, 0), window=self.inner, anchor='nw')
        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        canvas.bind('<MouseWheel>', lambda e: canvas.yview_scroll(-1 * (e.delta // 120), 'units'))


class LabeledField(ttk.Frame):
    def __init__(self, parent, label, var=None, width=280, **kw):
        super().__init__(parent, style='Surface.TFrame', **kw)
        ttk.Label(self, text=label, style='Dim.TLabel',
                  background=PALETTE['surface']).pack(anchor='w', pady=(0, 3))
        self.var = var or tk.StringVar()
        ttk.Entry(self, textvariable=self.var, width=width // 8).pack(fill='x', ipady=4)

    def get(self):
        return self.var.get()

    def set(self, v):
        self.var.set(v)


# ---------------------------------------------------------------------------
# Базовый диалог
# ---------------------------------------------------------------------------

class ModalForm(tk.Toplevel):
    def __init__(self, parent, title, width=560, height=600):
        super().__init__(parent)
        self.title(title)
        self.configure(bg=PALETTE['bg'])
        self.geometry(f'{width}x{height}')
        self.resizable(True, True)
        self.grab_set()
        self.result = None
        self._build()

    def _build(self):
        pass

    def _ok(self):
        if self._validate():
            self.result = self._collect()
            self.destroy()

    def _validate(self):
        return True

    def _collect(self):
        return {}

    def _footer(self, parent):
        bar = tk.Frame(parent, bg=PALETTE['bg'])
        bar.pack(fill='x', pady=12, padx=16)
        action_btn(bar, 'Сохранить', self._ok, 'primary').pack(side='right', padx=(8, 0))
        action_btn(bar, 'Отмена', self.destroy, 'secondary').pack(side='right')
        return bar


# ---------------------------------------------------------------------------
# Базовая панель каталога (CRUD)
# ---------------------------------------------------------------------------

class CatalogPanel(ttk.Frame):
    COLUMNS = []
    TITLE = 'Каталог'
    SUBTITLE = ''

    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self.configure(style='TFrame')
        self._build()
        self.load()

    def _build(self):
        _, body = page_card(self, self.TITLE, self.SUBTITLE)

        tb = tk.Frame(body, bg=PALETTE['surface'])
        tb.pack(fill='x', padx=16, pady=(12, 8))
        btns = tk.Frame(tb, bg=PALETTE['surface'])
        btns.pack(side='right')
        action_btn(btns, '+ Новая', self._add, 'primary').pack(side='left', padx=3)
        action_btn(btns, 'Изменить', self._edit, 'secondary').pack(side='left', padx=3)
        action_btn(btns, 'Удалить', self._delete, 'danger').pack(side='left', padx=3)

        sf = tk.Frame(tb, bg=PALETTE['surface'])
        sf.pack(side='left', fill='x', expand=True)
        self._search_var = tk.StringVar()
        self._search_var.trace_add('write', lambda *_: self.load())
        styled_entry(sf, textvariable=self._search_var, width=36).pack(side='left', ipady=4)
        tk.Label(sf, text='  Фильтр', bg=PALETTE['surface'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(side='left')

        tf = tk.Frame(body, bg=PALETTE['surface'])
        tf.pack(fill='both', expand=True, padx=16, pady=(0, 16))
        cols = [c[0] for c in self.COLUMNS]
        self.tree = ttk.Treeview(tf, columns=cols, show='headings', selectmode='browse', height=16)
        for key, disp, w in self.COLUMNS:
            self.tree.heading(key, text=disp)
            self.tree.column(key, width=w, minwidth=60)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self.tree.bind('<Double-1>', lambda e: self._edit())

    def load(self):
        q = self._search_var.get().lower() if hasattr(self, '_search_var') else ''
        self.tree.delete(*self.tree.get_children())
        for row in self._fetch_rows():
            vals = [str(row[c[0]] or '') for c in self.COLUMNS]
            if q and not any(q in v.lower() for v in vals):
                continue
            self.tree.insert('', 'end', iid=str(row['id']), values=vals)

    def _fetch_rows(self):
        return []

    def _add(self):
        pass

    def _edit(self):
        pass

    def _delete(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите запись для удаления')
            return
        if messagebox.askyesno('Удаление', 'Удалить выбранную запись?'):
            self._do_delete(int(sel[0]))
            self.load()

    def _do_delete(self, row_id):
        pass

    def _selected_id(self):
        sel = self.tree.selection()
        return int(sel[0]) if sel else None


# ---------------------------------------------------------------------------
# Сотрудники
# ---------------------------------------------------------------------------

class StaffEditor(ModalForm):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Сотрудник', width=500, height=500)

    def _build(self):
        tk.Label(self, text='Карточка сотрудника', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=PALETTE['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, var, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            styled_entry(f, textvariable=var, width=32).grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_last = tk.StringVar(value=self.data.get('last_name', ''))
        self.v_first = tk.StringVar(value=self.data.get('first_name', ''))
        self.v_mid = tk.StringVar(value=self.data.get('middle_name', ''))
        self.v_pos = tk.StringVar(value=self.data.get('position', ''))
        self.v_sal = tk.StringVar(value=str(self.data.get('salary', '')))
        self.v_tax = tk.StringVar(value=str(self.data.get('tax_rate', '30.2')))

        row('Фамилия *', self.v_last, 0)
        row('Имя *', self.v_first, 1)
        row('Отчество', self.v_mid, 2)
        row('Должность', self.v_pos, 3)
        row('Оклад (руб/мес) *', self.v_sal, 4)
        row('Налоговая ставка % *', self.v_tax, 5)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_last.get() or not self.v_first.get():
            messagebox.showwarning('Ошибка', 'Заполните Фамилию и Имя')
            return False
        try:
            float(self.v_sal.get())
            float(self.v_tax.get())
        except ValueError:
            messagebox.showwarning('Ошибка', 'Оклад и ставка — числа')
            return False
        return True

    def _collect(self):
        return dict(last_name=self.v_last.get(), first_name=self.v_first.get(),
                    middle_name=self.v_mid.get(), position=self.v_pos.get(),
                    salary=float(self.v_sal.get()), tax_rate=float(self.v_tax.get()))


class StaffCatalog(CatalogPanel):
    TITLE = 'Штат'
    SUBTITLE = 'Сотрудники организации для расчёта трудозатрат'
    COLUMNS = [('id', '#', 40), ('fio', 'ФИО', 200), ('position', 'Должность', 160),
               ('salary', 'Оклад', 100), ('tax_rate', 'Нал. ставка %', 100)]

    def _fetch_rows(self):
        conn = connect_db()
        rows = conn.execute(
            "SELECT *, last_name||' '||first_name||COALESCE(' '||middle_name,'') as fio "
            "FROM employees ORDER BY last_name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = StaffEditor(self)
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "INSERT INTO employees(last_name,first_name,middle_name,position,salary,tax_rate) "
                "VALUES(?,?,?,?,?,?)",
                (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                 d.result['position'], d.result['salary'], d.result['tax_rate']))
            conn.commit()
            conn.close()
            self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid:
            messagebox.showwarning('Выбор', 'Выберите запись')
            return
        conn = connect_db()
        row = conn.execute("SELECT * FROM employees WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = StaffEditor(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "UPDATE employees SET last_name=?,first_name=?,middle_name=?,position=?,salary=?,tax_rate=? "
                "WHERE id=?",
                (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                 d.result['position'], d.result['salary'], d.result['tax_rate'], rid))
            conn.commit()
            conn.close()
            self.load()

    def _do_delete(self, rid):
        conn = connect_db()
        conn.execute("DELETE FROM employees WHERE id=?", (rid,))
        conn.commit()
        conn.close()


# ---------------------------------------------------------------------------
# Исполнители (физ. лица)
# ---------------------------------------------------------------------------

class FreelancerEditor(ModalForm):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Исполнитель', width=500, height=520)

    def _build(self):
        tk.Label(self, text='Исполнитель (физ. лицо)', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=PALETTE['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_last = tk.StringVar(value=self.data.get('last_name', ''))
        self.v_first = tk.StringVar(value=self.data.get('first_name', ''))
        self.v_mid = tk.StringVar(value=self.data.get('middle_name', ''))
        self.v_ctype = tk.StringVar(value=self.data.get('contract_type', 'ГПХ'))
        self.v_tax = tk.StringVar(value=str(self.data.get('tax_rate', '13')))
        self.v_unit = tk.StringVar(value=self.data.get('unit', 'часы'))
        self.v_rate = tk.StringVar(value=str(self.data.get('rate', '')))

        row('Фамилия *', styled_entry(f, textvariable=self.v_last, width=28), 0)
        row('Имя *', styled_entry(f, textvariable=self.v_first, width=28), 1)
        row('Отчество', styled_entry(f, textvariable=self.v_mid, width=28), 2)

        cb_type = styled_combo(f, ['ГПХ', 'НПД'], textvariable=self.v_ctype, width=12)
        cb_type.bind('<<ComboboxSelected>>', self._on_type)
        row('Тип оформления *', cb_type, 3)

        self.tax_entry = styled_entry(f, textvariable=self.v_tax, width=12)
        row('Налоговая ставка %', self.tax_entry, 4)
        row('Единица', styled_combo(f, ['часы', 'дни', 'полная стоимость'],
                                    textvariable=self.v_unit, width=16), 5)
        row('Ставка за ед.', styled_entry(f, textvariable=self.v_rate, width=16), 6)
        f.columnconfigure(1, weight=1)
        self._on_type()
        self._footer(self)

    def _on_type(self, e=None):
        if self.v_ctype.get() == 'НПД':
            self.v_tax.set('0')
            self.tax_entry.config(state='disabled')
        else:
            self.tax_entry.config(state='normal')

    def _validate(self):
        if not self.v_last.get() or not self.v_first.get():
            messagebox.showwarning('Ошибка', 'Заполните Фамилию и Имя')
            return False
        return True

    def _collect(self):
        return dict(last_name=self.v_last.get(), first_name=self.v_first.get(),
                    middle_name=self.v_mid.get(), contract_type=self.v_ctype.get(),
                    tax_rate=float(self.v_tax.get() or 0),
                    unit=self.v_unit.get(), rate=float(self.v_rate.get() or 0))


class FreelancerCatalog(CatalogPanel):
    TITLE = 'Внешние исполнители'
    SUBTITLE = 'Физлица по договорам ГПХ и НПД'
    COLUMNS = [('id', '#', 40), ('fio', 'ФИО', 200), ('contract_type', 'Тип', 80),
               ('unit', 'Ед.', 80), ('rate', 'Ставка', 100), ('tax_rate', 'Нал. %', 80)]

    def _fetch_rows(self):
        conn = connect_db()
        rows = conn.execute(
            "SELECT *, last_name||' '||first_name||COALESCE(' '||middle_name,'') as fio "
            "FROM contractors ORDER BY last_name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = FreelancerEditor(self)
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "INSERT INTO contractors(last_name,first_name,middle_name,contract_type,tax_rate,unit,rate) "
                "VALUES(?,?,?,?,?,?,?)",
                tuple(d.result[k] for k in ['last_name', 'first_name', 'middle_name',
                                            'contract_type', 'tax_rate', 'unit', 'rate']))
            conn.commit()
            conn.close()
            self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid:
            messagebox.showwarning('Выбор', 'Выберите запись')
            return
        conn = connect_db()
        row = conn.execute("SELECT * FROM contractors WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = FreelancerEditor(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "UPDATE contractors SET last_name=?,first_name=?,middle_name=?,contract_type=?,"
                "tax_rate=?,unit=?,rate=? WHERE id=?",
                (*[d.result[k] for k in ['last_name', 'first_name', 'middle_name',
                                        'contract_type', 'tax_rate', 'unit', 'rate']], rid))
            conn.commit()
            conn.close()
            self.load()

    def _do_delete(self, rid):
        conn = connect_db()
        conn.execute("DELETE FROM contractors WHERE id=?", (rid,))
        conn.commit()
        conn.close()


# ---------------------------------------------------------------------------
# Субподрядчики
# ---------------------------------------------------------------------------

class VendorEditor(ModalForm):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Субподрядчик', width=500, height=560)

    def _build(self):
        tk.Label(self, text='Субподрядчик (юр. лицо / ИП)', bg=PALETTE['bg'],
                 fg=PALETTE['text'], font=TYPEFACE['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=PALETTE['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_inn = tk.StringVar(value=self.data.get('inn', ''))
        self.v_type = tk.StringVar(value=self.data.get('type', 'ООО'))
        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_dir = tk.StringVar(value=self.data.get('director_name', ''))
        self.v_email = tk.StringVar(value=self.data.get('email', ''))
        self.v_phone = tk.StringVar(value=self.data.get('phone', ''))
        self.v_unit = tk.StringVar(value=self.data.get('unit', 'часы'))
        self.v_rate = tk.StringVar(value=str(self.data.get('rate', '')))

        row('ИНН', styled_entry(f, textvariable=self.v_inn, width=20), 0)
        row('Тип', styled_combo(f, ['ООО', 'АО', 'ИП', 'ПАО'], textvariable=self.v_type, width=12), 1)
        row('Название *', styled_entry(f, textvariable=self.v_name, width=28), 2)
        row('Руководитель', styled_entry(f, textvariable=self.v_dir, width=28), 3)
        row('Email', styled_entry(f, textvariable=self.v_email, width=28), 4)
        row('Телефон', styled_entry(f, textvariable=self.v_phone, width=20), 5)
        row('Единица', styled_combo(f, ['часы', 'дни', 'полная стоимость'],
                                    textvariable=self.v_unit, width=16), 6)
        row('Ставка за ед.', styled_entry(f, textvariable=self.v_rate, width=16), 7)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Заполните Название')
            return False
        return True

    def _collect(self):
        return dict(inn=self.v_inn.get(), type=self.v_type.get(), name=self.v_name.get(),
                    director_name=self.v_dir.get(), email=self.v_email.get(),
                    phone=self.v_phone.get(), unit=self.v_unit.get(),
                    rate=float(self.v_rate.get() or 0))


class VendorCatalog(CatalogPanel):
    TITLE = 'Партнёры'
    SUBTITLE = 'Субподрядные организации и ИП'
    COLUMNS = [('id', '#', 40), ('type', 'Тип', 60), ('name', 'Название', 220),
               ('director_name', 'Руководитель', 160), ('phone', 'Телефон', 120)]

    def _fetch_rows(self):
        conn = connect_db()
        rows = conn.execute("SELECT * FROM subcontractors ORDER BY name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = VendorEditor(self)
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "INSERT INTO subcontractors(inn,type,name,director_name,email,phone,unit,rate) "
                "VALUES(?,?,?,?,?,?,?,?)",
                tuple(d.result[k] for k in ['inn', 'type', 'name', 'director_name',
                                            'email', 'phone', 'unit', 'rate']))
            conn.commit()
            conn.close()
            self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid:
            messagebox.showwarning('Выбор', 'Выберите запись')
            return
        conn = connect_db()
        row = conn.execute("SELECT * FROM subcontractors WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = VendorEditor(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "UPDATE subcontractors SET inn=?,type=?,name=?,director_name=?,email=?,phone=?,unit=?,rate=? "
                "WHERE id=?",
                (*[d.result[k] for k in ['inn', 'type', 'name', 'director_name',
                                        'email', 'phone', 'unit', 'rate']], rid))
            conn.commit()
            conn.close()
            self.load()

    def _do_delete(self, rid):
        conn = connect_db()
        conn.execute("DELETE FROM subcontractors WHERE id=?", (rid,))
        conn.commit()
        conn.close()


# ---------------------------------------------------------------------------
# Оборудование
# ---------------------------------------------------------------------------

class AssetEditor(ModalForm):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Оборудование', width=500, height=480)

    def _build(self):
        tk.Label(self, text='Оборудование', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=PALETTE['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_desc = tk.StringVar(value=self.data.get('description', ''))
        self.v_own = tk.StringVar(value=self.data.get('ownership', 'Собственное'))
        self.v_rent = tk.StringVar(value=str(self.data.get('rental_cost', '0')))
        self.v_unit = tk.StringVar(value=self.data.get('unit', 'часы'))
        self.v_rate = tk.StringVar(value=str(self.data.get('rate', '')))

        row('Название *', styled_entry(f, textvariable=self.v_name, width=28), 0)
        row('Описание', styled_entry(f, textvariable=self.v_desc, width=28), 1)
        row('Тип владения', styled_combo(f, ['Собственное', 'В аренде'],
                                         textvariable=self.v_own, width=16), 2)
        row('Эксплуат. стоимость', styled_entry(f, textvariable=self.v_rent, width=16), 3)
        row('Единица', styled_combo(f, ['часы', 'дни', 'полная стоимость'],
                                    textvariable=self.v_unit, width=16), 4)
        row('Ставка за ед.', styled_entry(f, textvariable=self.v_rate, width=16), 5)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Введите Название')
            return False
        return True

    def _collect(self):
        return dict(name=self.v_name.get(), description=self.v_desc.get(),
                    ownership=self.v_own.get(), rental_cost=float(self.v_rent.get() or 0),
                    unit=self.v_unit.get(), rate=float(self.v_rate.get() or 0))


class AssetCatalog(CatalogPanel):
    TITLE = 'Техника'
    SUBTITLE = 'Оборудование и активы проекта'
    COLUMNS = [('id', '#', 40), ('name', 'Название', 200), ('ownership', 'Тип', 100),
               ('unit', 'Ед.', 80), ('rate', 'Ставка', 100)]

    def _fetch_rows(self):
        conn = connect_db()
        rows = conn.execute("SELECT * FROM equipment ORDER BY name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = AssetEditor(self)
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "INSERT INTO equipment(name,description,ownership,rental_cost,unit,rate) VALUES(?,?,?,?,?,?)",
                tuple(d.result[k] for k in ['name', 'description', 'ownership', 'rental_cost', 'unit', 'rate']))
            conn.commit()
            conn.close()
            self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid:
            messagebox.showwarning('Выбор', 'Выберите запись')
            return
        conn = connect_db()
        row = conn.execute("SELECT * FROM equipment WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = AssetEditor(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "UPDATE equipment SET name=?,description=?,ownership=?,rental_cost=?,unit=?,rate=? WHERE id=?",
                (*[d.result[k] for k in ['name', 'description', 'ownership', 'rental_cost', 'unit', 'rate']], rid))
            conn.commit()
            conn.close()
            self.load()

    def _do_delete(self, rid):
        conn = connect_db()
        conn.execute("DELETE FROM equipment WHERE id=?", (rid,))
        conn.commit()
        conn.close()


# ---------------------------------------------------------------------------
# Клиенты (заказчики)
# ---------------------------------------------------------------------------

class ClientEditor(ModalForm):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Клиент', width=500, height=500)

    def _build(self):
        tk.Label(self, text='Клиент', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['title']).pack(pady=(20, 10))
        f = tk.Frame(self, bg=PALETTE['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='w', pady=4)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=4)

        self.v_inn = tk.StringVar(value=self.data.get('inn', ''))
        self.v_type = tk.StringVar(value=self.data.get('type', 'Юридическое лицо'))
        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_dir = tk.StringVar(value=self.data.get('director_name', ''))
        self.v_email = tk.StringVar(value=self.data.get('email', ''))
        self.v_phone = tk.StringVar(value=self.data.get('phone', ''))

        row('ИНН', styled_entry(f, textvariable=self.v_inn, width=20), 0)
        row('Тип', styled_combo(f, ['Физическое лицо', 'Индивидуальный предприниматель', 'Юридическое лицо'],
                                textvariable=self.v_type, width=28), 1)
        row('Название / ФИО *', styled_entry(f, textvariable=self.v_name, width=28), 2)
        row('ФИО руководителя', styled_entry(f, textvariable=self.v_dir, width=28), 3)
        row('Email', styled_entry(f, textvariable=self.v_email, width=28), 4)
        row('Телефон', styled_entry(f, textvariable=self.v_phone, width=20), 5)
        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Введите Название / ФИО')
            return False
        return True

    def _collect(self):
        return dict(inn=self.v_inn.get(), type=self.v_type.get(), name=self.v_name.get(),
                    director_name=self.v_dir.get(), email=self.v_email.get(), phone=self.v_phone.get())


class ClientCatalog(CatalogPanel):
    TITLE = 'Контрагенты'
    SUBTITLE = 'Заказчики и клиенты для коммерческих предложений'
    COLUMNS = [('id', '#', 40), ('type', 'Тип', 120), ('name', 'Название / ФИО', 200),
               ('director_name', 'Руководитель', 160), ('email', 'Email', 160)]

    def _fetch_rows(self):
        conn = connect_db()
        rows = conn.execute("SELECT * FROM customers ORDER BY name").fetchall()
        conn.close()
        return rows

    def _add(self):
        d = ClientEditor(self)
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "INSERT INTO customers(inn,type,name,director_name,email,phone) VALUES(?,?,?,?,?,?)",
                tuple(d.result[k] for k in ['inn', 'type', 'name', 'director_name', 'email', 'phone']))
            conn.commit()
            conn.close()
            self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid:
            messagebox.showwarning('Выбор', 'Выберите запись')
            return
        conn = connect_db()
        row = conn.execute("SELECT * FROM customers WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = ClientEditor(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "UPDATE customers SET inn=?,type=?,name=?,director_name=?,email=?,phone=? WHERE id=?",
                (*[d.result[k] for k in ['inn', 'type', 'name', 'director_name', 'email', 'phone']], rid))
            conn.commit()
            conn.close()
            self.load()

    def _do_delete(self, rid):
        conn = connect_db()
        conn.execute("DELETE FROM customers WHERE id=?", (rid,))
        conn.commit()
        conn.close()


# ---------------------------------------------------------------------------
# Сессия пользователя и роли
# ---------------------------------------------------------------------------

ALL_ROLES = ['Глобальный администратор', 'Коммерческий директор', 'Бухгалтер', 'Кадровик']

MARGIN_VISIBLE_ROLES = {'Глобальный администратор', 'Коммерческий директор'}


class UserSession:
    current_user = None

    @classmethod
    def set_user(cls, user_row):
        cls.current_user = dict(user_row) if user_row is not None else None

    @classmethod
    def get_roles(cls):
        if not cls.current_user:
            return []
        roles = cls.current_user.get('roles', '[]')
        if isinstance(roles, str):
            try:
                roles = json.loads(roles)
            except Exception:
                roles = []
        return roles or []

    @classmethod
    def can_see_margin(cls):
        if not cls.current_user:
            return True
        roles = set(cls.get_roles())
        return bool(roles & MARGIN_VISIBLE_ROLES)

    @classmethod
    def display_name(cls):
        if not cls.current_user:
            return 'Гость'
        u = cls.current_user
        return f"{u.get('last_name', '')} {u.get('first_name', '')}".strip() or 'Пользователь'


# ---------------------------------------------------------------------------
# Расчёт стоимости
# ---------------------------------------------------------------------------

def _workdays_in_month(yr, mo):
    _, days_in_month = calendar.monthrange(yr, mo)
    return sum(1 for d in range(1, days_in_month + 1)
               if calendar.weekday(yr, mo, d) < 5)


def workdays_count(date_start, date_end):
    if not date_start or not date_end:
        return 0
    try:
        d1 = (datetime.strptime(date_start, '%Y-%m-%d').date()
              if isinstance(date_start, str) else date_start)
        d2 = (datetime.strptime(date_end, '%Y-%m-%d').date()
              if isinstance(date_end, str) else date_end)
    except Exception:
        return 0
    if d2 < d1:
        return 0
    days = (d2 - d1).days + 1
    full_weeks, rem = divmod(days, 7)
    count = full_weeks * 5
    for i in range(rem):
        d = d1 + timedelta(days=full_weeks * 7 + i)
        if d.weekday() < 5:
            count += 1
    return count


def calc_employee_cost(emp, units, date_start, date_end):
    zp = emp['salary']
    ns = emp['tax_rate'] / 100

    if not date_start or not date_end:
        if units and units > 0:
            return zp + zp * ns
        return 0.0

    try:
        d1 = datetime.strptime(date_start, '%Y-%m-%d').date()
        d2 = datetime.strptime(date_end, '%Y-%m-%d').date()
    except Exception:
        if units and units > 0:
            return zp + zp * ns
        return 0.0

    if d2 < d1:
        return 0.0

    total = 0.0
    cur = d1
    while cur <= d2:
        yr, mo = cur.year, cur.month
        cwdm = _workdays_in_month(yr, mo)
        if cwdm == 0:
            cwdm = 1
        day_rate = (zp + zp * ns) / cwdm
        _, days_in_month = calendar.monthrange(yr, mo)
        mo_end = date(yr, mo, days_in_month)
        seg_end = min(d2, mo_end)
        cwd = sum(1 for d in range(cur.day, seg_end.day + 1)
                  if calendar.weekday(yr, mo, d) < 5)
        total += day_rate * cwd
        if mo == 12:
            cur = date(yr + 1, 1, 1)
        else:
            cur = date(yr, mo + 1, 1)
        if cur > d2:
            break
    return total


def calc_contractor_cost(con, units):
    cez = units
    suz = con['rate']
    ns = con['tax_rate'] / 100
    if con['contract_type'] == 'НПД':
        return cez * suz
    return cez * suz + (cez * suz) * ns


def calc_equipment_cost(eq, units):
    return eq['rate'] * units


def calc_subcontractor_cost(sub, units):
    return sub['rate'] * units


class BudgetItemEditor(ModalForm):
    def __init__(self, parent, project_dates=None, data=None):
        self.project_dates = project_dates or ('', '')
        self.data = data or {}
        self._emp_map = {}
        self._con_map = {}
        self._sub_map = {}
        self._eq_map = {}
        super().__init__(parent, 'Ресурс проекта', width=560, height=620)

    def _build(self):
        tk.Label(self, text='Ресурс проекта', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['title']).pack(pady=(16, 8))
        f = tk.Frame(self, bg=PALETTE['bg'])
        f.pack(fill='both', expand=True, padx=20)

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='w', pady=5)
            widget.grid(row=row_n, column=1, sticky='ew', padx=(10, 0), pady=5)

        self.v_name = tk.StringVar(value=self.data.get('name', ''))
        self.v_rtype = tk.StringVar(value=self.data.get('resource_type', 'Сотрудник'))
        self.v_exec = tk.StringVar(value='')
        self.v_service = tk.StringVar(value=self.data.get('service', ''))
        self.v_ds = tk.StringVar(value=self.data.get('date_start', self.project_dates[0]))
        self.v_de = tk.StringVar(value=self.data.get('date_end', self.project_dates[1]))
        self.v_units = tk.StringVar(value=str(self.data.get('units', '1')))
        self.v_margin = tk.StringVar(value=str(self.data.get('margin_pct', '0')))

        row('Название в документах *', styled_entry(f, textvariable=self.v_name, width=30), 0)

        self.cb_type = styled_combo(f, ['Сотрудник', 'Исполнитель', 'Субподрядчик', 'Оборудование'],
                                    textvariable=self.v_rtype, width=20)
        self.cb_type.bind('<<ComboboxSelected>>', self._on_type)
        row('Тип ресурса *', self.cb_type, 1)

        self.cb_exec = styled_combo(f, [], textvariable=self.v_exec, width=26)
        row('Исполнитель *', self.cb_exec, 2)

        row('Название услуги', styled_entry(f, textvariable=self.v_service, width=30), 3)
        self.e_ds = styled_entry(f, textvariable=self.v_ds, width=14)
        row('Дата начала', self.e_ds, 4)
        self.e_de = styled_entry(f, textvariable=self.v_de, width=14)
        row('Дата окончания', self.e_de, 5)
        self.e_units = styled_entry(f, textvariable=self.v_units, width=10)
        self.lbl_units_row = tk.Label(f, text='Кол-во единиц (часы/дни)', bg=PALETTE['bg'],
                                      fg=PALETTE['text_dim'], font=TYPEFACE['small'])
        self.lbl_units_row.grid(row=6, column=0, sticky='w', pady=5)
        self.e_units.grid(row=6, column=1, sticky='ew', padx=(10, 0), pady=5)
        self.lbl_units_hint = tk.Label(f, text='', bg=PALETTE['bg'],
                                       fg=PALETTE['text_dim'], font=TYPEFACE['small'])
        self.lbl_units_hint.grid(row=6, column=1, sticky='e', padx=(0, 4))
        self._can_margin = UserSession.can_see_margin()
        self.lbl_margin_row = tk.Label(f, text='Маржинальность %', bg=PALETTE['bg'],
                                       fg=PALETTE['text_dim'], font=TYPEFACE['small'])
        self.e_margin = styled_entry(f, textvariable=self.v_margin, width=10)
        if self._can_margin:
            self.lbl_margin_row.grid(row=7, column=0, sticky='w', pady=5)
            self.e_margin.grid(row=7, column=1, sticky='ew', padx=(10, 0), pady=5)

        self.lbl_cost = tk.Label(f, text='Себестоимость: —', bg=PALETTE['bg'],
                                 fg=PALETTE['success'], font=TYPEFACE['heading'])
        self.lbl_cost.grid(row=8, column=0, columnspan=2, sticky='w', pady=(10, 0))

        f.columnconfigure(1, weight=1)
        self._load_executor_list()

        if self.data.get('executor_id'):
            self.v_exec.set(str(self.data.get('executor_display', '')))

        self.v_units.trace_add('write', lambda *_: self._update_cost())
        self.v_margin.trace_add('write', lambda *_: self._update_cost())
        self.v_ds.trace_add('write', lambda *_: self._sync_units())
        self.v_de.trace_add('write', lambda *_: self._sync_units())
        self.cb_exec.bind('<<ComboboxSelected>>', lambda e: self._update_cost())
        self._sync_units()
        self._update_cost()
        self._footer(self)

    def _on_type(self, e=None):
        self._load_executor_list()
        self._sync_units()
        self._update_cost()

    def _sync_units(self):
        if self.v_rtype.get() == 'Сотрудник':
            wd = workdays_count(self.v_ds.get(), self.v_de.get())
            self.v_units.set(str(wd))
            self.e_units.config(state='disabled')
            self.lbl_units_row.config(text='Кол-во рабочих дней (из интервала)')
            self.lbl_units_hint.config(text='авто' if wd else '')
        else:
            self.e_units.config(state='normal')
            self.lbl_units_row.config(text='Кол-во единиц (часы/дни)')
            self.lbl_units_hint.config(text='')

    def _load_executor_list(self):
        rtype = self.v_rtype.get()
        conn = connect_db()
        names = []
        if rtype == 'Сотрудник':
            rows = conn.execute("SELECT id, last_name||' '||first_name as n FROM employees").fetchall()
            self._emp_map = {r['n']: dict(r) for r in rows}
            names = list(self._emp_map.keys())
        elif rtype == 'Исполнитель':
            rows = conn.execute(
                "SELECT id, last_name||' '||first_name as n, contract_type, tax_rate, rate, unit "
                "FROM contractors").fetchall()
            self._con_map = {r['n']: dict(r) for r in rows}
            names = list(self._con_map.keys())
        elif rtype == 'Субподрядчик':
            rows = conn.execute("SELECT id, name as n, rate, unit FROM subcontractors").fetchall()
            self._sub_map = {r['n']: dict(r) for r in rows}
            names = list(self._sub_map.keys())
        elif rtype == 'Оборудование':
            rows = conn.execute("SELECT id, name as n, rate, unit FROM equipment").fetchall()
            self._eq_map = {r['n']: dict(r) for r in rows}
            names = list(self._eq_map.keys())
        conn.close()
        self.cb_exec['values'] = names
        if names:
            self.cb_exec.current(0)

    def _get_cost(self):
        try:
            units = float(self.v_units.get() or 0)
        except ValueError:
            return 0.0
        rtype = self.v_rtype.get()
        name = self.v_exec.get()
        if rtype == 'Сотрудник' and name in self._emp_map:
            conn = connect_db()
            emp = conn.execute("SELECT * FROM employees WHERE id=?",
                               (self._emp_map[name]['id'],)).fetchone()
            conn.close()
            return calc_employee_cost(dict(emp), units, self.v_ds.get(), self.v_de.get())
        if rtype == 'Исполнитель' and name in self._con_map:
            conn = connect_db()
            con = conn.execute("SELECT * FROM contractors WHERE id=?",
                               (self._con_map[name]['id'],)).fetchone()
            conn.close()
            return calc_contractor_cost(dict(con), units)
        if rtype == 'Субподрядчик' and name in self._sub_map:
            conn = connect_db()
            sub = conn.execute("SELECT * FROM subcontractors WHERE id=?",
                               (self._sub_map[name]['id'],)).fetchone()
            conn.close()
            return calc_subcontractor_cost(dict(sub), units)
        if rtype == 'Оборудование' and name in self._eq_map:
            conn = connect_db()
            eq = conn.execute("SELECT * FROM equipment WHERE id=?",
                              (self._eq_map[name]['id'],)).fetchone()
            conn.close()
            return calc_equipment_cost(dict(eq), units)
        return 0.0

    def _update_cost(self):
        try:
            cost = self._get_cost()
            if self._can_margin:
                try:
                    mp = float(self.v_margin.get() or 0) / 100
                except ValueError:
                    mp = 0.0
                total = cost * (1 + mp)
                self.lbl_cost.config(
                    text=f'Себестоимость: {cost:,.2f} ₽   |   Итого (с маржой): {total:,.2f} ₽')
            else:
                self.lbl_cost.config(text=f'Себестоимость: {cost:,.2f} ₽')
        except Exception:
            pass

    def _get_exec_id(self):
        rtype = self.v_rtype.get()
        name = self.v_exec.get()
        m = {'Сотрудник': self._emp_map, 'Исполнитель': self._con_map,
             'Субподрядчик': self._sub_map, 'Оборудование': self._eq_map}
        mapping = m.get(rtype, {})
        return mapping.get(name, {}).get('id')

    def _validate(self):
        if not self.v_name.get():
            messagebox.showwarning('Ошибка', 'Введите название ресурса')
            return False
        if not self.v_exec.get():
            messagebox.showwarning('Ошибка', 'Выберите исполнителя')
            return False
        if self.v_rtype.get() == 'Сотрудник' and (not self.v_ds.get() or not self.v_de.get()):
            messagebox.showwarning(
                'Ошибка',
                'Для типа «Сотрудник» укажите дату начала и окончания —\n'
                'количество дней и стоимость рассчитываются по интервалу.')
            return False
        return True

    def _collect(self):
        cost = self._get_cost()
        try:
            mp = float(self.v_margin.get() or 0)
        except ValueError:
            mp = 0.0
        total = cost * (1 + mp / 100)
        rtype = self.v_rtype.get()
        if rtype == 'Сотрудник':
            unit_label = 'дни'
        else:
            m = {'Исполнитель': self._con_map, 'Субподрядчик': self._sub_map,
                 'Оборудование': self._eq_map}.get(rtype, {})
            unit_label = m.get(self.v_exec.get(), {}).get('unit', '')
        return dict(
            name=self.v_name.get(),
            resource_type=rtype,
            executor_id=self._get_exec_id(),
            executor_display=self.v_exec.get(),
            service=self.v_service.get(),
            date_start=self.v_ds.get(),
            date_end=self.v_de.get(),
            units=float(self.v_units.get() or 1),
            unit_label=unit_label,
            margin_pct=mp,
            cost=cost,
            total=total,
        )


# ---------------------------------------------------------------------------
# Окно проекта
# ---------------------------------------------------------------------------

class ProjectWindow(tk.Toplevel):
    def __init__(self, parent, project_id=None, workspace_id=None, on_save=None):
        super().__init__(parent)
        self.project_id = project_id
        self.workspace_id = workspace_id
        self.on_save = on_save
        self.resources = []
        self.configure(bg=PALETTE['bg'])
        self.title(f'Карточка проекта — {APP_NAME}')
        self.geometry('960x780')
        self._load_project()
        self._build()
        self.grab_set()

    def _load_project(self):
        if self.project_id:
            conn = connect_db()
            row = conn.execute("SELECT * FROM projects WHERE id=?", (self.project_id,)).fetchone()
            conn.close()
            self.proj = dict(row)
            self.resources = json.loads(self.proj.get('resources', '[]'))
        else:
            self.proj = {'name': '', 'date_start': '', 'date_end': '',
                         'description': '', 'customer_id': None, 'tax_rate': 0, 'status': 'Активный'}

    def _build(self):
        body = tk.Frame(self, bg=PALETTE['bg'])
        body.pack(fill='both', expand=True)

        nav = tk.Frame(body, bg=PALETTE['sidebar'], width=200)
        nav.pack(side='left', fill='y')
        nav.pack_propagate(False)

        tk.Label(nav, text='Разделы', bg=PALETTE['sidebar'], fg=PALETTE['sidebar_section'],
                 font=TYPEFACE['section']).pack(anchor='w', padx=16, pady=(20, 8))

        self._proj_content = tk.Frame(body, bg=PALETTE['bg'])
        self._proj_content.pack(side='right', fill='both', expand=True)

        self._proj_pages = {}
        self._proj_nav_btns = {}

        for label in ('Общие сведения', 'Расчёт затрат'):
            btn = tk.Button(
                nav, text=label, anchor='w', relief='flat', cursor='hand2',
                bg=PALETTE['sidebar'], fg=PALETTE['sidebar_text'],
                activebackground=PALETTE['sidebar_active'], activeforeground=PALETTE['white'],
                font=TYPEFACE['normal'], padx=16, pady=10,
                command=lambda l=label: self._show_proj_page(l),
            )
            btn.pack(fill='x')
            self._proj_nav_btns[label] = btn

        page1 = tk.Frame(self._proj_content, bg=PALETTE['bg'])
        page2 = tk.Frame(self._proj_content, bg=PALETTE['bg'])
        self._proj_pages['Общие сведения'] = page1
        self._proj_pages['Расчёт затрат'] = page2

        self._build_details(page1)
        self._build_resources(page2)
        self._show_proj_page('Общие сведения')
        self._build_bottom()

    def _show_proj_page(self, label):
        for name, frame in self._proj_pages.items():
            frame.pack_forget()
        self._proj_pages[label].pack(fill='both', expand=True)
        for name, btn in self._proj_nav_btns.items():
            if name == label:
                btn.config(bg=PALETTE['sidebar_active'], fg=PALETTE['white'],
                           font=TYPEFACE['heading'])
            else:
                btn.config(bg=PALETTE['sidebar'], fg=PALETTE['sidebar_text'],
                           font=TYPEFACE['normal'])

    def _build_details(self, parent):
        sf = ScrollPanel(parent)
        sf.pack(fill='both', expand=True)
        f = sf.inner
        f.configure(style='TFrame')
        pad = {'padx': 20, 'pady': 6}

        divider_heading(f, 'Основная информация')

        row_f = ttk.Frame(f)
        row_f.pack(fill='x', **pad)
        ttk.Label(row_f, text='Название *').pack(side='left', padx=(0, 8))
        self.v_name = tk.StringVar(value=self.proj.get('name', ''))
        styled_entry(row_f, textvariable=self.v_name, width=44).pack(side='left')

        row2 = ttk.Frame(f)
        row2.pack(fill='x', **pad)
        ttk.Label(row2, text='Начало').pack(side='left', padx=(0, 8))
        self.v_ds = tk.StringVar(value=self.proj.get('date_start', ''))
        styled_entry(row2, textvariable=self.v_ds, width=14).pack(side='left', padx=(0, 20))
        ttk.Label(row2, text='Окончание').pack(side='left', padx=(0, 8))
        self.v_de = tk.StringVar(value=self.proj.get('date_end', ''))
        styled_entry(row2, textvariable=self.v_de, width=14).pack(side='left')

        row3 = ttk.Frame(f)
        row3.pack(fill='x', **pad)
        ttk.Label(row3, text='Статус').pack(side='left', padx=(0, 8))
        self.v_status = tk.StringVar(value=self.proj.get('status', 'Активный'))
        styled_combo(row3, ['Активный', 'Завершён', 'Приостановлен'],
                     textvariable=self.v_status, width=18).pack(side='left')

        divider_heading(f, 'Описание')
        desc_f = tk.Frame(f, bg=PALETTE['bg'])
        desc_f.pack(fill='x', padx=20, pady=4)
        self.txt_desc = tk.Text(desc_f, height=4, bg=PALETTE['surface'], fg=PALETTE['text'],
                                font=TYPEFACE['normal'], relief='solid', bd=1,
                                insertbackground=PALETTE['text'],
                                highlightthickness=1, highlightbackground=PALETTE['border'])
        self.txt_desc.insert('1.0', self.proj.get('description', ''))
        self.txt_desc.pack(fill='x')

        divider_heading(f, 'Финансы')

        row4 = ttk.Frame(f)
        row4.pack(fill='x', **pad)
        ttk.Label(row4, text='Клиент').pack(side='left', padx=(0, 8))

        conn = connect_db()
        cust_rows = conn.execute("SELECT id, name FROM customers ORDER BY name").fetchall()
        conn.close()
        self._cust_map = {r['name']: r['id'] for r in cust_rows}
        cust_names = ['(не выбран)'] + list(self._cust_map.keys())
        self.v_cust = tk.StringVar()
        if self.proj.get('customer_id'):
            conn = connect_db()
            cr = conn.execute("SELECT name FROM customers WHERE id=?",
                              (self.proj['customer_id'],)).fetchone()
            conn.close()
            if cr:
                self.v_cust.set(cr['name'])
        if not self.v_cust.get():
            self.v_cust.set('(не выбран)')
        styled_combo(row4, cust_names, textvariable=self.v_cust, width=30).pack(side='left')

        row5 = ttk.Frame(f)
        row5.pack(fill='x', **pad)
        ttk.Label(row5, text='Налоговая ставка %').pack(side='left', padx=(0, 8))
        self.v_tax = tk.StringVar(value=str(self.proj.get('tax_rate', '0')))
        styled_entry(row5, textvariable=self.v_tax, width=8).pack(side='left')

        divider_heading(f, 'Итоги (авторасчёт)')
        self.lbl_summary = tk.Label(f, text='', bg=PALETTE['bg'], fg=PALETTE['text'],
                                    font=TYPEFACE['normal'], justify='left')
        self.lbl_summary.pack(anchor='w', padx=20, pady=4)
        self._update_summary()

    def _build_resources(self, parent):
        hdr = tk.Frame(parent, bg=PALETTE['bg'])
        hdr.pack(fill='x', padx=20, pady=(16, 8))
        tk.Label(hdr, text='Строки сметы', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(side='left')
        action_btn(hdr, 'Добавить', self._add_resource, 'primary').pack(side='right', padx=4)
        action_btn(hdr, 'Удалить', self._del_resource, 'danger').pack(side='right', padx=4)
        action_btn(hdr, 'Изменить', self._edit_resource, 'secondary').pack(side='right', padx=4)

        can_margin = UserSession.can_see_margin()
        if can_margin:
            cols = ('name', 'type', 'executor', 'service', 'units', 'cost', 'margin', 'total')
            hdrs = ('Название', 'Тип', 'Исполнитель', 'Услуга', 'Ед.', 'Себест.', 'Маржа %', 'Итого')
            widths = (160, 100, 140, 120, 60, 100, 80, 110)
        else:
            cols = ('name', 'type', 'executor', 'service', 'units', 'cost')
            hdrs = ('Название', 'Тип', 'Исполнитель', 'Услуга', 'Ед.', 'Себест.')
            widths = (160, 100, 140, 120, 60, 100)
        tf = tk.Frame(parent, bg=PALETTE['bg'])
        tf.pack(fill='both', expand=True, padx=20, pady=(0, 10))
        self.res_tree = ttk.Treeview(tf, columns=cols, show='headings')
        for c, h, w in zip(cols, hdrs, widths):
            self.res_tree.heading(c, text=h)
            self.res_tree.column(c, width=w, minwidth=50)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.res_tree.yview)
        self.res_tree.configure(yscrollcommand=sb.set)
        self.res_tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self.res_tree.bind('<Double-1>', lambda e: self._edit_resource())
        self._reload_res_tree()

    def _reload_res_tree(self):
        self.res_tree.delete(*self.res_tree.get_children())
        can_margin = UserSession.can_see_margin()
        for i, r in enumerate(self.resources):
            base_vals = [r.get('name', ''), r.get('resource_type', ''), r.get('executor_display', ''),
                         r.get('service', ''), r.get('units', ''), f"{r.get('cost', 0):,.2f}"]
            if can_margin:
                base_vals += [f"{r.get('margin_pct', 0):.1f}%", f"{r.get('total', 0):,.2f}"]
            self.res_tree.insert('', 'end', iid=str(i), values=tuple(base_vals))
        self._update_summary()

    def _add_resource(self):
        d = BudgetItemEditor(self, project_dates=(self.v_ds.get(), self.v_de.get()))
        self.wait_window(d)
        if d.result:
            self.resources.append(d.result)
            self._reload_res_tree()

    def _edit_resource(self):
        sel = self.res_tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите ресурс')
            return
        idx = int(sel[0])
        d = BudgetItemEditor(self, project_dates=(self.v_ds.get(), self.v_de.get()),
                             data=self.resources[idx])
        self.wait_window(d)
        if d.result:
            self.resources[idx] = d.result
            self._reload_res_tree()

    def _del_resource(self):
        sel = self.res_tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите ресурс')
            return
        idx = int(sel[0])
        if messagebox.askyesno('Удаление', 'Удалить ресурс?'):
            self.resources.pop(idx)
            self._reload_res_tree()

    def _update_summary(self):
        cost_sum = sum(r.get('cost', 0) for r in self.resources)
        total_sum = sum(r.get('total', 0) for r in self.resources)
        try:
            ns = float(getattr(self, 'v_tax', tk.StringVar()).get() or 0) / 100
        except Exception:
            ns = 0.0
        if UserSession.can_see_margin():
            final = total_sum * (1 + ns)
            profit = total_sum - cost_sum
            txt = (f'Себестоимость:          {cost_sum:>15,.2f} ₽\n'
                   f'Стоимость с маржой:     {total_sum:>15,.2f} ₽\n'
                   f'Налог ({ns * 100:.0f}%):              {total_sum * ns:>15,.2f} ₽\n'
                   f'Итоговая стоимость:     {final:>15,.2f} ₽\n'
                   f'Чистая прибыль:         {profit:>15,.2f} ₽')
        else:
            txt = (f'Себестоимость:          {cost_sum:>15,.2f} ₽\n'
                   f'Налог на себестоимость ({ns * 100:.0f}%): {cost_sum * ns:>15,.2f} ₽\n'
                   f'(стоимость с маржой и прибыль видны только руководителю)')
        if hasattr(self, 'lbl_summary'):
            self.lbl_summary.config(text=txt, font=TYPEFACE['mono'])

    def _build_bottom(self):
        bf = tk.Frame(self, bg=PALETTE['surface'], pady=10,
                      highlightthickness=1, highlightbackground=PALETTE['border'])
        bf.pack(fill='x', side='bottom')
        inner = tk.Frame(bf, bg=PALETTE['surface'])
        inner.pack(padx=20)
        action_btn(inner, 'Сохранить', self._save, 'primary').pack(side='left', padx=6)
        action_btn(inner, 'НМА', self._export_nma, 'secondary').pack(side='left', padx=6)
        action_btn(inner, 'КП', self._export_kp, 'secondary').pack(side='left', padx=6)
        action_btn(inner, 'Закрыть', self.destroy, 'danger').pack(side='right', padx=6)

    def _save(self):
        name = self.v_name.get().strip()
        if not name:
            messagebox.showwarning('Ошибка', 'Введите название проекта')
            return
        cid = self._cust_map.get(self.v_cust.get())
        conn = connect_db()
        if self.project_id:
            conn.execute("""UPDATE projects SET name=?,date_start=?,date_end=?,description=?,
                            customer_id=?,tax_rate=?,status=?,resources=? WHERE id=?""",
                         (name, self.v_ds.get(), self.v_de.get(),
                          self.txt_desc.get('1.0', 'end').strip(),
                          cid, float(self.v_tax.get() or 0), self.v_status.get(),
                          json.dumps(self.resources, ensure_ascii=False), self.project_id))
        else:
            conn.execute("""INSERT INTO projects(workspace_id,name,date_start,date_end,description,
                            customer_id,tax_rate,status,resources) VALUES(?,?,?,?,?,?,?,?,?)""",
                         (self.workspace_id, name, self.v_ds.get(), self.v_de.get(),
                          self.txt_desc.get('1.0', 'end').strip(),
                          cid, float(self.v_tax.get() or 0), self.v_status.get(),
                          json.dumps(self.resources, ensure_ascii=False)))
        conn.commit()
        conn.close()
        messagebox.showinfo('Готово', 'Проект сохранён')
        if self.on_save:
            self.on_save()

    def _choose_export_format(self, title):
        d = tk.Toplevel(self)
        d.title(title)
        d.configure(bg=PALETTE['bg'])
        d.geometry('300x180')
        d.grab_set()
        result = tk.StringVar(value='')
        tk.Label(d, text='Формат экспорта', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(pady=(20, 10))
        bf = tk.Frame(d, bg=PALETTE['bg'])
        bf.pack()
        for fmt, label, avail in [('pdf', 'PDF', HAS_PDF), ('xlsx', 'Excel', HAS_XLSX), ('docx', 'Word', HAS_DOCX)]:
            state = 'normal' if avail else 'disabled'
            tk.Button(bf, text=label, width=8, bg=PALETTE['surface2'], fg=PALETTE['text'],
                      relief='flat', font=TYPEFACE['normal'], state=state,
                      command=lambda f=fmt: (result.set(f), d.destroy())).pack(side='left', padx=6)
        d.wait_window()
        return result.get()

    def _export_nma(self):
        fmt = self._choose_export_format('Экспорт НМА')
        if not fmt:
            return
        name = self.v_name.get() or 'Проект'
        resources = self.resources
        cost_sum = sum(r.get('cost', 0) for r in resources)
        path = filedialog.asksaveasfilename(
            defaultextension=f'.{fmt}',
            filetypes=[(fmt.upper(), f'*.{fmt}')],
            initialfile=f'НМА_{name}.{fmt}')
        if not path:
            return
        if fmt == 'xlsx':
            export_nma_xlsx(path, name, self.v_ds.get(), self.v_de.get(), cost_sum, resources)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'docx':
            export_nma_docx(path, name, self.v_ds.get(), self.v_de.get(), cost_sum, resources)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'pdf':
            export_nma_pdf(path, name, self.v_ds.get(), self.v_de.get(), cost_sum, resources)

    def _export_kp(self):
        if self.v_cust.get() == '(не выбран)' or not self.v_cust.get():
            messagebox.showwarning('КП', 'Выберите клиента для формирования КП')
            return
        fmt = self._choose_export_format('Экспорт КП')
        if not fmt:
            return
        name = self.v_name.get() or 'Проект'
        cid = self._cust_map.get(self.v_cust.get())
        conn = connect_db()
        cust = dict(conn.execute("SELECT * FROM customers WHERE id=?", (cid,)).fetchone()) if cid else {}
        conn.close()
        total_sum = sum(r.get('total', 0) for r in self.resources)
        try:
            ns = float(self.v_tax.get() or 0) / 100
        except Exception:
            ns = 0.0
        final = total_sum * (1 + ns)
        path = filedialog.asksaveasfilename(
            defaultextension=f'.{fmt}',
            filetypes=[(fmt.upper(), f'*.{fmt}')],
            initialfile=f'КП_{name}.{fmt}')
        if not path:
            return
        company = read_setting('company_name')
        director = read_setting('director_name')
        dir_pos = read_setting('director_position')
        if fmt == 'xlsx':
            export_kp_xlsx(path, name, cust, self.v_ds.get(), self.v_de.get(),
                           final, self.resources, company, director, dir_pos)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'docx':
            export_kp_docx(path, name, cust, self.v_ds.get(), self.v_de.get(),
                           final, self.resources, company, director, dir_pos)
            messagebox.showinfo('Экспорт', f'Файл сохранён:\n{path}')
        elif fmt == 'pdf':
            export_kp_pdf(path, name, cust, self.v_ds.get(), self.v_de.get(),
                          final, self.resources, company, director, dir_pos)


# ---------------------------------------------------------------------------
# Рабочая область
# ---------------------------------------------------------------------------

class WorkspacePanel(ttk.Frame):
    def __init__(self, parent, workspace_id, **kw):
        super().__init__(parent, **kw)
        self.workspace_id = workspace_id
        self.configure(style='TFrame')
        self._build()
        self.load()

    def _build(self):
        conn = connect_db()
        ws = conn.execute("SELECT * FROM workspaces WHERE id=?", (self.workspace_id,)).fetchone()
        conn.close()
        ws_name = ws['name'] if ws else '—'

        _, body = page_card(self, ws_name, 'Проекты в выбранной зоне')
        tb = tk.Frame(body, bg=PALETTE['surface'])
        tb.pack(fill='x', padx=16, pady=(12, 8))
        btns = tk.Frame(tb, bg=PALETTE['surface'])
        btns.pack(side='right')
        action_btn(btns, '+ Проект', self._add_project, 'primary').pack(side='left', padx=3)
        action_btn(btns, 'Открыть', self._open_project, 'secondary').pack(side='left', padx=3)
        action_btn(btns, 'Удалить', self._delete_project, 'danger').pack(side='left', padx=3)
        action_btn(btns, 'Участники', self._manage_members, 'secondary').pack(side='left', padx=3)

        cols = ('id', 'name', 'status', 'date_start', 'date_end', 'cost')
        hdrs = ('#', 'Проект', 'Статус', 'Старт', 'Финиш', 'Сумма')
        widths = (40, 240, 100, 100, 100, 140)
        tf = tk.Frame(body, bg=PALETTE['surface'])
        tf.pack(fill='both', expand=True, padx=16, pady=(0, 16))
        self.tree = ttk.Treeview(tf, columns=cols, show='headings', height=14)
        for c, h, w in zip(cols, hdrs, widths):
            self.tree.heading(c, text=h)
            self.tree.column(c, width=w, minwidth=50)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self.tree.bind('<Double-1>', lambda e: self._open_project())

    def load(self):
        self.tree.delete(*self.tree.get_children())
        conn = connect_db()
        rows = conn.execute("SELECT * FROM projects WHERE workspace_id=? ORDER BY id DESC",
                            (self.workspace_id,)).fetchall()
        conn.close()
        for row in rows:
            resources = json.loads(row['resources'] or '[]')
            total = sum(r.get('total', 0) for r in resources)
            try:
                ns = (row['tax_rate'] or 0) / 100
            except Exception:
                ns = 0
            final = total * (1 + ns)
            self.tree.insert('', 'end', iid=str(row['id']), values=(
                row['id'], row['name'], row['status'],
                row['date_start'] or '', row['date_end'] or '',
                f'{final:,.2f} ₽'
            ))

    def _add_project(self):
        ProjectWindow(self, workspace_id=self.workspace_id, on_save=self.load)

    def _open_project(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите проект')
            return
        ProjectWindow(self, project_id=int(sel[0]), on_save=self.load)

    def _delete_project(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите проект')
            return
        if messagebox.askyesno('Удаление', 'Удалить проект со всеми ресурсами?'):
            conn = connect_db()
            conn.execute("DELETE FROM projects WHERE id=?", (int(sel[0]),))
            conn.commit()
            conn.close()
            self.load()

    def _manage_members(self):
        d = MemberManager(self, self.workspace_id)
        self.wait_window(d)


# ---------------------------------------------------------------------------
# Параметры компании
# ---------------------------------------------------------------------------

class CompanySettings(ttk.Frame):
    TABS = (
        ('company', 'Компания'),
        ('leadership', 'Руководство'),
        ('contacts', 'Контакты'),
        ('branding', 'Логотип'),
    )

    FIELD_GROUPS = {
        'company': [
            ('company_name', 'Наименование организации',
             'Отображается в шапке документов и на главной странице'),
        ],
        'leadership': [
            ('director_name', 'ФИО подписанта',
             'Используется при формировании НМА и коммерческих предложений'),
            ('director_position', 'Должность подписанта',
             'Например: «Генеральный директор»'),
        ],
        'contacts': [
            ('phone', 'Телефон', 'Будет указан в экспортируемых файлах'),
            ('email', 'E-mail', 'Контактный адрес для документов'),
        ],
    }

    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self.configure(style='TFrame')
        self._logo_img = None
        self.vars = {}
        self._pending_logo = read_setting('logo_path')
        self._tab_btns = {}
        self._pages = {}
        self._active_tab = None
        self._build()

    def _build(self):
        outer = tk.Frame(self, bg=PALETTE['bg'])
        outer.pack(fill='both', expand=True, padx=28, pady=24)

        header = tk.Frame(outer, bg=PALETTE['bg'])
        header.pack(fill='x', pady=(0, 20))
        tk.Label(header, text='Настройки организации', bg=PALETTE['bg'],
                 fg=PALETTE['text'], font=TYPEFACE['title']).pack(side='left')
        self._status_lbl = tk.Label(header, text='', bg=PALETTE['bg'],
                                    fg=PALETTE['success'], font=TYPEFACE['small'])
        self._status_lbl.pack(side='right', padx=8)

        shell = panel_box(outer)
        shell.pack(fill='both', expand=True)
        body = tk.Frame(shell, bg=PALETTE['surface'])
        body.pack(fill='both', expand=True)

        rail = tk.Frame(body, bg=PALETTE['surface2'], width=200)
        rail.pack(side='left', fill='y')
        rail.pack_propagate(False)

        tk.Label(rail, text='РАЗДЕЛЫ', bg=PALETTE['surface2'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['section']).pack(anchor='w', padx=16, pady=(18, 8))

        for key, label in self.TABS:
            btn = tk.Button(
                rail, text=f'  {label}', anchor='w', relief='flat', cursor='hand2',
                bg=PALETTE['surface2'], fg=PALETTE['text'],
                activebackground=PALETTE['hover'], activeforeground=PALETTE['primary'],
                font=TYPEFACE['normal'], padx=14, pady=11,
                command=lambda k=key: self._show_tab(k),
            )
            btn.pack(fill='x', padx=8, pady=2)
            self._tab_btns[key] = btn

        tk.Frame(rail, bg=PALETTE['border'], height=1).pack(fill='x', padx=16, pady=(16, 12))
        tk.Label(rail, text='Предпросмотр шапки\nдокумента', bg=PALETTE['surface2'],
                 fg=PALETTE['text_dim'], font=TYPEFACE['small'], justify='left').pack(
            anchor='w', padx=16, pady=(0, 6))
        self._mini_preview = tk.Frame(rail, bg=PALETTE['stat_bg'],
                                      highlightthickness=1, highlightbackground=PALETTE['border'])
        self._mini_preview.pack(fill='x', padx=12, pady=(0, 16), ipady=6)
        self._preview_labels = {}
        for slot in ('name', 'phone', 'email'):
            lbl = tk.Label(self._mini_preview, text='—', bg=PALETTE['stat_bg'],
                           fg=PALETTE['text'], font=TYPEFACE['small'], anchor='w')
            lbl.pack(fill='x', padx=10, pady=2)
            self._preview_labels[slot] = lbl

        self._content = tk.Frame(body, bg=PALETTE['surface'])
        self._content.pack(side='right', fill='both', expand=True, padx=28, pady=24)

        for key, _ in self.TABS:
            page = tk.Frame(self._content, bg=PALETTE['surface'])
            self._pages[key] = page
            if key == 'branding':
                self._build_branding_tab(page)
            else:
                self._build_fields_tab(page, key)

        footer = tk.Frame(shell, bg=PALETTE['surface2'], height=56)
        footer.pack(fill='x', side='bottom')
        footer.pack_propagate(False)
        foot_inner = tk.Frame(footer, bg=PALETTE['surface2'])
        foot_inner.pack(fill='both', expand=True, padx=20, pady=10)
        tk.Label(foot_inner, text='Изменения применяются после сохранения',
                 bg=PALETTE['surface2'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(side='left')
        action_btn(foot_inner, 'Сохранить всё', self._save, 'primary').pack(side='right', padx=(8, 0))
        action_btn(foot_inner, 'Сбросить', self._reload, 'secondary').pack(side='right')

        self._show_tab('company')
        self._update_preview()

    def _build_fields_tab(self, page, group_key):
        fields = self.FIELD_GROUPS.get(group_key, [])
        tab_title = next(l for k, l in self.TABS if k == group_key)

        tk.Label(page, text=tab_title, bg=PALETTE['surface'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(anchor='w', pady=(0, 4))
        tk.Label(page, text='Заполните поля — они попадут в экспорт НМА и КП',
                 bg=PALETTE['surface'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(anchor='w', pady=(0, 20))

        form = tk.Frame(page, bg=PALETTE['surface'])
        form.pack(fill='x')
        form.columnconfigure(1, weight=1)

        for row, (key, label, hint) in enumerate(fields):
            base = row * 2
            lbl_col = tk.Frame(form, bg=PALETTE['surface'])
            lbl_col.grid(row=base, column=0, sticky='nw', padx=(0, 20), pady=(14, 6))
            tk.Label(lbl_col, text=label, bg=PALETTE['surface'], fg=PALETTE['text'],
                     font=TYPEFACE['normal']).pack(anchor='w')
            tk.Label(lbl_col, text=hint, bg=PALETTE['surface'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small'], wraplength=220, justify='left').pack(anchor='w', pady=(3, 0))

            v = tk.StringVar(value=read_setting(key))
            self.vars[key] = v
            v.trace_add('write', lambda *_: self._update_preview())
            styled_entry(form, textvariable=v, width=44).grid(
                row=base, column=1, sticky='ew', pady=(14, 6))
            if row < len(fields) - 1:
                tk.Frame(form, bg=PALETTE['border'], height=1).grid(
                    row=base + 1, column=0, columnspan=2, sticky='ew', pady=(0, 4))

    def _build_branding_tab(self, page):
        tk.Label(page, text='Логотип', bg=PALETTE['surface'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(anchor='w', pady=(0, 4))
        tk.Label(page, text='Изображение размещается в верхней части PDF, Word и Excel',
                 bg=PALETTE['surface'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(anchor='w', pady=(0, 20))

        layout = tk.Frame(page, bg=PALETTE['surface'])
        layout.pack(fill='both', expand=True)

        preview_box = tk.Frame(layout, bg=PALETTE['stat_bg'],
                               highlightthickness=2, highlightbackground=PALETTE['border'],
                               highlightcolor=PALETTE['primary'])
        preview_box.pack(side='left', fill='both', expand=True, padx=(0, 20), ipady=20, ipadx=20)

        tk.Label(preview_box, text='Область предпросмотра', bg=PALETTE['stat_bg'],
                 fg=PALETTE['text_dim'], font=TYPEFACE['small']).pack(anchor='nw', padx=16, pady=(12, 8))

        self.logo_preview = tk.Label(
            preview_box, bg=PALETTE['white'], text='Логотип не загружен',
            fg=PALETTE['text_dim'], font=TYPEFACE['small'],
            width=28, height=8, relief='flat',
            highlightthickness=1, highlightbackground=PALETTE['border'],
        )
        self.logo_preview.pack(padx=16, pady=(0, 12))

        side = tk.Frame(layout, bg=PALETTE['surface'], width=240)
        side.pack(side='right', fill='y')
        side.pack_propagate(False)

        tk.Label(side, text='Файл', bg=PALETTE['surface'], fg=PALETTE['text'],
                 font=TYPEFACE['normal']).pack(anchor='w', pady=(0, 6))
        self._logo_path_lbl = tk.Label(side, text='—', bg=PALETTE['surface'],
                                       fg=PALETTE['text_dim'], font=TYPEFACE['small'],
                                       wraplength=220, justify='left')
        self._logo_path_lbl.pack(anchor='w', pady=(0, 16))

        action_btn(side, 'Загрузить изображение', self._pick_logo, 'primary').pack(
            fill='x', pady=(0, 8))
        action_btn(side, 'Убрать логотип', self._clear_logo, 'danger').pack(fill='x')

        tk.Label(side, text='PNG · JPG · GIF · BMP\nРекомендуемый размер: 400×150 px',
                 bg=PALETTE['surface'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small'], justify='left').pack(anchor='w', pady=(20, 0))
        self._refresh_logo_preview()

    def _show_tab(self, key):
        for page in self._pages.values():
            page.pack_forget()
        self._pages[key].pack(fill='both', expand=True)
        for k, btn in self._tab_btns.items():
            if k == key:
                btn.config(bg=PALETTE['white'], fg=PALETTE['primary'],
                           font=TYPEFACE['heading'])
            else:
                btn.config(bg=PALETTE['surface2'], fg=PALETTE['text'],
                           font=TYPEFACE['normal'])
        self._active_tab = key

    def _update_preview(self):
        name = self.vars.get('company_name', tk.StringVar()).get() or '—'
        phone = self.vars.get('phone', tk.StringVar()).get()
        email = self.vars.get('email', tk.StringVar()).get()
        self._preview_labels['name'].config(text=name[:40])
        self._preview_labels['phone'].config(
            text=f'Тел.: {phone}' if phone else 'Тел.: —')
        self._preview_labels['email'].config(
            text=f'Email: {email}' if email else 'Email: —')

    def _refresh_logo_preview(self):
        path = self._pending_logo
        path_lbl = getattr(self, '_logo_path_lbl', None)
        preview = getattr(self, 'logo_preview', None)
        if not preview:
            return
        if path and os.path.exists(path):
            if path_lbl:
                path_lbl.config(text=os.path.basename(path))
            try:
                from PIL import Image, ImageTk
                img = Image.open(path)
                img.thumbnail((280, 140), Image.LANCZOS)
                self._logo_img = ImageTk.PhotoImage(img)
                preview.config(image=self._logo_img, text='', width=280, height=140,
                               bg=PALETTE['white'])
                return
            except ImportError:
                preview.config(image='', text=os.path.basename(path),
                               width=30, height=3, bg=PALETTE['white'])
                return
        if path_lbl:
            path_lbl.config(text='Файл не выбран')
        preview.config(image='', text='Загрузите логотип компании',
                       width=28, height=8, bg=PALETTE['white'])
        self._logo_img = None

    def _pick_logo(self):
        path = filedialog.askopenfilename(
            title='Выберите логотип',
            filetypes=[('Изображения', '*.png *.jpg *.jpeg *.gif *.bmp'),
                       ('Все файлы', '*.*')])
        if not path:
            return
        self._pending_logo = path
        self._refresh_logo_preview()
        self._flash_status('Логотип выбран — нажмите «Сохранить всё»')

    def _clear_logo(self):
        self._pending_logo = ''
        self._refresh_logo_preview()
        self._flash_status('Логотип убран — нажмите «Сохранить всё»')

    def _reload(self):
        for key in self.vars:
            self.vars[key].set(read_setting(key))
        self._pending_logo = read_setting('logo_path')
        self._refresh_logo_preview()
        self._update_preview()
        self._flash_status('Значения восстановлены из базы')

    def _flash_status(self, text):
        self._status_lbl.config(text=text, fg=PALETTE['accent'])
        self.after(3500, lambda: self._status_lbl.config(text=''))

    def _save(self):
        for key, v in self.vars.items():
            write_setting(key, v.get())
        write_setting('logo_path', self._pending_logo or '')
        self._update_preview()
        self._flash_status('Все настройки сохранены')
        messagebox.showinfo('Готово', 'Параметры организации сохранены')


# ---------------------------------------------------------------------------
# Учётные записи
# ---------------------------------------------------------------------------

class AccountEditor(ModalForm):
    def __init__(self, parent, data=None):
        self.data = data or {}
        super().__init__(parent, 'Учётная запись', width=540, height=620)

    def _build(self):
        tk.Label(self, text='Учётная запись', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['title']).pack(pady=(20, 10))
        sf = ScrollPanel(self)
        sf.pack(fill='both', expand=True)
        f = sf.inner
        f.configure(style='TFrame')

        def row(label, widget, row_n):
            tk.Label(f, text=label, bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                     font=TYPEFACE['small']).grid(row=row_n, column=0, sticky='nw', pady=6, padx=(20, 8))
            widget.grid(row=row_n, column=1, sticky='ew', pady=6, padx=(0, 20))

        self.v_last = tk.StringVar(value=self.data.get('last_name', ''))
        self.v_first = tk.StringVar(value=self.data.get('first_name', ''))
        self.v_mid = tk.StringVar(value=self.data.get('middle_name', ''))
        self.v_email = tk.StringVar(value=self.data.get('email', ''))
        self.v_pos = tk.StringVar(value=self.data.get('position', ''))
        self.v_pwd = tk.StringVar(value=self.data.get('password', 'password'))

        row('Фамилия *', styled_entry(f, textvariable=self.v_last, width=30), 0)
        row('Имя *', styled_entry(f, textvariable=self.v_first, width=30), 1)
        row('Отчество', styled_entry(f, textvariable=self.v_mid, width=30), 2)
        row('Email', styled_entry(f, textvariable=self.v_email, width=30), 3)
        row('Должность', styled_entry(f, textvariable=self.v_pos, width=30), 4)
        row('Пароль', styled_entry(f, textvariable=self.v_pwd, width=30), 5)

        tk.Label(f, text='Роли', bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).grid(row=6, column=0, sticky='nw', pady=6, padx=(20, 8))
        roles_f = tk.Frame(f, bg=PALETTE['bg'])
        roles_f.grid(row=6, column=1, sticky='ew', pady=6, padx=(0, 20))
        cur_roles = (json.loads(self.data.get('roles', '[]'))
                     if isinstance(self.data.get('roles'), str)
                     else (self.data.get('roles') or []))
        self.role_vars = {}
        for role in ALL_ROLES:
            v = tk.BooleanVar(value=role in cur_roles)
            self.role_vars[role] = v
            tk.Checkbutton(roles_f, text=role, variable=v,
                           bg=PALETTE['bg'], fg=PALETTE['text'], selectcolor=PALETTE['surface2'],
                           activebackground=PALETTE['bg'], activeforeground=PALETTE['text'],
                           font=TYPEFACE['small']).pack(anchor='w')

        f.columnconfigure(1, weight=1)
        self._footer(self)

    def _validate(self):
        if not self.v_last.get() or not self.v_first.get():
            messagebox.showwarning('Ошибка', 'Заполните Фамилию и Имя')
            return False
        return True

    def _collect(self):
        roles = [r for r, v in self.role_vars.items() if v.get()]
        return dict(
            last_name=self.v_last.get(), first_name=self.v_first.get(),
            middle_name=self.v_mid.get(), email=self.v_email.get(),
            position=self.v_pos.get(), password=self.v_pwd.get(),
            roles=json.dumps(roles, ensure_ascii=False),
        )


class AccountCatalog(CatalogPanel):
    TITLE = 'Пользователи'
    SUBTITLE = 'Учётные записи и роли доступа'
    COLUMNS = [('id', '#', 40), ('fio', 'ФИО', 220), ('email', 'Email', 180),
               ('position', 'Должность', 140), ('roles_display', 'Роли', 200)]

    def _fetch_rows(self):
        conn = connect_db()
        rows = conn.execute(
            "SELECT *, last_name||' '||first_name||COALESCE(' '||middle_name,'') as fio "
            "FROM users ORDER BY last_name").fetchall()
        conn.close()
        result = []
        for r in rows:
            d = dict(r)
            try:
                roles = json.loads(d.get('roles', '[]'))
                d['roles_display'] = ', '.join(roles) if roles else '—'
            except Exception:
                d['roles_display'] = '—'
            result.append(d)
        return result

    def _add(self):
        d = AccountEditor(self)
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "INSERT INTO users(last_name,first_name,middle_name,email,position,password,roles) "
                "VALUES(?,?,?,?,?,?,?)",
                (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                 d.result['email'], d.result['position'], d.result['password'], d.result['roles']))
            conn.commit()
            conn.close()
            self.load()

    def _edit(self):
        rid = self._selected_id()
        if not rid:
            messagebox.showwarning('Выбор', 'Выберите запись')
            return
        conn = connect_db()
        row = conn.execute("SELECT * FROM users WHERE id=?", (rid,)).fetchone()
        conn.close()
        d = AccountEditor(self, dict(row))
        self.wait_window(d)
        if d.result:
            conn = connect_db()
            conn.execute(
                "UPDATE users SET last_name=?,first_name=?,middle_name=?,email=?,position=?,password=?,roles=? "
                "WHERE id=?",
                (d.result['last_name'], d.result['first_name'], d.result['middle_name'],
                 d.result['email'], d.result['position'], d.result['password'], d.result['roles'], rid))
            conn.commit()
            conn.close()
            self.load()

    def _do_delete(self, rid):
        conn = connect_db()
        conn.execute("DELETE FROM users WHERE id=?", (rid,))
        conn.commit()
        conn.close()


# ---------------------------------------------------------------------------
# Участники рабочей области
# ---------------------------------------------------------------------------

class MemberManager(tk.Toplevel):
    RIGHTS = ['Просмотр', 'Редактирование проектов']

    def __init__(self, parent, workspace_id):
        super().__init__(parent)
        self.workspace_id = workspace_id
        self.title('Участники рабочей области')
        self.configure(bg=PALETTE['bg'])
        self.geometry('640x520')
        self.grab_set()
        self._load()
        self._build()

    def _load(self):
        conn = connect_db()
        ws = conn.execute("SELECT * FROM workspaces WHERE id=?", (self.workspace_id,)).fetchone()
        self.ws = dict(ws) if ws else {}
        try:
            self.members = json.loads(self.ws.get('members', '[]'))
        except Exception:
            self.members = []
        self.all_users = conn.execute(
            "SELECT id, last_name||' '||first_name as fio FROM users ORDER BY last_name").fetchall()
        conn.close()

    def _build(self):
        tk.Label(self, text=f'Участники: {self.ws.get("name", "")}',
                 bg=PALETTE['bg'], fg=PALETTE['text'], font=TYPEFACE['title']).pack(
            pady=(16, 4), padx=20, anchor='w')

        add_f = tk.Frame(self, bg=PALETTE['bg'])
        add_f.pack(fill='x', padx=20, pady=8)
        tk.Label(add_f, text='Добавить:', bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(side='left', padx=(0, 8))
        self.v_add_user = tk.StringVar()
        user_names = [f"{u['fio']} (#{u['id']})" for u in self.all_users]
        self.user_id_map = {f"{u['fio']} (#{u['id']})": u['id'] for u in self.all_users}
        self.cb_user = styled_combo(add_f, user_names, textvariable=self.v_add_user, width=28)
        self.cb_user.pack(side='left', padx=(0, 10))
        self.v_right = tk.StringVar(value='Просмотр')
        styled_combo(add_f, self.RIGHTS, textvariable=self.v_right, width=22).pack(side='left', padx=(0, 10))
        action_btn(add_f, 'Добавить', self._add_member, 'primary').pack(side='left')

        tf = tk.Frame(self, bg=PALETTE['bg'])
        tf.pack(fill='both', expand=True, padx=20, pady=(0, 8))
        self.tree = ttk.Treeview(tf, columns=('uid', 'fio', 'right'), show='headings')
        self.tree.heading('uid', text='#')
        self.tree.column('uid', width=40)
        self.tree.heading('fio', text='ФИО')
        self.tree.column('fio', width=280)
        self.tree.heading('right', text='Права')
        self.tree.column('right', width=200)
        sb = ttk.Scrollbar(tf, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        self._reload_tree()

        bf = tk.Frame(self, bg=PALETTE['bg'])
        bf.pack(fill='x', padx=20, pady=10)
        action_btn(bf, 'Удалить выбранного', self._remove_member, 'danger').pack(side='left', padx=(0, 10))
        action_btn(bf, 'Сохранить', self._save, 'primary').pack(side='left')
        action_btn(bf, 'Закрыть', self.destroy, 'secondary').pack(side='right')

    def _reload_tree(self):
        self.tree.delete(*self.tree.get_children())
        conn = connect_db()
        for m in self.members:
            uid = m.get('user_id')
            right = m.get('right', 'Просмотр')
            u = conn.execute(
                "SELECT last_name||' '||first_name as fio FROM users WHERE id=?", (uid,)).fetchone()
            fio = u['fio'] if u else f'ID={uid}'
            self.tree.insert('', 'end', iid=str(uid), values=(uid, fio, right))
        conn.close()

    def _add_member(self):
        name = self.v_add_user.get()
        uid = self.user_id_map.get(name)
        if not uid:
            messagebox.showwarning('Ошибка', 'Выберите пользователя')
            return
        if any(m['user_id'] == uid for m in self.members):
            messagebox.showwarning('Ошибка', 'Пользователь уже добавлен')
            return
        self.members.append({'user_id': uid, 'right': self.v_right.get()})
        self._reload_tree()

    def _remove_member(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите участника')
            return
        uid = int(sel[0])
        self.members = [m for m in self.members if m['user_id'] != uid]
        self._reload_tree()

    def _save(self):
        conn = connect_db()
        conn.execute("UPDATE workspaces SET members=? WHERE id=?",
                     (json.dumps(self.members, ensure_ascii=False), self.workspace_id))
        conn.commit()
        conn.close()
        messagebox.showinfo('Готово', 'Состав участников сохранён')


# ---------------------------------------------------------------------------
# Окно авторизации
# ---------------------------------------------------------------------------

class AuthWindow(tk.Toplevel):
    MODES = ('login', 'guest', 'info')

    def __init__(self, parent):
        super().__init__(parent)
        self.title(f'{APP_NAME} — авторизация')
        self.configure(bg='#E2E8F0')
        self.geometry('460x520')
        self.resizable(False, False)
        self.result = False
        self._mode = tk.StringVar(value='login')
        self.protocol('WM_DELETE_WINDOW', self._on_close)
        self._build()
        self.grab_set()

    def _build(self):
        card = tk.Frame(self, bg=PALETTE['surface'],
                        highlightthickness=1, highlightbackground=PALETTE['border'])
        card.place(relx=0.5, rely=0.5, anchor='center', width=400, height=460)

        tk.Label(card, text=APP_NAME, bg=PALETTE['surface'], fg=PALETTE['primary'],
                 font=TYPEFACE['brand']).pack(pady=(28, 2))
        tk.Label(card, text=APP_TAGLINE, bg=PALETTE['surface'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(pady=(0, 20))

        mode_bar = tk.Frame(card, bg=PALETTE['surface2'], padx=4, pady=4)
        mode_bar.pack(fill='x', padx=24)
        for val, caption in (('login', 'Вход'), ('guest', 'Гость'), ('info', 'О системе')):
            tk.Radiobutton(
                mode_bar, text=caption, variable=self._mode, value=val,
                indicatoron=False, width=10, padx=8, pady=6,
                bg=PALETTE['surface2'], fg=PALETTE['text'],
                selectcolor=PALETTE['primary'], activebackground=PALETTE['surface2'],
                activeforeground=PALETTE['white'], font=TYPEFACE['small'],
                command=self._switch_mode,
            ).pack(side='left', expand=True, fill='x', padx=2)

        self._mode_body = tk.Frame(card, bg=PALETTE['surface'])
        self._mode_body.pack(fill='both', expand=True, padx=24, pady=(16, 0))

        conn = connect_db()
        rows = conn.execute(
            "SELECT id, last_name||' '||first_name as n, password FROM users ORDER BY last_name"
        ).fetchall()
        conn.close()
        self._user_map = {r['n']: dict(r) for r in rows}

        self._switch_mode()

    def _clear_mode_body(self):
        for w in self._mode_body.winfo_children():
            w.destroy()

    def _switch_mode(self):
        self._clear_mode_body()
        mode = self._mode.get()
        if mode == 'login':
            self._render_login()
        elif mode == 'guest':
            self._render_guest()
        else:
            self._render_info()

    def _render_login(self):
        p = self._mode_body
        tk.Label(p, text='Учётная запись', bg=PALETTE['surface'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(anchor='w', pady=(0, 12))

        tk.Label(p, text='Выберите пользователя', bg=PALETTE['surface'],
                 fg=PALETTE['text_dim'], font=TYPEFACE['small']).pack(anchor='w')
        self.v_user = tk.StringVar()
        names = list(self._user_map.keys())
        cb = styled_combo(p, names, textvariable=self.v_user, width=32)
        cb.pack(fill='x', pady=(4, 12))
        if names:
            cb.current(0)

        tk.Label(p, text='Пароль доступа', bg=PALETTE['surface'],
                 fg=PALETTE['text_dim'], font=TYPEFACE['small']).pack(anchor='w')
        self.v_pwd = tk.StringVar()
        styled_entry(p, textvariable=self.v_pwd, width=32, show='•').pack(fill='x', pady=(4, 8))

        if not names:
            tk.Label(p, text='Пользователи не заданы — перейдите в режим «Гость».',
                     bg=PALETTE['surface'], fg=PALETTE['warning'],
                     font=TYPEFACE['small'], wraplength=320, justify='left').pack(anchor='w', pady=4)

        action_btn(p, 'Продолжить', self._login, 'primary').pack(pady=(16, 0))

    def _render_guest(self):
        p = self._mode_body
        tk.Label(p, text='Работа без входа', bg=PALETTE['surface'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(anchor='w', pady=(0, 12))
        tk.Label(
            p,
            text='Открывает полный доступ к справочникам и проектам.\n'
                 'Наценка (маржа) отображается без ограничений роли.',
            bg=PALETTE['surface'], fg=PALETTE['text_dim'], font=TYPEFACE['normal'],
            justify='left', wraplength=320,
        ).pack(anchor='w', pady=(0, 20))
        action_btn(p, 'Начать работу', self._skip, 'accent').pack(anchor='w')

    def _render_info(self):
        p = self._mode_body
        tk.Label(p, text=APP_NAME, bg=PALETTE['surface'], fg=PALETTE['primary'],
                 font=TYPEFACE['title']).pack(anchor='w', pady=(0, 8))
        tk.Label(
            p,
            text=f'{APP_TAGLINE}.\n\n'
                 'Функции:\n'
                 '  — ведение проектных зон и смет\n'
                 '  — справочники ресурсов и контрагентов\n'
                 '  — формирование НМА и коммерческих предложений\n'
                 '  — разграничение прав по ролям\n\n'
                 'Хранилище: it_cost.db (SQLite)',
            bg=PALETTE['surface'], fg=PALETTE['text_dim'], font=TYPEFACE['normal'],
            justify='left', wraplength=320,
        ).pack(anchor='w')

    def _login(self):
        name = self.v_user.get()
        u = self._user_map.get(name)
        if not u:
            self._skip()
            return
        if (u.get('password') or 'password') != self.v_pwd.get():
            messagebox.showwarning('Ошибка', 'Неверный пароль')
            return
        conn = connect_db()
        full = conn.execute("SELECT * FROM users WHERE id=?", (u['id'],)).fetchone()
        conn.close()
        UserSession.set_user(full)
        self.result = True
        self.destroy()

    def _skip(self):
        UserSession.set_user(None)
        self.result = True
        self.destroy()

    def _on_close(self):
        self.result = False
        self.destroy()


# ---------------------------------------------------------------------------
# Главное приложение
# ---------------------------------------------------------------------------

class TehzadApp(tk.Tk):
    NAV_GROUPS = [
        ('ОБЗОР', [
            ('Сводка', '_show_home'),
            ('Проектные зоны', '_show_workspaces'),
        ]),
        ('СПРАВОЧНИКИ', [
            ('Штат', '_show_staff'),
            ('Внешние исполнители', '_show_freelancers'),
            ('Партнёры', '_show_vendors'),
            ('Техника', '_show_assets'),
            ('Контрагенты', '_show_clients'),
        ]),
        ('АДМИНИСТРИРОВАНИЕ', [
            ('Пользователи', '_show_accounts'),
            ('Настройки', '_show_settings'),
        ]),
    ]

    def __init__(self):
        super().__init__()
        self.title(f'{APP_NAME} — {APP_TAGLINE}')
        self.geometry('1320x820')
        self.minsize(1024, 680)
        self.configure(bg=PALETTE['bg'])
        configure_theme()
        self._nav_btns = {}
        self._active_label = None
        self._current_frame = None
        self._stat_labels = {}
        self._build()
        self._refresh_stats()

    def _build(self):
        body = tk.Frame(self, bg=PALETTE['bg'])
        body.pack(fill='both', expand=True)

        sidebar = tk.Frame(body, bg=PALETTE['sidebar'], width=240)
        sidebar.pack(side='left', fill='y')
        sidebar.pack_propagate(False)

        brand = tk.Frame(sidebar, bg=PALETTE['sidebar'])
        brand.pack(fill='x', padx=16, pady=(20, 16))
        tk.Label(brand, text=APP_NAME, bg=PALETTE['sidebar'], fg=PALETTE['white'],
                 font=TYPEFACE['brand']).pack(anchor='w')
        tk.Label(brand, text=APP_TAGLINE, bg=PALETTE['sidebar'], fg=PALETTE['sidebar_text'],
                 font=TYPEFACE['small'], wraplength=200, justify='left').pack(anchor='w', pady=(4, 0))

        tk.Frame(sidebar, bg='#334155', height=1).pack(fill='x', padx=16, pady=8)

        nav_scroll = tk.Frame(sidebar, bg=PALETTE['sidebar'])
        nav_scroll.pack(fill='both', expand=True)

        for section, items in self.NAV_GROUPS:
            tk.Label(nav_scroll, text=section, bg=PALETTE['sidebar'],
                     fg=PALETTE['sidebar_section'], font=TYPEFACE['section']).pack(
                anchor='w', padx=16, pady=(12, 4))
            for label, method in items:
                btn = ttk.Button(nav_scroll, text=f'  {label}', style='Side.TButton',
                                 command=lambda l=label, m=method: self._navigate(l, m))
                btn.pack(fill='x', padx=8, pady=1)
                self._nav_btns[label] = btn

        stats_box = tk.Frame(sidebar, bg=PALETTE['sidebar_active'],
                             highlightthickness=1, highlightbackground='#475569')
        stats_box.pack(fill='x', padx=12, pady=(8, 12))
        tk.Label(stats_box, text='Показатели', bg=PALETTE['sidebar_active'],
                 fg=PALETTE['sidebar_section'], font=TYPEFACE['section']).pack(
            anchor='w', padx=12, pady=(10, 6))
        for key, caption in (('projects', 'Активных проектов'),
                             ('employees', 'Записей в штате'),
                             ('customers', 'Контрагентов')):
            row = tk.Frame(stats_box, bg=PALETTE['sidebar_active'])
            row.pack(fill='x', padx=12, pady=3)
            val_lbl = tk.Label(row, text='—', bg=PALETTE['sidebar_active'],
                               fg=PALETTE['accent_light'], font=TYPEFACE['heading'])
            val_lbl.pack(side='right')
            tk.Label(row, text=caption, bg=PALETTE['sidebar_active'],
                     fg=PALETTE['sidebar_text'], font=TYPEFACE['small']).pack(side='left')
            self._stat_labels[key] = val_lbl
        tk.Frame(stats_box, height=8, bg=PALETTE['sidebar_active']).pack()

        user_box = tk.Frame(sidebar, bg=PALETTE['sidebar'])
        user_box.pack(fill='x', side='bottom', padx=16, pady=(0, 16))
        tk.Frame(sidebar, bg='#334155', height=1).pack(fill='x', side='bottom', padx=16, pady=(0, 8))
        roles = UserSession.get_roles()
        roles_txt = ', '.join(roles) if roles else 'гость'
        margin_txt = 'наценка доступна' if UserSession.can_see_margin() else 'наценка скрыта'
        tk.Label(user_box, text=UserSession.display_name(), bg=PALETTE['sidebar'],
                 fg=PALETTE['white'], font=TYPEFACE['small']).pack(anchor='w')
        tk.Label(user_box, text=f'{roles_txt} · {margin_txt}', bg=PALETTE['sidebar'],
                 fg=PALETTE['sidebar_section'], font=('Verdana', 8)).pack(anchor='w')

        main = tk.Frame(body, bg=PALETTE['bg'])
        main.pack(side='right', fill='both', expand=True)

        topbar = tk.Frame(main, bg=PALETTE['header'], height=52,
                          highlightthickness=0, highlightbackground=PALETTE['border'])
        topbar.pack(fill='x')
        topbar.pack_propagate(False)
        self._page_title = tk.Label(topbar, text='Сводка', bg=PALETTE['header'],
                                    fg=PALETTE['header_text'], font=TYPEFACE['heading'])
        self._page_title.pack(side='left', padx=24, pady=14)
        self._breadcrumb = tk.Label(topbar, text='', bg=PALETTE['header'],
                                      fg=PALETTE['text_dim'], font=TYPEFACE['small'])
        self._breadcrumb.pack(side='left', pady=14)

        self.content = tk.Frame(main, bg=PALETTE['bg'])
        self.content.pack(fill='both', expand=True)

        self._navigate('Сводка', '_show_home')

    def _refresh_stats(self):
        conn = connect_db()
        counts = {
            'projects': conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0],
            'employees': conn.execute("SELECT COUNT(*) FROM employees").fetchone()[0],
            'customers': conn.execute("SELECT COUNT(*) FROM customers").fetchone()[0],
        }
        conn.close()
        for key, val in counts.items():
            if key in self._stat_labels:
                self._stat_labels[key].config(text=str(val))

    def _activate_nav(self, label):
        if self._active_label and self._active_label in self._nav_btns:
            self._nav_btns[self._active_label].configure(style='Side.TButton')
        btn = self._nav_btns.get(label)
        if btn:
            btn.configure(style='SideActive.TButton')
            self._active_label = label
        self._page_title.config(text=label)

    def _navigate(self, label, method):
        self._activate_nav(label)
        getattr(self, method)()

    def _switch(self, frame):
        if self._current_frame:
            self._current_frame.destroy()
        self._current_frame = frame
        frame.pack(fill='both', expand=True)

    def _show_home(self):
        f = tk.Frame(self.content, bg=PALETTE['bg'])
        company = read_setting('company_name', 'Организация не указана')
        self._breadcrumb.config(text=APP_TAGLINE)

        hero = tk.Frame(f, bg=PALETTE['primary'], height=140)
        hero.pack(fill='x', padx=24, pady=(20, 0))
        hero.pack_propagate(False)
        hi = tk.Frame(hero, bg=PALETTE['primary'])
        hi.pack(fill='both', expand=True, padx=28, pady=22)
        tk.Label(hi, text=f'Добрый день, {UserSession.display_name()}',
                 bg=PALETTE['primary'], fg='#BFDBFE', font=TYPEFACE['normal']).pack(anchor='w')
        tk.Label(hi, text=company, bg=PALETTE['primary'], fg=PALETTE['white'],
                 font=('Verdana', 26, 'bold')).pack(anchor='w', pady=(4, 6))
        tk.Label(hi, text='Управляйте сметами, ресурсами и документами из одного окна',
                 bg=PALETTE['primary'], fg='#DBEAFE', font=TYPEFACE['small']).pack(anchor='w')

        body = tk.Frame(f, bg=PALETTE['bg'])
        body.pack(fill='both', expand=True, padx=24, pady=24)

        tk.Label(body, text='Перейти к разделам', bg=PALETTE['bg'], fg=PALETTE['text'],
                 font=TYPEFACE['heading']).pack(anchor='w', pady=(0, 12))
        tiles = tk.Frame(body, bg=PALETTE['bg'])
        tiles.pack(anchor='w', fill='x')
        nav_tile(tiles, 'Проектные зоны', 'Создание и ведение проектов',
                 lambda: self._navigate('Проектные зоны', '_show_workspaces'))
        nav_tile(tiles, 'Штат', 'Сотрудники и оклады',
                 lambda: self._navigate('Штат', '_show_staff'))
        nav_tile(tiles, 'Контрагенты', 'Клиенты для КП',
                 lambda: self._navigate('Контрагенты', '_show_clients'))
        nav_tile(tiles, 'Настройки', 'Реквизиты и логотип',
                 lambda: self._navigate('Настройки', '_show_settings'))

        tk.Label(body, text='Подсказка', bg=PALETTE['bg'], fg=PALETTE['text_dim'],
                 font=TYPEFACE['small']).pack(anchor='w', pady=(28, 6))
        tip = panel_box(body)
        tip.pack(fill='x')
        tk.Label(tip, text='Создайте зону → добавьте проект → заполните смету → экспортируйте НМА или КП',
                 bg=PALETTE['surface'], fg=PALETTE['text_dim'], font=TYPEFACE['normal'],
                 padx=16, pady=14).pack(anchor='w')

        self._refresh_stats()
        self._switch(f)

    def _show_workspaces(self):
        f = tk.Frame(self.content, bg=PALETTE['bg'])
        self._breadcrumb.config(text='Управление проектными зонами и проектами')

        split = tk.Frame(f, bg=PALETTE['bg'])
        split.pack(fill='both', expand=True, padx=20, pady=16)

        left = tk.Frame(split, bg=PALETTE['sidebar'], width=220)
        left.pack(side='left', fill='y')
        left.pack_propagate(False)

        tk.Label(left, text='Зоны', bg=PALETTE['sidebar'], fg=PALETTE['sidebar_section'],
                 font=TYPEFACE['section']).pack(anchor='w', padx=14, pady=(14, 6))

        self._ws_list = tk.Listbox(
            left, bg=PALETTE['sidebar'], fg=PALETTE['sidebar_text'],
            selectbackground=PALETTE['primary'], selectforeground=PALETTE['white'],
            font=TYPEFACE['normal'], relief='flat', highlightthickness=0,
            activestyle='none', bd=0,
        )
        self._ws_list.pack(fill='both', expand=True, padx=8, pady=(0, 8))
        self._ws_list.bind('<<ListboxSelect>>', lambda e: self._on_ws_select())

        action_btn(left, 'Новая зона',
                   lambda: self._add_workspace_from_list(), 'accent').pack(
            fill='x', padx=10, pady=(0, 12))

        self._ws_panel = tk.Frame(split, bg=PALETTE['bg'])
        self._ws_panel.pack(side='right', fill='both', expand=True, padx=(12, 0))

        self._ws_data = []
        self._ws_nb = None
        self._ws_parent = f
        self._refresh_ws_list()
        self._switch(f)

    def _refresh_ws_list(self):
        self._ws_list.delete(0, 'end')
        self._ws_data = [('__mgmt__', 'Управление зонами')]
        self._ws_list.insert('end', '  Управление зонами')
        conn = connect_db()
        for ws in conn.execute("SELECT id, name FROM workspaces ORDER BY name").fetchall():
            self._ws_data.append((ws['id'], ws['name']))
            self._ws_list.insert('end', f'  {ws["name"]}')
        conn.close()
        self._ws_list.selection_set(0)
        self._on_ws_select()

    def _on_ws_select(self):
        sel = self._ws_list.curselection()
        if not sel:
            return
        idx = sel[0]
        for w in self._ws_panel.winfo_children():
            w.destroy()
        key, name = self._ws_data[idx]
        if key == '__mgmt__':
            self._build_ws_mgmt_panel(self._ws_panel)
        else:
            panel = WorkspacePanel(self._ws_panel, workspace_id=key)
            panel.pack(fill='both', expand=True)

    def _build_ws_mgmt_panel(self, parent):
        _, body = page_card(parent, 'Управление зонами', 'Список всех проектных зон')
        tb = tk.Frame(body, bg=PALETTE['surface'])
        tb.pack(fill='x', padx=16, pady=(12, 8))
        btns = tk.Frame(tb, bg=PALETTE['surface'])
        btns.pack(side='right')
        action_btn(btns, '+ Зона', lambda: self._add_workspace_from_list(), 'primary').pack(
            side='left', padx=3)
        action_btn(btns, 'Изменить', lambda: self._edit_workspace_from_list(tree), 'secondary').pack(
            side='left', padx=3)
        action_btn(btns, 'Удалить', lambda: self._del_workspace_from_list(tree), 'danger').pack(
            side='left', padx=3)

        tf = tk.Frame(body, bg=PALETTE['surface'])
        tf.pack(fill='both', expand=True, padx=16, pady=(0, 16))
        tree = ttk.Treeview(tf, columns=('id', 'name', 'subdomain'), show='headings', height=14)
        tree.heading('id', text='#')
        tree.column('id', width=40)
        tree.heading('name', text='Название')
        tree.column('name', width=220)
        tree.heading('subdomain', text='Код')
        tree.column('subdomain', width=160)
        sb = ttk.Scrollbar(tf, orient='vertical', command=tree.yview)
        tree.configure(yscrollcommand=sb.set)
        tree.pack(side='left', fill='both', expand=True)
        sb.pack(side='right', fill='y')
        conn = connect_db()
        for ws in conn.execute("SELECT * FROM workspaces ORDER BY name").fetchall():
            tree.insert('', 'end', iid=str(ws['id']),
                        values=(ws['id'], ws['name'], ws['subdomain'] or ''))
        conn.close()

    def _add_workspace_from_list(self):
        name = simpledialog.askstring('Новая зона', 'Название проектной зоны:', parent=self)
        if not name:
            return
        sub = simpledialog.askstring('Код зоны', 'Краткий код (необязательно):', parent=self)
        conn = connect_db()
        conn.execute("INSERT INTO workspaces(name,subdomain) VALUES(?,?)", (name, sub or ''))
        conn.commit()
        conn.close()
        self._refresh_ws_list()

    def _edit_workspace_from_list(self, tree):
        sel = tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите зону')
            return
        wid = int(sel[0])
        conn = connect_db()
        ws = conn.execute("SELECT * FROM workspaces WHERE id=?", (wid,)).fetchone()
        conn.close()
        name = simpledialog.askstring('Изменить', 'Название:', initialvalue=ws['name'], parent=self)
        if not name:
            return
        sub = simpledialog.askstring('Код', 'Код зоны:', initialvalue=ws['subdomain'] or '', parent=self)
        conn = connect_db()
        conn.execute("UPDATE workspaces SET name=?,subdomain=? WHERE id=?", (name, sub or '', wid))
        conn.commit()
        conn.close()
        self._refresh_ws_list()

    def _del_workspace_from_list(self, tree):
        sel = tree.selection()
        if not sel:
            messagebox.showwarning('Выбор', 'Выберите зону')
            return
        if messagebox.askyesno('Удаление', 'Удалить зону и все проекты в ней?'):
            wid = int(sel[0])
            conn = connect_db()
            conn.execute("DELETE FROM projects WHERE workspace_id=?", (wid,))
            conn.execute("DELETE FROM workspaces WHERE id=?", (wid,))
            conn.commit()
            conn.close()
            self._refresh_ws_list()

    def _show_staff(self):
        self._breadcrumb.config(text='Справочник / Штат')
        self._switch(StaffCatalog(self.content))

    def _show_freelancers(self):
        self._breadcrumb.config(text='Справочник / Внешние исполнители')
        self._switch(FreelancerCatalog(self.content))

    def _show_vendors(self):
        self._breadcrumb.config(text='Справочник / Партнёры')
        self._switch(VendorCatalog(self.content))

    def _show_assets(self):
        self._breadcrumb.config(text='Справочник / Техника')
        self._switch(AssetCatalog(self.content))

    def _show_clients(self):
        self._breadcrumb.config(text='Справочник / Контрагенты')
        self._switch(ClientCatalog(self.content))

    def _show_accounts(self):
        self._breadcrumb.config(text='Администрирование / Пользователи')
        self._switch(AccountCatalog(self.content))

    def _show_settings(self):
        self._breadcrumb.config(text='Администрирование / Настройки')
        self._switch(CompanySettings(self.content))


# ---------------------------------------------------------------------------
# Экспорт XLSX
# ---------------------------------------------------------------------------

def export_nma_xlsx(path, project_name, ds, de, cost_sum, resources):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'НМА'
    hdr_fill = PatternFill(fgColor='2563EB', fill_type='solid')
    hdr_font = Font(color='FFFFFF', bold=True)
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    company = read_setting('company_name')
    director = read_setting('director_name')
    dir_pos = read_setting('director_position')
    phone = read_setting('phone')
    email_val = read_setting('email')
    logo_path = read_setting('logo_path')

    start_row = 1
    if logo_path and os.path.exists(logo_path):
        try:
            from openpyxl.drawing.image import Image as XLImage
            img = XLImage(logo_path)
            img.height = 60
            img.width = 160
            ws.add_image(img, 'A1')
            ws.row_dimensions[1].height = 50
            start_row = 4
        except Exception:
            pass

    r = start_row
    ws.cell(r, 1, 'ФОРМА СТОИМОСТИ НЕМАТЕРИАЛЬНОГО АКТИВА').font = Font(bold=True, size=14)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
    r += 1
    ws.cell(r, 1, company)
    r += 1
    if phone:
        ws.cell(r, 1, f'Тел.: {phone}')
        r += 1
    if email_val:
        ws.cell(r, 1, f'Email: {email_val}')
        r += 1
    ws.cell(r, 1, f'Проект: {project_name}')
    r += 1
    ws.cell(r, 1, f'Период: {ds} — {de}')
    r += 1
    c = ws.cell(r, 1, f'Итоговая стоимость НМА: {cost_sum:,.2f} ₽')
    c.font = Font(bold=True)
    r += 1
    ws.append([])
    r += 1

    headers = ['Название', 'Тип', 'Исполнитель', 'Единиц', 'Себестоимость']
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(r, ci, h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.border = border
    r += 1

    for res in resources:
        for ci, val in enumerate([res.get('name'), res.get('resource_type'),
                                  res.get('executor_display'), res.get('units'),
                                  res.get('cost', 0)], 1):
            cell = ws.cell(r, ci, val)
            cell.border = border
        r += 1

    for ci, val in enumerate(['', '', '', 'ИТОГО', cost_sum], 1):
        cell = ws.cell(r, ci, val)
        cell.font = Font(bold=True)
    r += 2

    ws.cell(r, 1, f'{dir_pos}  _______________  {director}')

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 16
    ws.column_dimensions['C'].width = 24
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 18
    wb.save(path)


def export_kp_xlsx(path, project_name, cust, ds, de, final, resources, company, director, dir_pos):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'КП'
    hdr_fill = PatternFill(fgColor='2563EB', fill_type='solid')
    hdr_font = Font(color='FFFFFF', bold=True)
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    phone_val = read_setting('phone')
    email_val = read_setting('email')
    logo_path = read_setting('logo_path')

    start_row = 1
    if logo_path and os.path.exists(logo_path):
        try:
            from openpyxl.drawing.image import Image as XLImage
            img = XLImage(logo_path)
            img.height = 60
            img.width = 160
            ws.add_image(img, 'A1')
            ws.row_dimensions[1].height = 50
            start_row = 4
        except Exception:
            pass

    r = start_row
    ws.cell(r, 1, 'КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ').font = Font(bold=True, size=16)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    r += 1
    ws.cell(r, 1, company)
    r += 1
    if phone_val:
        ws.cell(r, 1, f'Тел.: {phone_val}')
        r += 1
    if email_val:
        ws.cell(r, 1, f'Email: {email_val}')
        r += 1
    ws.cell(r, 1, f'Заказчик: {cust.get("name", "")} | {cust.get("director_name", "")}')
    r += 1
    ws.cell(r, 1, f'Email заказчика: {cust.get("email", "")}')
    r += 1
    ws.cell(r, 1, f'Проект: {project_name}')
    r += 1
    ws.cell(r, 1, f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    r += 1
    c = ws.cell(r, 1, f'Итоговая стоимость: {final:,.2f} ₽')
    c.font = Font(bold=True, size=12)
    r += 1
    ws.append([])
    r += 1

    headers = ['Название услуги', 'Кол-во', 'Ед.', 'Начало', 'Окончание', 'Стоимость (с маржой)']
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(r, ci, h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.border = border
    r += 1

    for res in resources:
        service_name = res.get('service') or res.get('name')
        for ci, val in enumerate([service_name, res.get('units'), res.get('unit_label', ''),
                                  res.get('date_start'), res.get('date_end'), res.get('total', 0)], 1):
            ws.cell(r, ci, val).border = border
        r += 1

    for ci, val in enumerate(['', '', '', '', 'ИТОГО', final], 1):
        ws.cell(r, ci, val).font = Font(bold=True)
    r += 2

    ws.cell(r, 1, f'{dir_pos}  _______________  {director}')

    for col, w in zip('ABCDEF', [30, 10, 10, 12, 12, 20]):
        ws.column_dimensions[col].width = w
    wb.save(path)


# ---------------------------------------------------------------------------
# Экспорт DOCX
# ---------------------------------------------------------------------------

def export_nma_docx(path, project_name, ds, de, cost_sum, resources):
    doc = Document()

    company = read_setting('company_name')
    director = read_setting('director_name')
    dir_pos = read_setting('director_position')
    phone_val = read_setting('phone')
    email_val = read_setting('email')
    logo_path = read_setting('logo_path')

    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Cm(5))
        except Exception:
            pass

    doc.add_heading('ФОРМА СТОИМОСТИ НЕМАТЕРИАЛЬНОГО АКТИВА', 0)
    doc.add_paragraph(company)
    if phone_val:
        doc.add_paragraph(f'Тел.: {phone_val}')
    if email_val:
        doc.add_paragraph(f'Email: {email_val}')
    doc.add_paragraph()
    doc.add_paragraph(f'Проект: {project_name}')
    doc.add_paragraph(f'Период: {ds} — {de}')
    p = doc.add_paragraph()
    p.add_run(f'Итоговая стоимость НМА: {cost_sum:,.2f} ₽').bold = True
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Название', 'Тип', 'Исполнитель', 'Единиц', 'Себестоимость']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for r in resources:
        row = table.add_row().cells
        row[0].text = str(r.get('name', ''))
        row[1].text = str(r.get('resource_type', ''))
        row[2].text = str(r.get('executor_display', ''))
        row[3].text = str(r.get('units', ''))
        row[4].text = f"{r.get('cost', 0):,.2f}"

    total_row = table.add_row().cells
    total_row[3].text = 'ИТОГО'
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[4].text = f'{cost_sum:,.2f} ₽'
    total_row[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f'{dir_pos}  _______________  {director}')
    doc.save(path)


def export_kp_docx(path, project_name, cust, ds, de, final, resources, company, director, dir_pos):
    doc = Document()

    phone_val = read_setting('phone')
    email_val = read_setting('email')
    logo_path = read_setting('logo_path')

    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Cm(5))
        except Exception:
            pass

    doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', 0)
    doc.add_paragraph(company)
    if phone_val:
        doc.add_paragraph(f'Тел.: {phone_val}')
    if email_val:
        doc.add_paragraph(f'Email: {email_val}')
    doc.add_paragraph()
    doc.add_paragraph(f'Заказчик: {cust.get("name", "")} | {cust.get("director_name", "")}')
    doc.add_paragraph(f'Email заказчика: {cust.get("email", "")}')
    doc.add_paragraph(f'Проект: {project_name}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    p = doc.add_paragraph()
    p.add_run(f'Итоговая стоимость: {final:,.2f} ₽').bold = True
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Услуга', 'Кол-во', 'Ед.', 'Начало', 'Окончание', 'Стоимость']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for r in resources:
        row = table.add_row().cells
        row[0].text = str(r.get('service') or r.get('name', ''))
        row[1].text = str(r.get('units', ''))
        row[2].text = str(r.get('unit_label', ''))
        row[3].text = str(r.get('date_start', ''))
        row[4].text = str(r.get('date_end', ''))
        row[5].text = f"{r.get('total', 0):,.2f}"

    t_row = table.add_row().cells
    t_row[4].text = 'ИТОГО'
    t_row[4].paragraphs[0].runs[0].bold = True
    t_row[5].text = f'{final:,.2f} ₽'
    t_row[5].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f'{dir_pos}  _______________  {director}')
    doc.save(path)


# ---------------------------------------------------------------------------
# Экспорт PDF
# ---------------------------------------------------------------------------

def register_cyrillic_font():
    try:
        pdfmetrics.getFont('Arial')
        return True
    except Exception:
        pass
    try:
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/Arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            os.path.expanduser("~/Library/Fonts/Arial.ttf"),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'arial.ttf'),
            "arial.ttf",
        ]
        try:
            import importlib.util
            spec = importlib.util.find_spec('matplotlib')
            if spec:
                mpl_dir = os.path.dirname(spec.origin)
                font_paths.insert(0, os.path.join(mpl_dir, 'mpl-data', 'fonts', 'ttf', 'DejaVuSans.ttf'))
        except Exception:
            pass
        for path in font_paths:
            if path and os.path.exists(path):
                pdfmetrics.registerFont(TTFont('Arial', path))
                pdfmetrics.registerFontFamily('Arial', normal='Arial')
                return True
        return False
    except Exception:
        return False


def export_nma_pdf(path, project_name, ds, de, cost_sum, resources):
    if not HAS_PDF:
        messagebox.showerror("Ошибка", "ReportLab не установлен")
        return

    register_cyrillic_font()

    company = read_setting('company_name')
    director = read_setting('director_name')
    dir_pos = read_setting('director_position')
    phone_val = read_setting('phone')
    email_val = read_setting('email')
    logo_path = read_setting('logo_path')

    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    try:
        styles.add(ParagraphStyle(name='Russian', fontName='Arial', fontSize=11, leading=14))
        styles.add(ParagraphStyle(name='RussianBold', fontName='Arial', fontSize=12, leading=14, alignment=1))
        styles.add(ParagraphStyle(name='RussianSm', fontName='Arial', fontSize=9, leading=12, textColor=colors.grey))
    except Exception:
        pass

    story = []

    if logo_path and os.path.exists(logo_path):
        try:
            from reportlab.platypus import Image as RLImage
            logo = RLImage(logo_path, width=4 * cm, height=2 * cm, kind='proportional')
            story.append(logo)
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.append(Paragraph("ФОРМА СТОИМОСТИ НЕМАТЕРИАЛЬНОГО АКТИВА", styles['RussianBold']))
    story.append(Spacer(1, 8))
    story.append(Paragraph(company, styles['Russian']))
    if phone_val:
        story.append(Paragraph(f"Тел.: {phone_val}", styles.get('RussianSm', styles['Russian'])))
    if email_val:
        story.append(Paragraph(f"Email: {email_val}", styles.get('RussianSm', styles['Russian'])))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Проект:</b> {project_name}", styles['Russian']))
    story.append(Paragraph(f"<b>Период:</b> {ds} — {de}", styles['Russian']))
    story.append(Paragraph(f"<b>Итоговая стоимость НМА:</b> {cost_sum:,.2f} руб.", styles['RussianBold']))
    story.append(Spacer(1, 20))

    data = [['Название', 'Тип', 'Исполнитель', 'Единиц', 'Себестоимость']]
    for r in resources:
        data.append([
            r.get('name', ''),
            r.get('resource_type', ''),
            r.get('executor_display', ''),
            str(r.get('units', '')),
            f"{r.get('cost', 0):,.2f}",
        ])
    data.append(['', '', '', 'ИТОГО', f"{cost_sum:,.2f}"])

    table = Table(data, colWidths=[200, 80, 140, 60, 90])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563EB')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('FONTNAME', (0, -1), (-1, -1), 'Arial'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F5F7FA')]),
    ]))
    story.append(table)
    story.append(Spacer(1, 40))
    story.append(Paragraph(f"{dir_pos}  _______________  {director}", styles['Russian']))
    doc.build(story)
    messagebox.showinfo("Успешно", f"PDF сохранён:\n{path}")


def export_kp_pdf(path, project_name, cust, ds, de, final, resources, company, director, dir_pos):
    if not HAS_PDF:
        messagebox.showerror("Ошибка", "ReportLab не установлен")
        return

    register_cyrillic_font()

    phone_val = read_setting('phone')
    email_val = read_setting('email')
    logo_path = read_setting('logo_path')

    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    try:
        styles.add(ParagraphStyle(name='Russian', fontName='Arial', fontSize=11, leading=14))
        styles.add(ParagraphStyle(name='RussianBold', fontName='Arial', fontSize=12, leading=14, alignment=1))
        styles.add(ParagraphStyle(name='RussianSm', fontName='Arial', fontSize=9, leading=12, textColor=colors.grey))
    except Exception:
        pass

    story = []

    if logo_path and os.path.exists(logo_path):
        try:
            from reportlab.platypus import Image as RLImage
            logo = RLImage(logo_path, width=4 * cm, height=2 * cm, kind='proportional')
            story.append(logo)
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.append(Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", styles['RussianBold']))
    story.append(Spacer(1, 8))
    story.append(Paragraph(company, styles['Russian']))
    if phone_val:
        story.append(Paragraph(f"Тел.: {phone_val}", styles.get('RussianSm', styles['Russian'])))
    if email_val:
        story.append(Paragraph(f"Email: {email_val}", styles.get('RussianSm', styles['Russian'])))
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        f"<b>Заказчик:</b> {cust.get('name', '')} | {cust.get('director_name', '')}", styles['Russian']))
    story.append(Paragraph(f"<b>Проект:</b> {project_name}", styles['Russian']))
    story.append(Paragraph(f"<b>Дата:</b> {datetime.now().strftime('%d.%m.%Y')}", styles['Russian']))
    story.append(Paragraph(f"<b>Итоговая стоимость:</b> {final:,.2f} руб.", styles['RussianBold']))
    story.append(Spacer(1, 20))

    data = [['Услуга', 'Кол-во', 'Ед.', 'Начало', 'Окончание', 'Стоимость']]
    for r in resources:
        data.append([
            r.get('service') or r.get('name', ''),
            str(r.get('units', '')),
            r.get('unit_label', ''),
            r.get('date_start', ''),
            r.get('date_end', ''),
            f"{r.get('total', 0):,.2f}",
        ])
    data.append(['', '', '', '', 'ИТОГО', f"{final:,.2f}"])

    table = Table(data, colWidths=[160, 110, 50, 65, 65, 90])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563EB')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F5F7FA')]),
    ]))
    story.append(table)
    story.append(Spacer(1, 40))
    story.append(Paragraph(f"{dir_pos} ____________________ {director}", styles['Russian']))
    doc.build(story)
    messagebox.showinfo("Успешно", f"PDF сохранён:\n{path}")


# ---------------------------------------------------------------------------
# Точка входа
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    init_db()
    _root = tk.Tk()
    _root.withdraw()
    configure_theme()
    _auth = AuthWindow(_root)
    _root.wait_window(_auth)
    _root.destroy()
    if _auth.result:
        app = TehzadApp()
        app.mainloop()
